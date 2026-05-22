from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse, HttpResponseForbidden, FileResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.utils.decorators import method_decorator
from django.db.models import Q, Count, Sum, Avg, Prefetch
from django.utils import timezone
from django.core.cache import cache
from django.core.mail import send_mail
from django.conf import settings
import random
import string
from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import JSONParser, MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework import status
import json
import uuid
import math
import os
import hashlib
import hmac
import zipfile
import tempfile
import re
from datetime import datetime, date, time, timedelta

# Import models & serializers using parent directory relative path
from ..models import (
    Employee, EmployeeProfile, OfficeLocation, DepartmentOfficeAccess,
    AttendanceRecord, EmployeeRequest, EmployeeDocument, Task, BirthdayWish, TaskComment, TaskStep, TaskAttachment, Team,
    TemporaryTag, TrainingLog, AvatarAsset, Memoji, Notification, TaskHistory, Project, Holiday, HolidayUpload, UserHoliday
)
from ..serializers import AvatarAssetSerializer, MemojiSerializer
from ..security import require_valid_token, require_gated_token_api
from django.contrib.auth.hashers import make_password, check_password
from .utils import get_current_user

# --- Function: _parse_holiday_pdf ---
def _parse_holiday_pdf(file_bytes):
    """Parse a PDF file and extract holiday rows using pdfplumber."""
    try:
        import pdfplumber
        import io
        import re
        rows = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            # Try to find year in text
            full_text = ""
            for page in pdf.pages:
                full_text += (page.extract_text() or "") + "\n"
            year_match = re.search(r'\b(20\d{2})\b', full_text)
            doc_year = int(year_match.group(1)) if year_match else None

            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    if not table: continue
                    header = [str(c).lower().strip() if c else '' for c in table[0]]
                    for row in table[1:]:
                        if row and any(cell for cell in row):
                            rows.append({'raw': row, 'header': header, 'doc_year': doc_year})
        return rows
    except Exception as e:
        raise ValueError(f"PDF parsing failed: {e}")




# --- Function: _parse_holiday_docx ---
def _parse_holiday_docx(file_bytes):
    """Parse a DOCX file and extract holiday rows with correct paragraph context."""
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        import io, re
        
        doc = Document(io.BytesIO(file_bytes))
        rows = []
        
        # Capture year from full text
        full_text = "\n".join([p.text for p in doc.paragraphs])
        year_match = re.search(r'\b(20\d{2})\b', full_text)
        doc_year = int(year_match.group(1)) if year_match else None

        # Helper to iterate elements in order
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P

        def iter_block_items(parent):
            # docx.Document is a factory; the actual class is docx.document.Document
            if hasattr(parent, 'element') and hasattr(parent.element, 'body'):
                parent_elm = parent.element.body
            else:
                parent_elm = parent._element
            
            for child in parent_elm.iterchildren():
                tag = child.tag.lower()
                if tag.endswith('}p'):
                    yield Paragraph(child, parent)
                elif tag.endswith('}tbl'):
                    yield Table(child, parent)

        current_context = ""
        for item in iter_block_items(doc):
            if isinstance(item, Paragraph):
                txt = item.text.strip()
                if txt and len(txt) < 200:
                    current_context = txt
            elif isinstance(item, Table):
                if not item.rows: continue
                is_optional_table = "optional" in current_context.lower()
                header = [cell.text.lower().strip() for cell in item.rows[0].cells]
                for row in item.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        rows.append({
                            'raw': cells, 
                            'header': header, 
                            'doc_year': doc_year,
                            'is_optional_context': is_optional_table,
                            'context_text': current_context
                        })

        if not rows:
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    rows.append({'raw': t.split(), 'header': None, 'doc_year': doc_year})
        return rows
    except Exception as e:
        raise ValueError(f"DOCX parsing failed: {e}")




# --- Function: _parse_holiday_excel ---
def _parse_holiday_excel(file_bytes, file_ext):
    """Parse an Excel/CSV file using pandas."""
    try:
        import pandas as pd
        import io
        import re
        if file_ext in ['.xlsx', '.xls']:
            df = pd.read_excel(io.BytesIO(file_bytes), dtype=str)
        else:
            df = pd.read_csv(io.BytesIO(file_bytes), dtype=str)
        df.columns = [str(c).lower().strip() for c in df.columns]
        
        all_text = " ".join(df.astype(str).values.flatten())
        year_match = re.search(r'\b(20\d{2})\b', all_text)
        doc_year = int(year_match.group(1)) if year_match else None
        
        rows = []
        for _, row in df.iterrows():
            rows.append({'raw': list(row.values), 'header': list(df.columns), 'doc_year': doc_year})
        return rows
    except Exception as e:
        raise ValueError(f"Excel/CSV parsing failed: {e}")




# --- Function: _parse_holiday_txt ---
def _parse_holiday_txt(file_bytes):
    """Parse a plain-text file into rows."""
    try:
        text = file_bytes.decode('utf-8', errors='replace')
        rows = []
        for line in text.split('\n'):
            line = line.strip()
            if line and not line.startswith('#'):
                parts = [p.strip() for p in line.replace('\t', ',').split(',')]
                rows.append({'raw': parts, 'header': None})
        return rows
    except Exception as e:
        raise ValueError(f"TXT parsing failed: {e}")




# --- Function: _normalize_rows ---
def _normalize_rows(raw_rows):
    """Normalise parsed rows into Holiday objects."""
    import re
    from datetime import datetime as dt
    from django.utils import timezone

    current_year = timezone.now().year

    def clean_date_str(s):
        if not s: return ""
        # Remove ordinal suffixes (1st -> 1, 22nd -> 22)
        s = re.sub(r'(\d+)(st|nd|rd|th)', r'\1', str(s), flags=re.IGNORECASE)
        # Remove special characters like dots or commas
        s = re.sub(r'[,.\(\)]', ' ', s)
        return s.strip()

    def smart_parse_date(s, year=None, month_hint=None):
        s = clean_date_str(s)
        if not s or s.lower() in ('nan', 'none', ''):
            return None
        
        # Heuristic: If it's just a bare number (1-31), it's likely a day within a month
        is_bare_number = re.fullmatch(r'\d{1,2}', s)
        if is_bare_number:
            if not month_hint: return None
            s = f"{s} {month_hint}"

        has_year = re.search(r'\d{4}', s)
        has_digits = re.search(r'\d+', s)
        has_month_name = any(m in s.lower() for m in ['jan','feb','mar','apr','may','jun','jul','aug','sep','oct','nov','dec'])
        
        # Prefer strings that actually look like dates (have digits and month or are complex)
        if not has_digits: return None

        formats = [
            '%d %B %Y', '%d %b %Y', '%B %d %Y', '%b %d %Y',
            '%d-%m-%Y', '%d/%m/%Y', '%Y-%m-%d',
            '%d %m %Y', '%Y %m %d',
            '%d %B', '%B %d', '%b %d', '%d %b'
        ]
        
        target_year = year or current_year
        valid_date = None

        for fmt in formats:
            try:
                if '%' not in fmt or 'Y' not in fmt:
                    d = dt.strptime(f"{s} {target_year}", f"{fmt} %Y").date()
                else:
                    d = dt.strptime(s, fmt).date()
                valid_date = d
                break
            except ValueError:
                continue

        if not valid_date:
            try:
                from dateutil import parser as dutil
                default_dt = dt(target_year, 1, 1)
                res = dutil.parse(s, default=default_dt, dayfirst=True).date()
                if not has_digits: return None
                valid_date = res
            except Exception:
                pass
        
        if valid_date and year and valid_date.year != year:
            valid_date = valid_date.replace(year=year)
            
        return valid_date

    holidays = []
    seen_keys = set()
    MONTH_NAMES = ['january','february','march','april','may','june','july','august','september','october','november','december']

    for item in raw_rows:
        raw_data = item.get('raw', [])
        header = item.get('header') or []
        doc_year = item.get('doc_year') or current_year
        # Handle different item keys for is_optional
        is_optional_context = item.get('is_optional', item.get('is_optional_context', False))

        if not raw_data: continue
        
        # Decide if raw_data is a single row or a list of rows
        if isinstance(raw_data, list) and len(raw_data) > 0 and not isinstance(raw_data[0], (list, tuple)):
            rows_to_process = [raw_data]
        elif isinstance(raw_data, list):
            rows_to_process = raw_data
        else:
            continue

        header_text = ' '.join(str(h).lower() for h in header)
        context_is_optional = is_optional_context or any(w in header_text for w in ['optional', 'restricted', 'rh'])

        for raw in rows_to_process:
            if not isinstance(raw, (list, tuple)): continue
            row_text = ' '.join(str(c).lower() for c in raw if c)
            if not row_text: continue

            month_hint = None
            for m in MONTH_NAMES:
                if m in row_text or m[:3] in row_text:
                    month_hint = m
                    break
            
            name = None
            date_val = None
            is_optional = context_is_optional or any(w in row_text for w in ['optional', 'restricted', 'rh'])
            is_working = any(w in row_text for w in ['working day', 'non-holiday', 'work day'])
            
            # Step 1: Specific Column Mapping
            col_map = {str(h).lower().strip(): i for i, h in enumerate(header)}
            def get_col(*keys):
                for k in keys:
                    for hk, idx in col_map.items():
                        if k in hk and idx < len(raw):
                            return str(raw[idx]).strip()
                return None

            h_name = get_col('holiday', 'festival', 'occasion', 'name')
            h_date = get_col('date', 'on_date')
            h_month = get_col('month')
            
            if h_date:
                date_val = smart_parse_date(h_date, year=doc_year, month_hint=(h_month or month_hint))
            if h_name:
                name = h_name

            # Step 2: Advanced Row Scanning
            potential_dates = []
            potential_names = []

            for i, cell in enumerate(raw):
                cs = str(cell).strip()
                if not cs or cs.lower() in ('nan', ''): continue
                
                # Rule out simple serial numbers (column 0, numeric, small)
                if i == 0 and re.fullmatch(r'\d{1,2}', cs):
                    continue
                
                # Try as date
                d = smart_parse_date(cs, year=doc_year, month_hint=month_hint)
                if d:
                    priority = 0
                    if any(m in cs.lower() for m in ['jan','feb','mar','apr','may','jun','jul','aug','sep','oct','nov','dec']):
                        priority += 2
                    if not re.fullmatch(r'\d{1,2}', cs):
                        priority += 1
                    potential_dates.append((d, priority))
                
                # Try as name
                cs_low = cs.lower()
                is_month = cs_low in MONTH_NAMES or any(m[:3] == cs_low for m in MONTH_NAMES if len(cs_low)==3)
                if not is_month and not any(wd in cs_low for wd in ['monday','tuesday','wednesday','thursday','friday','saturday','sunday']):
                    if len(cs) > 2 and not re.search(r'\d', cs):
                        potential_names.append(cs)
                    elif len(cs) > 3 and not re.match(r'^\d+$', cs) and not smart_parse_date(cs, year=doc_year):
                        potential_names.append(cs)

            if not date_val and potential_dates:
                potential_dates.sort(key=lambda x: x[1], reverse=True)
                date_val = potential_dates[0][0]
            
            if not name and potential_names:
                name_candidates = [n for n in potential_names if n.lower() not in MONTH_NAMES]
                if name_candidates:
                    name = max(name_candidates, key=len)

            if not name or not date_val: continue

            # Cleanup
            name = re.sub(r'^\d+[\.\s]*-?\s*', '', name).strip()
            if any(x in name.lower() for x in ['holiday', 'festival', 'sr no']): continue
            if name.lower() in MONTH_NAMES: continue

            import calendar as cal_mod
            day_str = cal_mod.day_name[date_val.weekday()]

            key = (name.lower(), str(date_val))
            if key in seen_keys: continue
            seen_keys.add(key)

            holidays.append({
                'name': name[:200],
                'date': str(date_val),
                'day': day_str,
                'is_optional': is_optional,
                'is_working_day': is_working,
                'year': date_val.year,
            })

    return sorted(holidays, key=lambda h: (h['date'], h['name']))





# --- Function: holiday_upload_parse ---
@api_view(['POST'])
@require_gated_token_api
@parser_classes([MultiPartParser, FormParser])
def holiday_upload_parse(request):
    """
    Admin uploads a holiday document.
    Parses it and returns a preview — no DB save yet.
    """
    user_id = request.data.get('user_id')
    employee = Employee.objects.filter(id=user_id, role='admin').first()
    if not employee:
        return Response({'success': False, 'message': 'Admin access required.'}, status=403)

    uploaded_file = request.FILES.get('file')
    if not uploaded_file:
        return Response({'success': False, 'message': 'No file provided.'}, status=400)

    ext = os.path.splitext(uploaded_file.name)[1].lower()
    allowed = ['.pdf', '.docx', '.xlsx', '.xls', '.csv', '.txt']
    if ext not in allowed:
        return Response({
            'success': False,
            'message': f'Unsupported format. Allowed: {", ".join(allowed)}'
        }, status=400)

    try:
        if not uploaded_file:
            return Response({'success': False, 'message': 'No file provided.'}, status=400)

        file_bytes = uploaded_file.read()
        raw_rows = []

        if ext == '.pdf':
            raw_rows = _parse_holiday_pdf(file_bytes)
        elif ext == '.docx':
            raw_rows = _parse_holiday_docx(file_bytes)
        elif ext in ['.xlsx', '.xls', '.csv']:
            raw_rows = _parse_holiday_excel(file_bytes, ext)
        else:  # .txt
            raw_rows = _parse_holiday_txt(file_bytes)

        holidays = _normalize_rows(raw_rows)

        # Create HolidayUpload audit record
        upload_obj = HolidayUpload.objects.create(
            file_name=uploaded_file.name[:255],
            uploaded_by=employee,
            status='parsed',
            parsed_count=len(holidays),
        )
        
        # Save file to media storage
        try:
            uploaded_file.seek(0)
            upload_obj.file.save(uploaded_file.name, uploaded_file, save=True)
        except Exception as file_err:
            print(f"Warning: Could not save holiday file to storage: {file_err}")

        return Response({
            'success': True,
            'upload_id': upload_obj.id,
            'file_name': uploaded_file.name,
            'parsed_count': len(holidays),
            'holidays': holidays,
        })

    except Exception as e:
        import traceback
        error_msg = f"Holiday Parse Error: {str(e)}"
        print(error_msg)
        print(traceback.format_exc())
        return Response({
            'success': False, 
            'message': error_msg,
            'debug_trace': traceback.format_exc() if settings.DEBUG else None
        }, status=500)




# --- Function: holiday_save ---
@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def holiday_save(request):
    """
    Admin approves parsed holidays and saves them to DB.
    Accepts { user_id, upload_id, holidays: [...] }
    Handles duplicate-date conflicts gracefully.
    """
    user_id = request.data.get('user_id')
    employee = Employee.objects.filter(id=user_id, role='admin').first()
    if not employee:
        return Response({'success': False, 'message': 'Admin access required.'}, status=403)

    upload_id = request.data.get('upload_id')
    holidays_data = request.data.get('holidays', [])

    created = 0
    updated = 0
    skipped = 0
    errors = []

    from datetime import datetime as dt
    for h in holidays_data:
        try:
            date_val = dt.strptime(h['date'], '%Y-%m-%d').date()
            obj, created_flag = Holiday.objects.update_or_create(
                date=date_val,
                defaults={
                    'name': h.get('name', 'Holiday')[:200],
                    'day': h.get('day', ''),
                    'is_optional': bool(h.get('is_optional', False)),
                    'is_working_day': bool(h.get('is_working_day', False)),
                    'year': date_val.year,
                    'description': h.get('description', ''),
                }
            )
            if created_flag:
                created += 1
            else:
                updated += 1
        except Exception as e:
            skipped += 1
            errors.append(str(e))

    # mark upload as approved
    if upload_id:
        HolidayUpload.objects.filter(id=upload_id).update(
            status='approved',
            saved_count=created + updated
        )

    # Audit log notification to admin
    Notification.objects.create(
        user=employee,
        type='holiday_upload',
        message=f'Holiday list saved: {created} added, {updated} updated, {skipped} skipped.',
    )

    return Response({
        'success': True,
        'created': created,
        'updated': updated,
        'skipped': skipped,
        'errors': errors[:5],
    })




# --- Function: update_holiday ---
@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def update_holiday(request):
    """
    Update an existing holiday. Admin only.
    """
    user_id = request.data.get('user_id')
    employee = Employee.objects.filter(id=user_id, role='admin').first()
    if not employee:
        return Response({'success': False, 'message': 'Admin access required.'}, status=403)

    holiday_id = request.data.get('holiday_id')
    name = request.data.get('name')
    date_str = request.data.get('date')
    day_str = request.data.get('day')
    is_optional = request.data.get('is_optional')
    description = request.data.get('description', '')

    if not holiday_id:
        return Response({'success': False, 'message': 'holiday_id is required.'}, status=400)

    try:
        holiday = Holiday.objects.get(id=holiday_id)
        if name: holiday.name = name
        if date_str:
            new_date = datetime.strptime(date_str, '%Y-%m-%d').date()
            if Holiday.objects.filter(date=new_date).exclude(id=holiday_id).exists():
                return Response({'success': False, 'message': f'A holiday already exists on {date_str}'}, status=400)
            holiday.date = new_date
            holiday.year = new_date.year
            
            if day_str:
                holiday.day = day_str
            else:
                # derive weekday if not provided
                import calendar as cal_mod
                holiday.day = cal_mod.day_name[new_date.weekday()]
        elif day_str:
            holiday.day = day_str
        
        if is_optional is not None:
            holiday.is_optional = bool(is_optional)
        
        is_working_day = request.data.get('is_working_day')
        if is_working_day is not None:
            holiday.is_working_day = bool(is_working_day)
        
        holiday.description = description
        holiday.save()

        return Response({'success': True, 'message': 'Holiday updated successfully.'})
    except Holiday.DoesNotExist:
        return Response({'success': False, 'message': 'Holiday not found.'}, status=404)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=500)




# --- Function: delete_holiday ---
@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def delete_holiday(request):
    """
    Delete a holiday. Admin only.
    """
    user_id = request.data.get('user_id')
    employee = Employee.objects.filter(id=user_id, role='admin').first()
    if not employee:
        return Response({'success': False, 'message': 'Admin access required.'}, status=403)

    holiday_id = request.data.get('holiday_id')
    if not holiday_id:
        return Response({'success': False, 'message': 'holiday_id is required.'}, status=400)

    try:
        holiday = Holiday.objects.get(id=holiday_id)
        holiday.delete()
        return Response({'success': True, 'message': 'Holiday deleted successfully.'})
    except Holiday.DoesNotExist:
        return Response({'success': False, 'message': 'Holiday not found.'}, status=404)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=500)





# --- Function: get_holidays ---
@api_view(['GET'])
@require_gated_token_api
def get_holidays(request):
    """
    Returns holidays for a given year (or current year).
    Optionally includes user's optional holiday selections.
    Query params: year, user_id
    """
    year = request.GET.get('year', timezone.now().year)
    user_id = request.GET.get('user_id')

    try:
        year = int(year)
    except (ValueError, TypeError):
        year = timezone.now().year

    holidays = Holiday.objects.filter(year=year)

    # get this user's selected optional holidays
    selected_ids = set()
    if user_id:
        selected_ids = set(
            UserHoliday.objects.filter(
                user_id=user_id, holiday__year=year
            ).values_list('holiday_id', flat=True)
        )

    data = []
    for h in holidays:
        data.append({
            'id': h.id,
            'name': h.name,
            'date': str(h.date),
            'day': h.day,
            'is_optional': h.is_optional,
            'is_working_day': h.is_working_day,
            'year': h.year,
            'description': h.description or '',
            'user_selected': h.id in selected_ids,
        })

    return Response({'success': True, 'holidays': data, 'year': year})




# --- Function: export_holidays_ics ---
@api_view(['GET'])
@require_gated_token_api
def export_holidays_ics(request):
    """
    Export holidays as an ICS calendar file.
    Query params: year, user_id (if provided, includes user's optional selections)
    """
    year = int(request.GET.get('year', timezone.now().year))
    user_id = request.GET.get('user_id')

    try:
        from icalendar import Calendar, Event as ICSEvent
        from datetime import datetime as dt
        import pytz

        cal = Calendar()
        cal.add('prodid', '-//HANUSPHERE Holiday Calendar//EN')
        cal.add('version', '2.0')
        cal.add('calscale', 'GREGORIAN')
        cal.add('x-wr-calname', f'Holidays {year}')

        holidays = Holiday.objects.filter(year=year)
        selected_ids = set()
        if user_id:
            selected_ids = set(
                UserHoliday.objects.filter(
                    user_id=user_id, holiday__year=year
                ).values_list('holiday_id', flat=True)
            )

        for h in holidays:
            # Skip optional holidays unless user selected them (or no user_id – export all)
            if h.is_optional and user_id and h.id not in selected_ids:
                continue

            event = ICSEvent()
            event.add('summary', h.name)
            event.add('dtstart', h.date)
            event.add('dtend', h.date)
            event.add('description', f"{'Optional' if h.is_optional else 'Holiday'}")
            cal.add_component(event)

        ics_bytes = cal.to_ical()
        response = HttpResponse(ics_bytes, content_type='text/calendar; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="holidays_{year}.ics"'
        return response

    except ImportError:
        return Response({'success': False, 'message': 'icalendar library not installed.'}, status=500)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=500)




# --- Function: holiday_upload_history ---
@api_view(['GET'])
@require_gated_token_api
def holiday_upload_history(request):
    """
    Returns the last 20 holiday upload audit records (admin only).
    """
    user_id = request.GET.get('user_id')
    employee = Employee.objects.filter(id=user_id, role='admin').first()
    if not employee:
        return Response({'success': False, 'message': 'Admin access required.'}, status=403)

    uploads = HolidayUpload.objects.all()[:20]
    data = []
    for u in uploads:
        data.append({
            'id': u.id,
            'file_name': u.file_name,
            'uploaded_by': u.uploaded_by.name if u.uploaded_by else 'Unknown',
            'uploaded_at': u.uploaded_at.strftime('%Y-%m-%d %H:%M'),
            'status': u.status,
            'parsed_count': u.parsed_count,
            'saved_count': u.saved_count,
        })
    return Response({'success': True, 'uploads': data})



