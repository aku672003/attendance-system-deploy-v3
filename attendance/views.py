from django.shortcuts import render, redirect
from django.http import HttpResponse, JsonResponse, HttpResponseForbidden, FileResponse
from .security import require_valid_token, require_gated_token_api
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.utils.decorators import method_decorator
from django.db.models import Q, Count, Sum, Avg, Prefetch
from django.utils import timezone
from django.core.cache import cache
from django.core.mail import send_mail
from django.conf import settings
import random
import string
from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import JSONParser, MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework import status
import json
import uuid
import math
import os
import hashlib
import hmac
import zipfile
import tempfile
from datetime import datetime, date, time, timedelta
from .models import (
    Employee, EmployeeProfile, OfficeLocation, DepartmentOfficeAccess,
    AttendanceRecord, EmployeeRequest, EmployeeDocument, Task, BirthdayWish, TaskComment, TaskStep, TaskAttachment, Team,
    TemporaryTag, TrainingLog, AvatarAsset, Memoji, Notification, TaskHistory
)
from .serializers import AvatarAssetSerializer, MemojiSerializer
from django.contrib.auth.hashers import make_password, check_password


def calculate_distance(lat1, lon1, lat2, lon2):
    """Calculate distance between two points using Haversine formula"""
    R = 6371000  # Earth radius in meters
    phi1 = math.radians(float(lat1))
    phi2 = math.radians(float(lat2))
    delta_phi = math.radians(float(lat2) - float(lat1))
    delta_lambda = math.radians(float(lon2) - float(lon1))

    a = math.sin(delta_phi / 2) ** 2 + \
        math.cos(phi1) * math.cos(phi2) * math.sin(delta_lambda / 2) ** 2
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))

    return R * c


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def send_otp(request):
    """
    Send a 6-digit OTP to the user's email for password reset.
    """
    username = request.data.get('username', '').strip()
    email = request.data.get('email', '').strip()
    
    if not username or not email:
        return Response({'success': False, 'message': 'Username and Email are required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        employee = Employee.objects.get(username__iexact=username, email__iexact=email, is_active=True)
    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'Account not found with this username and email combo'}, status=status.HTTP_404_NOT_FOUND)

    # Generate 6-digit OTP
    otp = ''.join(random.choices(string.digits, k=6))
    
    # Store OTP in cache for 5 minutes
    cache_key = f"otp_{email}"
    cache.set(cache_key, otp, timeout=300)

    # Send email
    subject = 'Password Reset OTP - HanuAI Attendance System'
    message = f'Your OTP for password reset is: {otp}. It is valid for 5 minutes.'
    from_email = settings.DEFAULT_FROM_EMAIL
    recipient_list = [email]

    try:
        send_mail(subject, message, from_email, recipient_list, fail_silently=False)
        return Response({
            'success': True, 
            'message': 'OTP sent to your email',
            'debug_info': {
                'recipient': email,
                'from': from_email,
                'backend': settings.EMAIL_BACKEND
            }
        })
    except Exception as e:
        print(f"Error sending email: {e}")
        return Response({'success': False, 'message': 'Failed to send OTP. Please try again later.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def reset_password(request):
    """
    Verify OTP and reset user password.
    """
    username = request.data.get('username')
    email = request.data.get('email')
    otp = request.data.get('otp')
    new_password = request.data.get('new_password')

    if not username or not email or not otp or not new_password:
        return Response({'success': False, 'message': 'Username, Email, OTP, and new password are required'}, status=status.HTTP_400_BAD_REQUEST)

    # Verify OTP from cache
    cache_key = f"otp_{email}"
    cached_otp = cache.get(cache_key)

    if not cached_otp:
        return Response({'success': False, 'message': 'OTP expired or not found'}, status=status.HTTP_400_BAD_REQUEST)

    if cached_otp != otp:
        return Response({'success': False, 'message': 'Invalid OTP'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        employee = Employee.objects.get(username__iexact=username, email__iexact=email, is_active=True)
        employee.password = make_password(new_password)
        employee.save()
        
        # Clear OTP from cache
        cache.delete(cache_key)
        
        return Response({'success': True, 'message': 'Password reset successful'})
    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        print(f"Password reset error: {e}")
        return Response({'success': False, 'message': 'Failed to reset password'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def get_current_user(request, requested_id=None, require_admin=False):
    """
    Helper to extract and validate the user for API requests.
    Prioritizes request.user (attached via gated token decorator).
    If a specific requested_id is provided, it verifies the caller is authorized.
    """
    # 1. Get the authenticated user from the token (attached by decorator)
    token_user = getattr(request, 'user', None)
    
    # 2. If no token user (legacy/development without token), fallback to requested_id
    if not isinstance(token_user, Employee):
        if not requested_id:
            return None
        return Employee.objects.filter(id=requested_id).first()

    # 3. If token user exists, enforce that they only access their own data
    # (Unless they are an admin or it's an admin-level request)
    if requested_id and str(requested_id) != str(token_user.id):
        if token_user.role != 'admin':
            return None # Unauthorized
            
    if require_admin and token_user.role != 'admin':
        return None

    return token_user

@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def login(request):
    """Authenticate user credentials and return profile data"""
    data = request.data
    username = data.get('username')
    password = data.get('password')

    if not username or not password:
        return Response({
            'success': False,
            'message': 'Username and password are required'
        }, status=status.HTTP_400_BAD_REQUEST)

    try:
        # Support login via both username OR email
        employee = Employee.objects.get(Q(username=username) | Q(email=username), is_active=True)
        
        # Check password — no backdoors, hashed comparison only
        if check_password(password, employee.password):
            profile = EmployeeProfile.objects.filter(employee=employee).first()
            assignment = employee.get_current_assignment()
            user_data = {
                'id': employee.id,
                'username': employee.username,
                'name': employee.name,
                'email': employee.email,
                'phone': employee.phone,
                'department': assignment['department'],
                'primary_office': employee.primary_office,
                'role': assignment['role'],
                'is_temporary': assignment['is_temporary'],
                'has_subordinates': employee.subordinates.exists(),
                'gender': profile.gender if profile else None,
                'date_of_birth': str(profile.date_of_birth) if profile and profile.date_of_birth else None,
                'avatar_emoji': profile.avatar_emoji if profile else "👤",
                'avatar_url': profile.avatar_url if profile else None,
                'theme_settings': profile.theme_settings if profile else {},
                'mentors': [{'id': m.id, 'name': m.name} for m in employee.mentors.all()],
                'total_cl': profile.total_cl if profile else 12,
                'taken_cl': profile.taken_cl if profile else 0,
            }
            return Response({
                'success': True,
                'user': user_data,
                'message': 'Login successful'
            })
        else:
            return Response({
                'success': False,
                'message': 'Invalid username/email or password'
            }, status=status.HTTP_401_UNAUTHORIZED)
    except Employee.DoesNotExist:
        # Check if the user exists but is inactive (via username or email)
        if Employee.objects.filter(Q(username=username) | Q(email=username), is_active=False).exists():
             return Response({
                'success': False,
                'message': 'Your account is inactive. Please contact the administrator.'
            }, status=status.HTTP_403_FORBIDDEN)

        return Response({
            'success': False,
            'message': 'Invalid username/email or password'
        }, status=status.HTTP_401_UNAUTHORIZED)
    except Exception as e:
        import logging
        logging.getLogger('attendance').exception('Login error')
        return Response({
            'success': False,
            'message': 'Login failed. Please try again.'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def register(request):
    """Register a new employee with validated details (admin only)"""
    # Verify caller is an admin
    caller = get_current_user(request, require_admin=True)
    if not caller:
        return Response({'success': False, 'message': 'Admin access required to register users'}, status=403)

    data = request.data
    required_fields = ['username', 'password', 'name', 'email', 'phone', 'department', 'primary_office']

    for field in required_fields:
        if not data.get(field):
            return Response({
                'success': False,
                'message': f"Field '{field}' is required"
            }, status=status.HTTP_400_BAD_REQUEST)

    # Validate phone number
    if not data['phone'].isdigit() or len(data['phone']) != 10:
        return Response({
            'success': False,
            'message': 'Phone number must be exactly 10 digits'
        }, status=status.HTTP_400_BAD_REQUEST)

    # Check if username or email already exists
    if Employee.objects.filter(Q(username__iexact=data['username']) | Q(email__iexact=data['email'])).exists():
        return Response({
            'success': False,
            'message': 'Username or email already exists'
        }, status=status.HTTP_400_BAD_REQUEST)

    try:
        employee = Employee.objects.create(
            username=data['username'],
            password=make_password(data['password']),
            name=data['name'],
            email=data['email'],
            phone=data['phone'],
            department=data['department'],
            primary_office=data['primary_office'],
            role=data.get('role', 'employee'),
            is_active=True
        )
        
        # Handle multiple Mentors if provided as list, or single mentor_id
        mentor_ids = data.get('mentor_ids') or []
        if not mentor_ids and data.get('mentor_id') and data.get('mentor_id') != 'none':
            mentor_ids = [data.get('mentor_id')]
            
        if mentor_ids:
            employee.mentors.set(Employee.objects.filter(id__in=mentor_ids))

        # Create Profile and set initial leaves
        joining_date_str = data.get('date_of_joining')
        total_cl = int(data.get('total_cl', 12))
        
        # If joining date provided and it's this year, calculate pro-rata CL
        if joining_date_str:
            try:
                joining_date = datetime.strptime(joining_date_str, '%Y-%m-%d').date()
                today = date.today()
                if joining_date.year == today.year:
                    # 1 CL per month remaining (including joining month)
                    total_cl = 12 - joining_date.month + 1
            except Exception as e:
                print(f"Error calculating pro-rata CL: {e}")

        profile = EmployeeProfile.objects.create(
            employee=employee,
            total_cl=total_cl,
            taken_cl=0,
            date_of_joining=joining_date_str
        )

        return Response({
            'success': True,
            'message': 'Account created successfully',
            'employee_id': employee.id
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Registration failed. Please try again.'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def offices_list(request):
    """Retrieve registered active office locations"""
    department = request.GET.get('department')
    active_param = request.GET.get('active')
    only_active = active_param not in ['0', 'false', 'False']

    try:
        offices = OfficeLocation.objects.all()
        if only_active:
            offices = offices.filter(is_active=True)
        offices = offices.order_by('name')

        offices_data = [{
            'id': office.id,
            'name': office.name,
            'address': office.address,
            'latitude': float(office.latitude),
            'longitude': float(office.longitude),
            'radius_meters': office.radius_meters,
            'is_active': office.is_active,
        } for office in offices]

        return Response({
            'success': True,
            'offices': offices_data
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to fetch office information'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def check_location(request):
    """Check if user location is within office geofence"""
    data = request.data
    user_lat = data.get('latitude')
    user_lng = data.get('longitude')
    office_id = data.get('office_id')

    if not all([user_lat, user_lng, office_id]):
        return Response({
            'success': False,
            'message': 'Latitude, longitude, and office_id are required'
        }, status=status.HTTP_400_BAD_REQUEST)

    try:
        office = OfficeLocation.objects.get(id=office_id)
        distance = calculate_distance(
            user_lat, user_lng,
            float(office.latitude), float(office.longitude)
        )

        return Response({
            'success': True,
            'distance': distance,
            'in_range': distance <= office.radius_meters,
            'office_location': {
                'latitude': float(office.latitude),
                'longitude': float(office.longitude),
                'radius_meters': office.radius_meters,
            }
        })
    except OfficeLocation.DoesNotExist:
        return Response({
            'success': False,
            'message': 'Office not found'
        }, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to check location'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def mark_attendance(request):
    data = request.data
    employee_id = data.get('employee_id')
    
    # Secure identity: Use token user if available
    user = get_current_user(request, requested_id=employee_id)
    if not user:
        return Response({'success': False, 'message': 'Unauthorized'}, status=403)
        
    now_local = timezone.localtime(timezone.now())
    att_date = now_local.date()
    
    # 0. Restriction check: 9 AM - 6 PM for non-Surveyors
    assignment = user.get_current_assignment()
    is_admin = user.role == 'admin'
    
    if assignment['department'] != 'Surveyors' and not is_admin:
        current_hour = now_local.hour
        if current_hour < 9 or current_hour >= 18:
            return Response({
                'success': False,
                'message': 'Non-surveyors can only check in between 9:00 AM and 6:00 PM.'
            }, status=status.HTTP_400_BAD_REQUEST)

    # 0.5 Missed Check-outs Penalty Check
    # Check the last 3 days where the user had a check-in
    past_3_records = list(AttendanceRecord.objects.filter(
        employee_id=user.id,
        date__lt=att_date,
        check_in_time__isnull=False
    ).exclude(status__in=['absent', 'leave']).order_by('-date')[:3])

    if len(past_3_records) == 3 and all(r.check_out_time is None for r in past_3_records):
        last_missed_date = past_3_records[0].date
        # Check if an unblock request exists that was created AFTER the last missed checkout
        unblock_req = EmployeeRequest.objects.filter(
            employee_id=user.id,
            request_type='unblock_attendance',
            created_at__date__gte=last_missed_date
        ).order_by('-created_at').first()
        
        if not unblock_req or unblock_req.status == 'rejected':
            return Response({
                'success': False,
                'error_code': 'ATTENDANCE_BLOCKED',
                'message': 'You have checked in but not checked out for 3 consecutive days. Please send a request to the Admin to unblock your attendance.'
            }, status=status.HTTP_403_FORBIDDEN)
        elif unblock_req.status == 'pending':
            return Response({
                'success': False,
                'error_code': 'ATTENDANCE_BLOCKED_PENDING',
                'message': 'Your request to unblock attendance is pending Admin approval.'
            }, status=status.HTTP_403_FORBIDDEN)
        # If unblock_req.status == 'approved', we allow check-in!

    # 0.7 Auto-close any unclosed records from previous days
    # This prevents sessions from spanning multiple days and causing "99+ hours" bugs.
    AttendanceRecord.objects.filter(
        employee_id=user.id,
        date__lt=att_date,
        check_in_time__isnull=False,
        check_out_time__isnull=True
    ).exclude(status__in=['absent', 'leave']).update(
        status='absent',
        notes="Auto-marked absent: New check-in started on a later date"
    )

    # 1. Check if they already have a SUCCESSFUL check-in TODAY
    # We look for a record that HAS a check-in time and matches TODAY's date
    today_record = AttendanceRecord.objects.filter(
        employee_id=user.id, 
        date=att_date
    ).exclude(status='absent').first()

    if today_record and today_record.check_in_time:
        return Response({
            'success': False,
            'message': 'Attendance already marked for today'
        }, status=status.HTTP_400_BAD_REQUEST)

    # 1.5 WFH Approval Check
    if data.get('status') == 'wfh' or data.get('type') == 'wfh':
        wfh_check = check_wfh_eligibility(user.id, att_date.isoformat())
        if not wfh_check.get('has_approved_request'):
            return Response({
                'success': False,
                'message': 'You need an approved WFH request to check in for Work From Home today.'
            }, status=status.HTTP_403_FORBIDDEN)

    # 2. If an 'absent' placeholder exists for today (from your auto-logic), 
    # we update it instead of creating a duplicate.
    absent_record = AttendanceRecord.objects.filter(employee_id=user.id, date=att_date, status='absent').first()
    
    try:
        if absent_record:
            absent_record.check_in_time = now_local.time().strftime('%H:%M:%S')
            absent_record.status = data.get('status')
            absent_record.type = data.get('type')
            absent_record.check_in_location = data.get('location')
            absent_record.save()
            record = absent_record
        else:
            record = AttendanceRecord.objects.create(
                employee_id=user.id,
                date=att_date,
                check_in_time=now_local.time().strftime('%H:%M:%S'),
                type=data.get('type'),
                status=data.get('status'),
                check_in_location=data.get('location'),
                check_in_photo=data.get('photo'),
                office_id=data.get('office_id')
            )
        return Response({'success': True, 'message': 'Checked in successfully'})
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=500)

@api_view(['GET'])
@require_gated_token_api
def get_server_time(request):
    """Return the current server time in IST for frontend synchronization"""
    now_local = timezone.localtime(timezone.now())
    return Response({
        'success': True,
        'timestamp': now_local.timestamp() * 1000, # Milliseconds
        'formatted': now_local.strftime('%Y-%m-%d %H:%M:%S'),
        'timezone': 'Asia/Kolkata'
    })





@api_view(['POST'])
@require_gated_token_api
def create_team(request):
    try:
        data = request.data
        mentor_id = data.get('mentor_id')
        name = data.get('name')
        member_ids = data.get('members', [])

        if not mentor_id:
             return Response({'success': False, 'message': 'Mentor ID required'})

        mentor_obj = Employee.objects.filter(id=mentor_id).first()
        if not mentor_obj:
            return Response({'success': False, 'message': 'Mentor not found'})

        team = Team.objects.create(name=name, mentor=mentor_obj)
        
        if member_ids:
            members = Employee.objects.filter(id__in=member_ids)
            team.members.set(members)

        return Response({'success': True, 'message': 'Team created successfully', 'team_id': team.id})
    except Exception as e:
        return Response({'success': False, 'message': str(e)})

@api_view(['POST'])
@require_gated_token_api
def update_team(request):
    try:
        data = request.data
        team_id = data.get('team_id')
        name = data.get('name')
        member_ids = data.get('members', [])

        team = Team.objects.filter(id=team_id).first()
        if not team:
            return Response({'success': False, 'message': 'Team not found'})

        if name:
            team.name = name
        
        if member_ids:
            members = Employee.objects.filter(id__in=member_ids)
            team.members.set(members)
        
        team.save()
        return Response({'success': True, 'message': 'Team updated successfully'})
    except Exception as e:
        return Response({'success': False, 'message': str(e)})

@api_view(['DELETE', 'POST']) # Support POST with method override for simplicity if needed
@require_gated_token_api
def delete_team(request):
    try:
        team_id = request.data.get('team_id') or request.query_params.get('team_id')
        team = Team.objects.filter(id=team_id).first()
        if not team:
            return Response({'success': False, 'message': 'Team not found'})
            
        team.delete()
        return Response({'success': True, 'message': 'Team deleted successfully'})
    except Exception as e:
        return Response({'success': False, 'message': str(e)})


@api_view(['GET'])
@require_gated_token_api
def get_teams(request):
    try:
        mentor_id = request.query_params.get('mentor_id')
        if not mentor_id:
             return Response({'success': False, 'message': 'Mentor ID required'})
             
        teams = Team.objects.filter(mentor_id=mentor_id).prefetch_related('members')
        data = []
        for t in teams:
            data.append({
                'id': t.id,
                'name': t.name,
                'members': list(t.members.values('id', 'name', 'username', 'role'))
            })
            
        return Response({'success': True, 'teams': data})
    except Exception as e:
        return Response({'success': False, 'message': str(e)})

def check_location_proximity(lat, lng, office_id):
    """Helper function to check location proximity"""
    try:
        office = OfficeLocation.objects.get(id=office_id)
        distance = calculate_distance(
            lat, lng,
            float(office.latitude), float(office.longitude)
        )
        return {
            'success': True,
            'distance': distance,
            'in_range': distance <= office.radius_meters,
        }
    except:
        return {'success': False, 'in_range': False}


def check_wfh_eligibility(employee_id, check_date):
    """Check WFH eligibility for an employee"""
    try:
        check_date_obj = datetime.strptime(check_date, '%Y-%m-%d').date()
        
        # Check if there is an APPROVED WFH request for this date
        has_approved_request = EmployeeRequest.objects.filter(
            employee_id=employee_id,
            request_type='wfh',
            start_date__lte=check_date_obj,
            end_date__gte=check_date_obj,
            status='approved'
        ).exists()

        # Count approved WFH requests for the current month (for dashboard stats)
        current_month_requests = EmployeeRequest.objects.filter(
            employee_id=employee_id,
            request_type='wfh',
            status='approved',
            start_date__year=check_date_obj.year,
            start_date__month=check_date_obj.month
        ).count()

        return {
            'has_approved_request': has_approved_request,
            'can_request': has_approved_request, # Only allow if approved
            'current_count': current_month_requests,
            'max_limit': 1 # Hardcoded limit as per frontend logic
        }
    except Exception as e:
        print(f"Error checking WFH eligibility: {e}")
        return {'has_approved_request': False, 'can_request': False, 'current_count': 0, 'max_limit': 1}


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def check_out(request):
    data = request.data
    employee_id = data.get('employee_id')
    
    # Secure identity: Use token user if available
    user = get_current_user(request, requested_id=employee_id)
    if not user:
        return Response({'success': False, 'message': 'Unauthorized'}, status=403)
        
    now_local = timezone.localtime(timezone.now())
    
    try:
        # 1. Try to find the specific record for the date provided by the frontend
        target_date = data.get('date')
        if target_date:
            record = AttendanceRecord.objects.filter(
                employee_id=user.id,
                date=target_date,
                check_in_time__isnull=False,
                check_out_time__isnull=True
            ).first()
        else:
            record = None

        # 2. Fallback to the latest unclosed record if not found for specific date
        if not record:
            record = AttendanceRecord.objects.filter(
                employee_id=user.id, 
                check_in_time__isnull=False,
                check_out_time__isnull=True
            ).exclude(status__in=['absent', 'leave']).order_by('-date', '-check_in_time').first()

        # 3. Extra fallback: If today's record was auto-marked absent (forgot to check out earlier)
        if not record:
            record = AttendanceRecord.objects.filter(
                employee_id=user.id,
                date=now_local.date(),
                check_in_time__isnull=False,
                check_out_time__isnull=True,
                status='absent'
            ).first()

        if not record:
            return Response({'success': False, 'message': 'No active session found.'}, status=404)

        check_in_t = datetime.strptime(str(record.check_in_time), '%H:%M:%S').time()
        check_in_dt = timezone.make_aware(datetime.combine(record.date, check_in_t))
        
        # Calculate hours and cap at a reasonable maximum to prevent "99+ hours" bug
        # We enforce a hard cap of 14 hours per day.
        raw_hours = (now_local - check_in_dt).total_seconds() / 3600
        
        # If check-out is on a different day, or hours are excessive (>14h), cap it.
        # This handles cases where user forgets to check out for days.
        if record.date != now_local.date() or raw_hours > 14.0:
            worked_hours = 14.0
        else:
            worked_hours = round(max(0.0, raw_hours), 2)
        
        # Min hours requirement (e.g. 4.5h for a valid day)
        # Note: If it's a legitimate short day, user might need to check out anyway, 
        # but system might flag it. 4.5h is standard for half-day or minimal presence.
        if worked_hours < 4.5:
            # We allow it but maybe flag it? Actually, user said fix logic to be simple.
            # Let's keep the min check if it's already there.
            if worked_hours < 0.1: # Practically zero
                 return Response({'success': False, 'message': 'Minimum 0.1h required for check-out.'}, status=400)

        record.check_out_time = now_local.time().strftime('%H:%M:%S')
        record.total_hours = worked_hours
        
        if record.type == 'wfh':
            record.status = 'wfh'
            # Check tasks completed for today
            from .models import Task
            today_tasks = Task.objects.filter(assignees=user.id, due_date=record.date)
            total_tasks = today_tasks.count()
            completed_tasks = today_tasks.filter(status='completed').count()
            
            task_status_msg = f'Tasks: {completed_tasks}/{total_tasks} completed.'
            if total_tasks > 0:
                percent = round((completed_tasks / total_tasks * 100), 1)
                task_status_msg += f' ({percent}% completion)'
                if completed_tasks == total_tasks:
                    task_status_msg += ' - ALL TASKS CLEAR!'
                else:
                    task_status_msg += ' - PENDING TASKS!'
            else:
                task_status_msg = 'Tasks: No tasks assigned/due for today.'
                
            # Automate an approval request for Mentor/Admin to verify tasks at the end of the day
            # If a manual request already exists, update its reason to include task status
            wfh_req, created = EmployeeRequest.objects.get_or_create(
                employee_id=employee_id,
                request_type='wfh',
                start_date=record.date,
                end_date=record.date,
                defaults={
                    'status': 'pending', 
                    'reason': f'WFH Session Log ({worked_hours}h worked). {task_status_msg}'
                }
            )
            if not created and wfh_req.status == 'pending':
                wfh_req.reason = f'WFH Session Log ({worked_hours}h worked). {task_status_msg}'
                wfh_req.save()
            
            # Notify mentors if a request was created or task status was updated
            notification_msg = f"{record.employee.name}: WFH Log ({worked_hours}h). {task_status_msg}"
            for mentor in record.employee.mentors.all():
                _send_task_notification(mentor, notification_msg, wfh_req.id, type="request")
        elif record.type == 'client':
            record.status = 'client'
        else:
            # 9 Hours strict: Less than 9 hours is marked as half day
            record.status = 'half_day' if worked_hours < 9.0 else 'present'
            record.is_half_day = True if worked_hours < 9.0 else False
            
        record.save()
        
        return Response({'success': True, 'message': 'Checked out successfully'})
    except AttendanceRecord.DoesNotExist:
        return Response({'success': False, 'message': 'No active session found.'}, status=404)
    except Exception as e:
        return Response({'success': False, 'message': f'Check-out failed: {str(e)}'}, status=500)

@api_view(['GET'])
@require_gated_token_api
def today_attendance(request):
    """Get today's attendance for an employee"""
    employee_id = request.GET.get('employee_id')

    if not employee_id:
        return Response({
            'success': False,
            'message': 'Employee ID is required'
        }, status=status.HTTP_400_BAD_REQUEST)

    try:
        # Use server-side local date (timezone aware) to match mark_attendance logic
        today = timezone.localtime(timezone.now()).date()
        record = AttendanceRecord.objects.filter(
            employee_id=employee_id,
            date=today
        ).select_related('office').first()

        if record:
            record_data = {
                'id': record.id,
                'employee_id': record.employee_id,
                'date': str(record.date),
                'check_in_time': str(record.check_in_time) if record.check_in_time else None,
                'check_out_time': str(record.check_out_time) if record.check_out_time else None,
                'type': record.type,
                'status': record.status,
                'office_id': record.office_id,
                'office_name': record.office.name if record.office else None,
                'office_address': record.office.address if record.office else None,
                'check_in_location': record.check_in_location,
                'check_out_location': record.check_out_location,
                'check_out_photo_url': record.check_out_photo if record.check_out_photo and not record.check_out_photo.startswith('data:') else None,
                'total_hours': float(record.total_hours),
                'gender': getattr(record.employee.profile, 'gender', 'other') if hasattr(record.employee, 'profile') else 'other',
            }
            return Response({
                'success': True,
                'record': record_data
            })
        else:
            return Response({
                'success': True,
                'record': None
            })
    except Exception as e:
        return Response({
            'success': False,
            'message': "Failed to fetch today's attendance"
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def attendance_records(request):
    """Get attendance records with filters"""
    employee_id = request.GET.get('employee_id')
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    att_type = request.GET.get('type')
    days_limit = request.GET.get('days_limit')
    days_offset = int(request.GET.get('days_offset', 0))

    user = getattr(request, 'user', None)
    requested_user_id = request.GET.get('user_id')
    
    # Prefer user from token (request.user) over query param
    if not isinstance(user, Employee):
        user = Employee.objects.filter(id=requested_user_id).first() if requested_user_id else None
    elif requested_user_id and str(requested_user_id) != str(user.id) and user.role != 'admin':
        # Security: if a specific user_id was requested but doesn't match token, only allow for admins
        return Response({'success': False, 'message': 'Unauthorized user_id access'}, status=403)

    # Include de-facto mentors (any user who has subordinates)
    is_mentor = user and (user.role == 'mentor' or (user.role != 'admin' and user.subordinates.exists()))

    search = request.GET.get('search', '').strip().lower()

    # Auto-mark absentees only after 6 PM (scheduler handles the main trigger)
    now = timezone.localtime(timezone.now())
    today = now.date()
    if now.hour >= 18:
        mark_absentees_for_date(today)

    try:
        records_qs = AttendanceRecord.objects.select_related('employee', 'office').all()

        # Apply global search if provided
        if search:
            from django.db.models import Q
            
            # Month mapping for searching by month name
            months = {
                'jan': 1, 'january': 1, 'feb': 2, 'february': 2, 'mar': 3, 'march': 3,
                'apr': 4, 'april': 4, 'may': 5, 'jun': 6, 'june': 6, 'jul': 7, 'july': 7,
                'aug': 8, 'august': 8, 'sep': 9, 'september': 9, 'oct': 10, 'october': 10,
                'nov': 11, 'november': 11, 'dec': 12, 'december': 12
            }
            
            # Check if search term is a month name
            month_val = months.get(search)
            
            search_filter = Q(employee__name__icontains=search) | \
                            Q(employee__username__icontains=search) | \
                            Q(employee__department__icontains=search) | \
                            Q(office__name__icontains=search) | \
                            Q(status__icontains=search) | \
                            Q(type__icontains=search)
            
            if month_val:
                search_filter |= Q(date__month=month_val)
                
            # Handle specific date formats (e.g. 23-04-2026 or 2026-04-23)
            import re
            if re.match(r'^\d{2}-\d{2}-\d{4}$', search):
                try:
                    from datetime import datetime
                    d_obj = datetime.strptime(search, '%d-%m-%Y').date()
                    search_filter |= Q(date=d_obj)
                except: pass
            elif re.match(r'^\d{4}-\d{2}-\d{2}$', search):
                search_filter |= Q(date=search)
            elif search.isdigit() and len(search) == 4:
                search_filter |= Q(date__year=int(search))
            
            records_qs = records_qs.filter(search_filter)

        if is_mentor:
            from django.db.models import Q
            records_qs = records_qs.filter(Q(employee__mentors=user) | Q(employee=user)).distinct()

        if employee_id:
            records_qs = records_qs.filter(employee_id=employee_id)
        if start_date:
            records_qs = records_qs.filter(date__gte=start_date)
        if end_date:
            records_qs = records_qs.filter(date__lte=end_date)
        elif not search:
            # Default: only show records up to today (hide future-dated records)
            # But if searching, don't restrict to today by default unless requested
            records_qs = records_qs.filter(date__lte=today)

        if att_type:
            records_qs = records_qs.filter(type=att_type)

        records_qs = records_qs.distinct()

        has_more = False
        if days_limit:
            days_limit = int(days_limit)
            # Get unique dates in DESC order
            unique_dates = records_qs.values_list('date', flat=True).distinct().order_by('-date')
            total_days = unique_dates.count()
            total_days = unique_dates.count()

            target_dates = unique_dates[days_offset : days_offset + days_limit]
            has_more = total_days > (days_offset + days_limit)

            records_qs = records_qs.filter(date__in=target_dates)

        records_qs = records_qs.order_by('-date', '-created_at')

        records_data = []
        for record in records_qs:
            records_data.append({
                'id': record.id,
                'employee_id': record.employee_id,
                'employee_name': record.employee.name,
                'department': record.employee.department,
                'date': str(record.date),
                'check_in_time': str(record.check_in_time) if record.check_in_time else None,
                'check_out_time': str(record.check_out_time) if record.check_out_time else None,
                'type': record.type,
                'status': str(record.status or '').lower(),
                'office_id': record.office_id,
                'office_name': record.office.name if record.office else None,
                'office_address': record.office.address if record.office else None,
                'check_in_location': record.check_in_location,
                'check_out_location': record.check_out_location,
                'check_in_photo': record.check_in_photo,
                'check_out_photo': record.check_out_photo,
                'photo_url': record.check_out_photo or record.check_in_photo or None,
                'total_hours': float(record.total_hours),
                'is_half_day': record.is_half_day,
                'role': record.employee.role,
                'date_of_joining': str(record.employee.profile.date_of_joining) if hasattr(record.employee, 'profile') and record.employee.profile.date_of_joining else None,
            })

        return Response({
            'success': True,
            'records': records_data,
            'has_more': has_more
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to fetch attendance records'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

def mark_absentees_for_date(target_date):
    try:

        all_employees = Employee.objects.filter(is_active=True).values_list('id', flat=True)
        existing_records = AttendanceRecord.objects.filter(date=target_date).values_list('employee_id', flat=True)

        absentees = set(all_employees) - set(existing_records)

        if absentees:
            new_records = [
                AttendanceRecord(
                    employee_id=emp_id,
                    date=target_date,
                    status='absent',
                    type='office',
                    total_hours=0
                ) for emp_id in absentees
            ]
            AttendanceRecord.objects.bulk_create(new_records, ignore_conflicts=True)

        # 2. Mark employees who checked in but forgot to check out as 'absent'
        # This applies if the day has ended (6 PM for today, or any past day)
        now = timezone.localtime(timezone.now())
        if target_date < now.date() or (target_date == now.date() and now.hour >= 18):
            AttendanceRecord.objects.filter(
                date=target_date,
                check_in_time__isnull=False,
                check_out_time__isnull=True
            ).exclude(status__in=['absent', 'leave']).update(
                status='absent',
                notes="Absent marked: Forgot to check out"
            )
    except Exception as e:
        print(f"Error marking absentees: {e}")


@api_view(['GET'])
@require_gated_token_api
def monthly_stats(request):
    """Get monthly attendance statistics"""
    employee_id = request.GET.get('employee_id')
    year = request.GET.get('year') or date.today().year
    month = request.GET.get('month') or date.today().month

    if not employee_id:
        return Response({
            'success': False,
            'message': 'Employee ID is required'
        }, status=status.HTTP_400_BAD_REQUEST)

    try:
        records = AttendanceRecord.objects.filter(
            employee_id=employee_id,
            date__year=year,
            date__month=month
        )

        # Get employee profile for joining date and leave settings
        from attendance.models import EmployeeProfile
        profile = EmployeeProfile.objects.filter(employee_id=employee_id).first()
        joining_date = profile.date_of_joining if profile else None
        
        # Calculate monthly allowance and rollover
        # Logic: 1 CL per month. Jan-Dec cycle.
        current_year = int(year)
        current_month = int(month)
        
        # Determine start month (either January or joining month if in the same year)
        start_month = 1
        if joining_date and joining_date.year == current_year:
            start_month = joining_date.month
        elif joining_date and joining_date.year > current_year:
            # Not yet joined in the requested year
            start_month = 13 # Impossible month to show 0 balance
            
        # Total allowed leaves up to current viewed month
        # If viewed month < start_month, allowed = 0
        months_active_ytd = max(0, current_month - start_month + 1) if current_month >= start_month else 0
        total_allowed_upto_now = months_active_ytd # 1 per month
        
        # Yearly allowance (total for the year based on joining date)
        yearly_allowance = max(0, 12 - start_month + 1) if start_month <= 12 else 0
        
        # Total taken leaves in the WHOLE year (Full = 1, Half = 0.5)
        yearly_leaves_qs = EmployeeRequest.objects.filter(
            employee_id=employee_id,
            request_type__in=['full_day', 'half_day'],
            status='approved',
            start_date__year=current_year
        ).order_by('start_date')
        
        yearly_taken = 0.0
        yearly_leave_dates = []
        for req in yearly_leaves_qs:
            val = 1.0 if req.request_type == 'full_day' else 0.5
            yearly_taken += val
            yearly_leave_dates.append({
                'date': req.start_date.strftime('%Y-%m-%d'),
                'type': req.request_type
            })

        # Total taken leaves in this year up to the viewed month
        # ... and so on
        full_ytd = EmployeeRequest.objects.filter(
            employee_id=employee_id, request_type='full_day', status='approved',
            start_date__year=current_year, start_date__month__lte=current_month
        ).count()
        half_ytd = EmployeeRequest.objects.filter(
            employee_id=employee_id, request_type='half_day', status='approved',
            start_date__year=current_year, start_date__month__lte=current_month
        ).count()
        taken_leaves_ytd = float(full_ytd) + (float(half_ytd) * 0.5)
        
        # Taken in the current viewed month
        leaves_qs = EmployeeRequest.objects.filter(
            employee_id=employee_id,
            request_type__in=['full_day', 'half_day'],
            status='approved',
            start_date__year=current_year,
            start_date__month=current_month
        ).order_by('start_date')
        
        taken_this_month = 0.0
        leave_dates = []
        for req in leaves_qs:
            val = 1.0 if req.request_type == 'full_day' else 0.5
            taken_this_month += val
            leave_dates.append({
                'date': req.start_date.strftime('%Y-%m-%d'),
                'type': req.request_type
            })
        
        # Rollover from previous months in the same year
        allowance_this_month = 1 if current_month >= start_month else 0
        rollover = max(0, (total_allowed_upto_now - allowance_this_month) - (taken_leaves_ytd - taken_this_month))
        
        total_left_ytd = max(0, float(total_allowed_upto_now) - taken_leaves_ytd)
        yearly_left = max(0, float(yearly_allowance) - yearly_taken)

        from django.db.models import Count, Case, When, Sum, Q
        stats_data = records.aggregate(
            total_working_days=Count(Case(When(status__in=['present', 'half_day', 'wfh', 'client'], then=1))),
            weekday_present_days=Count(Case(When(Q(status__in=['present', 'half_day', 'wfh', 'client']) & Q(date__week_day__in=[2, 3, 4, 5, 6, 7]), then=1))),
            total_hours_sum=Sum('total_hours'),
            half_day_records=Count(Case(When(is_half_day=True, then=1))),
            wfh_days=Count(Case(When(type='wfh', then=1))),
            office_days=Count(Case(When(type='office', status='present', then=1))),
            client_days=Count(Case(When(type='client', then=1))),
            leave_records=Count(Case(When(status='leave', then=1))),
        )

        half_day_requests = EmployeeRequest.objects.filter(
            employee_id=employee_id,
            request_type='half_day',
            status='approved',
            start_date__year=year,
            start_date__month=month
        ).count()

        total_half_days = max(stats_data['half_day_records'], half_day_requests)

        # Count optional holidays for the year (Max 2)
        optional_holidays_count = UserHoliday.objects.filter(
            user_id=employee_id,
            holiday__year=year
        ).count()

        stats = {
            'total_working_days': stats_data['total_working_days'],
            'weekday_present_days': stats_data['weekday_present_days'],
            'total_hours': float(stats_data['total_hours_sum'] or 0),
            'half_days': total_half_days,
            'wfh_days': stats_data['wfh_days'],
            'office_days': stats_data['office_days'],
            'client_days': stats_data['client_days'],
            'leave_days': taken_this_month,
            'leave_allowance': allowance_this_month,
            'leave_rollover': rollover,
            'leave_total_left': total_left_ytd,
            'yearly_taken': yearly_taken,
            'yearly_allowance': yearly_allowance,
            'yearly_left': yearly_left,
            'optional_holidays': optional_holidays_count,
            'max_optional': 2,
            'leave_dates': leave_dates,
            'yearly_leave_dates': yearly_leave_dates,
        }

        return Response({
            'success': True,
            'stats': stats
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to fetch monthly statistics'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def wfh_eligibility(request):
    """Check WFH eligibility"""
    employee_id = request.GET.get('employee_id')
    check_date = request.GET.get('date') or date.today().isoformat()

    if not employee_id:
        return Response({
            'success': False,
            'message': 'Employee ID is required'
        }, status=status.HTTP_400_BAD_REQUEST)

    result = check_wfh_eligibility(employee_id, check_date)
    return Response({
        'success': True,
        **result
    })


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def wfh_request(request):
    """Submit WFH request"""
    data = request.data
    employee_id = data.get('employee_id')
    requested_date = data.get('date') or date.today().isoformat()
    reason = data.get('reason')

    if not employee_id:
        return Response({
            'success': False,
            'message': 'Employee ID is required'
        }, status=status.HTTP_400_BAD_REQUEST)

    try:
        req = EmployeeRequest.objects.create(
            employee_id=employee_id,
            request_type='wfh',
            start_date=requested_date,
            end_date=requested_date,
            reason=reason,
            status='pending'
        )
        
        # Notify mentors
        emp = Employee.objects.get(id=employee_id)
        for mentor in emp.mentors.all():
            _send_task_notification(mentor, f"{emp.name} requested WFH for {requested_date}", req.id, type="request")

        return Response({
            'success': True,
            'message': 'Request submitted'
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to submit request'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# Profile Management Views
@api_view(['GET', 'POST', 'PATCH'])
@require_gated_token_api
@parser_classes([JSONParser])
def employee_profile(request, employee_id=None):
    """Get or save employee profile"""
    
    # Handle PATCH (partial update, e.g. for theme settings)
    if request.method == 'PATCH':
        data = request.data
        employee_id = employee_id or data.get('employee_id')

        if not employee_id:
            return Response({'success': False, 'message': 'Employee ID is required'}, status=status.HTTP_400_BAD_REQUEST)
            
        try:
            employee = Employee.objects.get(id=employee_id)
            profile, created = EmployeeProfile.objects.get_or_create(employee=employee)
            if 'theme_settings' in data:
                profile.theme_settings = data.get('theme_settings')
                profile.save()
            return Response({'success': True, 'message': 'Profile updated successfully'})
        except Exception as e:
            return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    # Handle POST (save profile)
    if request.method == 'POST':
        data = request.data
        employee_id = employee_id or data.get('employee_id')

        if not employee_id:
            return Response({
                'success': False,
                'message': 'Employee ID is required'
            }, status=status.HTTP_400_BAD_REQUEST)

        try:
            employee = Employee.objects.get(id=employee_id)
            profile, created = EmployeeProfile.objects.get_or_create(employee=employee)

            # Update profile fields
            if 'emergency_contact_name' in data:
                profile.emergency_contact_name = data.get('emergency_contact_name')
            if 'emergency_contact_phone' in data:
                profile.emergency_contact_phone = data.get('emergency_contact_phone')
            if 'alternate_number' in data:
                profile.alternate_number = data.get('alternate_number')
            if 'bank_account_number' in data:
                profile.bank_account_number = data.get('bank_account_number')
            if 'bank_ifsc' in data:
                profile.bank_ifsc = data.get('bank_ifsc')
            if 'bank_name' in data:
                profile.bank_bank_name = data.get('bank_name')
            if 'pan_number' in data:
                profile.pan_number = data.get('pan_number')
            if 'aadhar_number' in data:
                profile.aadhar_number = data.get('aadhar_number')
            if 'highest_qualification' in data:
                profile.qualification = data.get('highest_qualification')
            if 'qualification_notes' in data:
                profile.certificates_summary = data.get('qualification_notes')
            if 'home_address' in data:
                profile.home_address = data.get('home_address')
            if 'current_address' in data:
                profile.current_address = data.get('current_address')
            if 'date_of_joining' in data:
                profile.date_of_joining = data.get('date_of_joining') or None
            if 'skill_set' in data:
                profile.skill_set = data.get('skill_set')
            if 'reporting_mentor' in data:
                profile.reporting_mentor = data.get('reporting_mentor')
            if 'professional_training' in data:
                profile.professional_training = data.get('professional_training')
            if 'family_details' in data:
                profile.family_details = data.get('family_details')
            if 'marital_status' in data:
                profile.marital_status = data.get('marital_status')
            if 'personal_email' in data:
                profile.personal_email = data.get('personal_email')
            if 'gender' in data:
                profile.gender = data.get('gender')
            if 'date_of_birth' in data:
                profile.date_of_birth = data.get('date_of_birth') or None
            
            # Personalization fields
            if 'avatar_emoji' in data:
                profile.avatar_emoji = data['avatar_emoji']
            if 'theme_settings' in data:
                profile.theme_settings = data['theme_settings']
            if 'avatar_url' in data and data['avatar_url']:
                profile.avatar_url = data['avatar_url']
                
            profile.save()

            # Update employee basic info if provided
            if data.get('name'):
                employee.name = data['name']
            if data.get('email'):
                employee.email = data['email']
            if data.get('phone'):
                employee.phone = data['phone']
            if data.get('primary_office'):
                employee.primary_office = data['primary_office']
            if data.get('password'):
                employee.password = make_password(data['password'])
            employee.save()

            return Response({
                'success': True,
                'message': 'Profile saved successfully'
            })
        except Employee.DoesNotExist:
            return Response({
                'success': False,
                'message': 'Employee not found'
            }, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({
                'success': False,
                'message': 'Failed to save profile'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    # Handle GET (get profile)
    employee_id = employee_id or request.GET.get('employee_id')

    if not employee_id:
        return Response({
            'success': False,
            'message': 'Employee ID is required'
        }, status=status.HTTP_400_BAD_REQUEST)

    try:
        employee = Employee.objects.get(id=employee_id)
        profile, _ = EmployeeProfile.objects.get_or_create(employee=employee)

        # Get documents
        documents = EmployeeDocument.objects.filter(employee_id=employee_id).order_by('-uploaded_at')
        docs_data = []
        for doc in documents:
            docs_data.append({
                'id': doc.id,
                'doc_type': doc.doc_type,
                'doc_name': doc.doc_name,
                'doc_number': doc.doc_number,
                'file_name': doc.file_name,
                'file_path': doc.file_path,
                'url': request.build_absolute_uri(f'/api/serve-document/{doc.id}'),
                'uploaded_at': doc.uploaded_at.isoformat() if doc.uploaded_at else None,
            })

        profile_data = {
            'id': employee.id,
            'username': employee.username,
            'name': employee.name,
            'official_email': employee.email,
            'official_phone': employee.phone,
            'department': employee.department,
            'emergency_contact_name': profile.emergency_contact_name,
            'emergency_contact_phone': profile.emergency_contact_phone,
            'alternate_number': profile.alternate_number,
            'bank_account_number': profile.bank_account_number,
            'bank_ifsc': profile.bank_ifsc,
            'bank_name': profile.bank_bank_name,
            'pan_number': profile.pan_number,
            'aadhar_number': profile.aadhar_number,
            'qualification': profile.qualification,
            'certificates_summary': profile.certificates_summary,
            'home_address': profile.home_address,
            'current_address': profile.current_address,
            'date_of_joining': str(profile.date_of_joining) if profile.date_of_joining else None,
            'skill_set': profile.skill_set,
            'reporting_mentor': ", ".join([m.name for m in employee.mentors.all()]) if employee.mentors.exists() else profile.reporting_mentor,
            'professional_training': profile.professional_training,
            'family_details': profile.family_details,
            'marital_status': profile.marital_status,
            'personal_email': profile.personal_email,
            'gender': profile.gender,
            'date_of_birth': str(profile.date_of_birth) if profile.date_of_birth else None,
            'avatar_emoji': profile.avatar_emoji,
            'avatar_url': profile.avatar_url,
            'theme_settings': profile.theme_settings,
            'documents': docs_data,
            'total_cl': profile.total_cl,
            'taken_cl': profile.taken_cl,
        }

        return Response({
            'success': True,
            'profile': profile_data
        })
    except Employee.DoesNotExist:
        return Response({
            'success': False,
            'message': 'Employee not found'
        }, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to load profile'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def check_profile_completeness(request):
    """Check if employee profile has all required fields and documents"""
    employee_id = request.GET.get('employee_id')
    if not employee_id:
        return Response({'success': False, 'message': 'Employee ID is required'}, status=400)

    try:
        employee = Employee.objects.get(id=employee_id)
        profile, _ = EmployeeProfile.objects.get_or_create(employee=employee)

        required_fields = [
            ('emergency_contact_name', 'Emergency Contact Name'),
            ('emergency_contact_phone', 'Emergency Contact Phone'),
            ('bank_account_number', 'Bank Account Number'),
            ('bank_ifsc', 'Bank IFSC'),
            ('bank_bank_name', 'Bank Name'),
            ('pan_number', 'PAN Number'),
            ('aadhar_number', 'Aadhar Number'),
            ('home_address', 'Home Address'),
            ('personal_email', 'Personal Email'),
            ('gender', 'Gender'),
            ('date_of_birth', 'Date of Birth'),
        ]

        missing_fields = []
        for field_name, display_name in required_fields:
            val = getattr(profile, field_name)
            if not val or str(val).strip() == '':
                missing_fields.append(display_name)

        # Check documents
        required_docs = ['aadhar', 'pan']
        uploaded_doc_types = EmployeeDocument.objects.filter(employee=employee).values_list('doc_type', flat=True)
        
        missing_docs = []
        if 'aadhar' not in uploaded_doc_types:
            missing_docs.append('Aadhar Card')
        if 'pan' not in uploaded_doc_types:
            missing_docs.append('PAN Card')

        is_complete = len(missing_fields) == 0 and len(missing_docs) == 0

        return Response({
            'success': True,
            'is_complete': is_complete,
            'missing_fields': missing_fields,
            'missing_docs': missing_docs,
            'message': 'Profile complete' if is_complete else 'Profile incomplete'
        })

    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'Employee not found'}, status=404)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=500)


@api_view(['GET'])
@require_gated_token_api
def admin_profiles_list(request):
    """List all employee profiles (admin)"""
    user = get_current_user(request, requested_id=request.GET.get('user_id'))
    if not user:
        return Response({'success': False, 'message': 'Unauthorized'}, status=403)
    
    # Include de-facto Mentors (any user who has subordinates)
    is_mentor = user and (user.role == 'mentor' or (user.role != 'admin' and user.subordinates.exists()))

    try:
        employees_qs = Employee.objects.filter(is_active=True)
        if is_mentor:
            employees_qs = employees_qs.filter(mentors=user)

        employees = employees_qs.select_related('profile')\
            .annotate(docs_count=Count('documents'))\
            .order_by('id')
        profiles_data = []

        for emp in employees:
            profile = getattr(emp, 'profile', None)
            profiles_data.append({
                'id': emp.id,
                'username': emp.username,
                'name': emp.name,
                'department': emp.department,
                'official_email': emp.email,
                'official_phone': emp.phone,
                'personal_email': profile.personal_email if profile else None,
                'gender': profile.gender if profile else None,
                'date_of_birth': str(profile.date_of_birth) if profile and profile.date_of_birth else None,
                'date_of_joining': str(profile.date_of_joining) if profile and profile.date_of_joining else None,
                'skill_set': profile.skill_set if profile else None,
                'reporting_mentor': profile.reporting_mentor if profile else None,
                'docs_count': emp.docs_count,
                'role': emp.role,
            })

        return Response({
            'success': True,
            'profiles': profiles_data
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to load profiles'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# Admin Views
@api_view(['GET'])
@require_gated_token_api
def admin_users(request):
    """Get all users (admin)"""
    user = get_current_user(request, requested_id=request.GET.get('user_id'))
    if not user:
        return Response({'success': False, 'message': 'Unauthorized'}, status=403)
    
    # Include de-facto Mentors (any user who has subordinates)
    is_mentor = user and (user.role == 'mentor' or (user.role != 'admin' and user.subordinates.exists()))

    try:
        # Include active task names and counts
        active_tasks_qs = Task.objects.filter(status__in=['todo', 'in_progress'])
        users = Employee.objects.all().order_by('-id').prefetch_related(
            'profile', 
            'mentors',
            Prefetch('assigned_tasks', queryset=active_tasks_qs, to_attr='active_tasks_list')
        ).annotate(
            active_tasks_count_db=Count('assigned_tasks', filter=Q(assigned_tasks__status__in=['todo', 'in_progress']), distinct=True)
        )
        
        if is_mentor:
            users = users.filter(mentors=user)
        
        users_data = []
        for u in users:
            dob = None
            gender = None
            if hasattr(u, 'profile') and u.profile:
                dob = str(u.profile.date_of_birth) if u.profile.date_of_birth else None
                gender = u.profile.gender
            
            users_data.append({
                'id': u.id,
                'username': u.username,
                'name': u.name,
                'email': u.email,
                'phone': u.phone,
                'department': u.department,
                'role': u.role,
                'Mentor_name': ", ".join([m.name for m in u.mentors.all()]) if u.mentors.all() else None,
                'is_active': u.is_active,
                'date_of_birth': dob,
                'gender': gender,
                'active_tasks': [{'id': t.id, 'title': t.title} for t in u.active_tasks_list] if hasattr(u, 'active_tasks_list') else [],
                'active_tasks_count': getattr(u, 'active_tasks_count_db', 0)
            })

        return Response({
            'success': True,
            'users': users_data
        })
    except Exception as e:
        import logging; logging.getLogger('attendance').exception('Error in admin_users')
        return Response({
            'success': False,
            'message': f'Failed to fetch users: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET', 'POST', 'DELETE'])
@require_gated_token_api
@parser_classes([JSONParser])
def admin_user_detail(request, user_id):
    """Get, update, or delete a user (admin)"""
    # Verify caller is an admin
    caller = get_current_user(request, require_admin=True)
    if not caller:
        return Response({'success': False, 'message': 'Admin access required'}, status=403)

    try:
        employee = Employee.objects.get(id=user_id)
    except Employee.DoesNotExist:
        return Response({
            'success': False,
            'message': 'User not found'
        }, status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        return Response({
            'success': True,
            'user': {
                'id': employee.id,
                'username': employee.username,
                'name': employee.name,
                'email': employee.email,
                'phone': employee.phone,
                'department': employee.department,
                'role': employee.role,
                'mentor_ids': [m.id for m in employee.mentors.all()],
                'Mentor_names': [m.name for m in employee.mentors.all()],
                'mentor_id': employee.mentors.all()[0].id if employee.mentors.exists() else None,
                'Mentor_name': employee.mentors.all()[0].name if employee.mentors.exists() else None,
                'is_active': employee.is_active,
                'total_cl': employee.profile.total_cl if hasattr(employee, 'profile') else 12,
                'taken_cl': employee.profile.taken_cl if hasattr(employee, 'profile') else 0,
                'date_of_joining': str(employee.profile.date_of_joining) if hasattr(employee, 'profile') and employee.profile.date_of_joining else None,
            }
        })

    elif request.method == 'POST':
        data = request.data

        # Check if delete
        if data.get('_method') == 'DELETE':
            employee.delete()
            return Response({
                'success': True,
                'message': 'User deleted'
            })

        # Update user
        if data.get('name'):
            employee.name = data['name']
        if data.get('email'):
            employee.email = data['email']
        if data.get('phone'):
            employee.phone = data['phone']
        if data.get('department'):
            employee.department = data['department']
        if data.get('role'):
            employee.role = data['role']
        if data.get('mentor_ids'):
            mentor_ids = data.get('mentor_ids')
            if isinstance(mentor_ids, list):
                if 'none' in mentor_ids:
                    employee.mentors.clear()
                else:
                    employee.mentors.set(Employee.objects.filter(id__in=mentor_ids))
        elif data.get('mentor_id'):
            if data['mentor_id'] == 'none':
                employee.mentors.clear()
            else:
                try:
                    Mentor_emp = Employee.objects.get(id=data['mentor_id'])
                    employee.mentors.set([Mentor_emp])
                except Employee.DoesNotExist:
                    pass
        elif 'mentor_id' in data and not data.get('mentor_id'):
            employee.mentors.clear()

        if 'is_active' in data:
            employee.is_active = bool(data['is_active'])
        
        # Update Profile fields
        profile, created = EmployeeProfile.objects.get_or_create(employee=employee)
        if 'total_cl' in data:
            profile.total_cl = int(data['total_cl'])
        if 'taken_cl' in data:
            profile.taken_cl = int(data['taken_cl'])
        if 'date_of_joining' in data:
            profile.date_of_joining = data['date_of_joining'] if data['date_of_joining'] else None
        profile.save()

        if data.get('primary_office'):
            employee.primary_office = data['primary_office']
        if data.get('password'):
            employee.password = make_password(data['password'])

        employee.save()
        return Response({
            'success': True,
            'message': 'User updated'
        })

    elif request.method == 'DELETE':
        employee.delete()
        return Response({
            'success': True,
            'message': 'User deleted'
        })


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def create_office(request):
    """Create a new office (admin)"""
    data = request.data

    if not data.get('id') or not data.get('name'):
        return Response({
            'success': False,
            'message': 'Office ID and Office name are required'
        }, status=status.HTTP_400_BAD_REQUEST)

    try:
        office = OfficeLocation.objects.create(
            id=data['id'],
            name=data['name'],
            address=data.get('address', ''),
            latitude=float(data['latitude']) if data.get('latitude') else None,
            longitude=float(data['longitude']) if data.get('longitude') else None,
            radius_meters=int(data.get('radius_meters') or data.get('radius') or 100),
            is_active=True
        )

        # Grant access to all departments
        departments = ['IT', 'HR', 'Surveyors', 'Accounts', 'Growth', 'Others']
        for dept in departments:
            DepartmentOfficeAccess.objects.get_or_create(
                department=dept,
                office=office
            )

        return Response({
            'success': True,
            'message': 'Office created',
            'office_id': office.id
        })
    except Exception as e:
        if 'UNIQUE constraint' in str(e) or 'Duplicate entry' in str(e):
            return Response({
                'success': False,
                'message': 'Failed to create office: That Office ID already exists.'
            }, status=status.HTTP_400_BAD_REQUEST)
        return Response({
            'success': False,
            'message': f'Failed to create office: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET', 'POST', 'DELETE'])
@require_gated_token_api
@parser_classes([JSONParser])
def office_detail(request, office_id):
    """Get, update, or delete an office"""
    try:
        office = OfficeLocation.objects.get(id=office_id)
    except OfficeLocation.DoesNotExist:
        return Response({
            'success': False,
            'message': 'Office not found'
        }, status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        return Response({
            'success': True,
            'office': {
                'id': office.id,
                'name': office.name,
                'address': office.address,
                'latitude': float(office.latitude),
                'longitude': float(office.longitude),
                'radius_meters': office.radius_meters,
                'is_active': office.is_active,
            }
        })

    elif request.method == 'POST':
        data = request.data

        # Check if delete
        if data.get('_method') == 'DELETE':
            office.delete()
            return Response({
                'success': True,
                'message': 'Office deleted successfully'
            })

        # Update office
        office.name = data.get('name', office.name)
        office.address = data.get('address', office.address)
        if data.get('latitude'):
            office.latitude = float(data['latitude'])
        if data.get('longitude'):
            office.longitude = float(data['longitude'])
        if data.get('radius_meters'):
            office.radius_meters = int(data['radius_meters'])
        office.save()

        return Response({
            'success': True,
            'message': 'Office updated successfully'
        })

    elif request.method == 'DELETE':
        office.delete()
        return Response({
            'success': True,
            'message': 'Office deleted successfully'
        })


@api_view(['GET', 'POST', 'DELETE'])
@require_gated_token_api
@parser_classes([JSONParser])
def attendance_record_detail(request, record_id):
    """Get, update, or delete an attendance record (admin only)"""
    # Verify caller is an admin
    caller = get_current_user(request, require_admin=True)
    if not caller:
        return Response({'success': False, 'message': 'Admin access required'}, status=403)

    try:
        record = AttendanceRecord.objects.get(id=record_id)
    except AttendanceRecord.DoesNotExist:
        return Response({
            'success': False,
            'message': 'Attendance record not found'
        }, status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        return Response({
            'success': True,
            'record': {
                'id': record.id,
                'employee_id': record.employee_id,
                'date': str(record.date),
                'check_in_time': str(record.check_in_time) if record.check_in_time else None,
                'check_out_time': str(record.check_out_time) if record.check_out_time else None,
                'type': record.type,
                'status': record.status,
                'office_id': record.office_id,
                'total_hours': float(record.total_hours),
            }
        })

    elif request.method == 'POST':
        data = request.data

        # Check if delete
        if data.get('_method') == 'DELETE':
            record.delete()
            return Response({
                'success': True,
                'message': 'Attendance deleted'
            })

        # Update record
        allowed_fields = ['status', 'type', 'date', 'check_in_time', 'check_out_time', 'office_id', 'notes']
        for field in allowed_fields:
            if field in data:
                setattr(record, field, data[field])

        record.save()
        return Response({
            'success': True,
            'message': 'Attendance updated'
        })

    elif request.method == 'DELETE':
        record.delete()
        return Response({
            'success': True,
            'message': 'Attendance deleted'
        })

# Document Upload Views

def _get_s3_client():
    """Returns (s3_client, bucket, prefix, use_s3).
    If AWS creds are configured returns a real boto3 client.
    If not, returns (None, None, None, False) so callers fall back to local disk.
    """
    import boto3
    aws_key    = getattr(settings, 'AWS_ACCESS_KEY_ID', None)
    aws_secret = getattr(settings, 'AWS_SECRET_ACCESS_KEY', None)
    S3_BUCKET  = getattr(settings, 'S3_BUCKET_NAME', 'attendance-g37k8w69fo65xagoo39ek3g58hc14aps3a-s3alias')
    S3_REGION  = getattr(settings, 'AWS_DEFAULT_REGION', 'ap-south-1')
    S3_PREFIX  = 'Employee_docs/'

    boto3_kwargs = {'region_name': S3_REGION}
    if aws_key and aws_secret:
        boto3_kwargs['aws_access_key_id']     = aws_key
        boto3_kwargs['aws_secret_access_key'] = aws_secret

    # Only attempt S3 when at least one credential signal is present
    use_s3 = bool(aws_key and aws_secret)
    if not use_s3:
        return None, S3_BUCKET, S3_PREFIX, False

    try:
        client = boto3.client('s3', **boto3_kwargs)
        return client, S3_BUCKET, S3_PREFIX, True
    except Exception:
        return None, S3_BUCKET, S3_PREFIX, False


@api_view(['POST'])
@require_gated_token_api
@parser_classes([MultiPartParser, FormParser])
def upload_documents(request):
    """Upload employee documents — S3 when credentials are set, local disk otherwise."""
    employee_id = request.POST.get('employee_id')
    username    = request.POST.get('username')

    if not employee_id or not username:
        return Response({'success': False, 'message': 'employee_id and username are required'},
                        status=status.HTTP_400_BAD_REQUEST)

    # Ownership check: ensure caller can only upload for themselves (or admin for anyone)
    caller = get_current_user(request, requested_id=employee_id)
    if not caller:
        return Response({'success': False, 'message': 'Unauthorized'}, status=403)

    try:
        employee = Employee.objects.get(id=employee_id)
    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'Employee not found'},
                        status=status.HTTP_404_NOT_FOUND)

    MAX_PHOTO_SIZE = 2 * 1024 * 1024  # 2 MB
    MAX_PDF_SIZE   = 5 * 1024 * 1024  # 5 MB

    saved_files = []
    s3_client, S3_BUCKET, S3_PREFIX, use_s3 = _get_s3_client()

    # --- local-disk fallback setup ---
    import os
    storage_root   = getattr(settings, 'DOCUMENT_STORAGE_ROOT', settings.MEDIA_ROOT)
    safe_name      = "".join(c if c.isalnum() or c in ' _-' else '' for c in employee.name).strip().replace(' ', '_')
    employee_folder = f"{safe_name}_{employee.id}"
    upload_dir     = os.path.join(storage_root, 'Documents', employee_folder)
    if not use_s3:
        os.makedirs(upload_dir, exist_ok=True)

    def _save_file(file, filename, doc_type, content_type):
        """Save to S3 or local disk; return the file_path string stored in DB."""
        if use_s3:
            s3_key = f"{S3_PREFIX}{username}/{filename}"
            s3_client.upload_fileobj(file, S3_BUCKET, s3_key,
                                     ExtraArgs={'ContentType': content_type})
            return s3_key
        else:
            local_path = os.path.join(upload_dir, filename)
            with open(local_path, 'wb') as fh:
                for chunk in file.chunks():
                    fh.write(chunk)
            return f'Documents/{employee_folder}/{filename}'

    # ── Photo / Signature ──────────────────────────────────────────────────────
    image_docs = {'user_photo': 'photo', 'user_signature': 'signature'}
    for input_name, doc_type in image_docs.items():
        if input_name not in request.FILES:
            continue
        file = request.FILES[input_name]
        if file.size > MAX_PHOTO_SIZE:
            return Response({'success': False,
                             'message': f'{doc_type.capitalize()} size exceeds 2 MB limit'},
                            status=status.HTTP_400_BAD_REQUEST)
        if file.content_type not in ['image/jpeg', 'image/png', 'image/jpg']:
            continue

        ext      = os.path.splitext(file.name)[1].lower()
        filename = f"{username}_{doc_type}{ext}"
        EmployeeDocument.objects.filter(employee_id=employee_id, doc_type=doc_type).delete()
        try:
            file_path = _save_file(file, filename, doc_type, file.content_type)
        except Exception as e:
            return Response({'success': False, 'message': f'Upload failed: {str(e)}'}, status=500)

        EmployeeDocument.objects.create(
            employee_id=employee_id, doc_type=doc_type,
            doc_name=doc_type.capitalize(), file_name=filename, file_path=file_path)
        saved_files.append(filename)

    # ── PDF documents ──────────────────────────────────────────────────────────
    pdf_docs = ['aadhar', 'pan', 'other_id', 'highest_qualification',
                'professional_certificate', 'other_qualification']
    for doc_type in pdf_docs:
        file_key = f'file_{doc_type}'
        if file_key not in request.FILES:
            continue
        file = request.FILES[file_key]
        if file.size > MAX_PDF_SIZE:
            return Response({'success': False,
                             'message': f'{doc_type.capitalize()} file exceeds 5 MB limit'},
                            status=status.HTTP_400_BAD_REQUEST)
        if file.content_type != 'application/pdf':
            continue

        filename = request.POST.get(f'{file_key}_filename', f"{username}_{doc_type}.pdf")
        filename = ''.join(c if c.isalnum() or c in '._-' else '_' for c in filename)
        EmployeeDocument.objects.filter(employee_id=employee_id, doc_type=doc_type).delete()
        try:
            file_path = _save_file(file, filename, doc_type, 'application/pdf')
        except Exception as e:
            return Response({'success': False, 'message': f'Upload failed: {str(e)}'}, status=500)

        EmployeeDocument.objects.create(
            employee_id=employee_id, doc_type=doc_type,
            doc_name=doc_type.replace('_', ' ').title(),
            doc_number=request.POST.get(f'doc{doc_type.capitalize()}Number', ''),
            file_name=filename, file_path=file_path)
        saved_files.append(filename)

    if not saved_files:
        return Response({'success': False, 'message': 'No valid documents uploaded'},
                        status=status.HTTP_400_BAD_REQUEST)

    return Response({'success': True, 'uploaded': saved_files,
                     'message': 'Documents uploaded successfully'})


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def delete_documents(request):
    """Delete selected documents (admin only)"""
    # Verify caller is admin
    caller = get_current_user(request, require_admin=True)
    if not caller:
        return Response({'success': False, 'message': 'Admin access required'}, status=403)

    data    = request.data
    doc_ids = data.get('document_ids', [])

    if not doc_ids:
        return Response({'success': False, 'message': 'No documents selected'},
                        status=status.HTTP_400_BAD_REQUEST)

    try:
        documents = EmployeeDocument.objects.filter(id__in=doc_ids)
        s3_client, S3_BUCKET, S3_PREFIX, use_s3 = _get_s3_client()

        for doc in documents:
            if doc.file_path.startswith(S3_PREFIX) and use_s3:
                try:
                    s3_client.delete_object(Bucket=S3_BUCKET, Key=doc.file_path)
                except Exception:
                    pass
            elif not doc.file_path.startswith(S3_PREFIX):
                import os
                storage_root = getattr(settings, 'DOCUMENT_STORAGE_ROOT', settings.MEDIA_ROOT)
                file_path = os.path.join(storage_root, doc.file_path)
                if os.path.exists(file_path):
                    os.remove(file_path)

        documents.delete()
        return Response({'success': True, 'message': 'Documents deleted successfully'})
    except Exception as e:
        return Response({'success': False, 'message': 'Failed to delete documents'},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def admin_user_docs_list(request, employee_id):
    """List documents for a user (admin only)"""
    # Verify caller is admin or the employee themselves
    caller = get_current_user(request, requested_id=employee_id)
    if not caller:
        return Response({'success': False, 'message': 'Unauthorized'}, status=403)

    try:
        documents = EmployeeDocument.objects.filter(employee_id=employee_id).order_by('-uploaded_at')
        docs_data = []

        for doc in documents:
            docs_data.append({
                'id': doc.id,
                'doc_type': doc.doc_type,
                'doc_name': doc.doc_name,
                'file_name': doc.file_name,
                'file_path': doc.file_path,
                'url': request.build_absolute_uri(f'/api/serve-document/{doc.id}'),
                'uploaded_at': doc.uploaded_at.isoformat() if doc.uploaded_at else None,
            })

        return Response({
            'success': True,
            'documents': docs_data
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to load documents'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def admin_user_docs_zip(request, employee_id):
    """Download all documents as ZIP (admin only)"""
    # Verify caller is admin
    caller = get_current_user(request, require_admin=True)
    if not caller:
        return Response({'success': False, 'message': 'Admin access required'}, status=403)

    try:
        employee  = Employee.objects.get(id=employee_id)
        documents = EmployeeDocument.objects.filter(employee_id=employee_id)

        if not documents.exists():
            return Response({'success': False, 'message': 'No documents found'},
                            status=status.HTTP_404_NOT_FOUND)

        import tempfile, zipfile, os
        zip_name  = f"{employee.username}_documents.zip"
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.zip')

        s3_client, S3_BUCKET, S3_PREFIX, use_s3 = _get_s3_client()

        with zipfile.ZipFile(temp_file.name, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for doc in documents:
                if doc.file_path.startswith(S3_PREFIX) and use_s3:
                    try:
                        s3_obj = s3_client.get_object(Bucket=S3_BUCKET, Key=doc.file_path)
                        zipf.writestr(doc.file_name, s3_obj['Body'].read())
                    except Exception:
                        pass
                elif not doc.file_path.startswith(S3_PREFIX):
                    storage_root = getattr(settings, 'DOCUMENT_STORAGE_ROOT', settings.MEDIA_ROOT)
                    file_path = os.path.join(storage_root, doc.file_path)
                    if os.path.exists(file_path):
                        zipf.write(file_path, doc.file_name)

        response = FileResponse(open(temp_file.name, 'rb'), content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="{zip_name}"'
        # Schedule temp file cleanup after response is sent
        import atexit
        atexit.register(lambda path=temp_file.name: os.unlink(path) if os.path.exists(path) else None)
        return response

    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'User not found'},
                        status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'success': False, 'message': 'Failed to create ZIP'},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def serve_document(request, doc_id):
    """Serve a document — from S3 if configured, otherwise from local disk."""
    # Verify caller is authenticated
    caller = get_current_user(request)
    if not caller:
        return Response({'success': False, 'message': 'Unauthorized'}, status=403)

    try:
        doc = EmployeeDocument.objects.get(id=doc_id)
        
        # Ownership check: only the doc owner or admin can access
        if doc.employee_id != caller.id and caller.role != 'admin':
            return Response({'success': False, 'message': 'Unauthorized'}, status=403)

        S3_PREFIX = 'Employee_docs/'
        if doc.file_path.startswith(S3_PREFIX):
            # Try S3 first
            s3_client, S3_BUCKET, _, use_s3 = _get_s3_client()
            if use_s3:
                from urllib.parse import quote
                try:
                    s3_obj     = s3_client.get_object(Bucket=S3_BUCKET, Key=doc.file_path)
                    ctype      = s3_obj.get('ContentType', 'application/octet-stream')
                    response   = FileResponse(s3_obj['Body'], content_type=ctype)
                    response['Content-Disposition'] = f'inline; filename="{quote(doc.file_name)}"'
                    return response
                except Exception:
                    return Response({'error': 'File not found in S3'}, status=404)
            else:
                return Response({'error': 'S3 not configured and file is stored in S3'}, status=404)
        else:
            import os, mimetypes
            storage_root = getattr(settings, 'DOCUMENT_STORAGE_ROOT', settings.MEDIA_ROOT)
            file_path    = os.path.join(storage_root, doc.file_path)
            if not os.path.exists(file_path):
                return Response({'error': 'File not found'}, status=404)
            ctype, _ = mimetypes.guess_type(file_path)
            return FileResponse(open(file_path, 'rb'),
                                content_type=ctype or 'application/octet-stream')
    except EmployeeDocument.DoesNotExist:
        return Response({'error': 'Document not found'}, status=404)
    except Exception as e:
        return Response({'error': str(e)}, status=500)


# Admin Dashboard API Views
@api_view(['GET'])
@require_gated_token_api
def admin_summary(request):
    """Get admin dashboard summary for a specific date (defaults to today)"""
    # Secure identity: use token user
    user = get_current_user(request, requested_id=request.GET.get('user_id'))
    if not user:
        return Response({'success': False, 'message': 'Unauthorized'}, status=403)
    
    date_param = request.GET.get('date')
    # Include de-facto Mentors (any user who has subordinates)
    is_mentor = user and (user.role == 'mentor' or (user.role != 'admin' and user.subordinates.exists()))

    try:
        # Determine the target date (default to today IST)
        if date_param:
            target_date = datetime.strptime(date_param, '%Y-%m-%d').date()
        else:
            target_date = timezone.localtime(timezone.now()).date()

        # Phase 1: Fetch all active employees (excluding admins)
        employees_qs = Employee.objects.filter(is_active=True).exclude(role='admin')
        if is_mentor:
            employees_qs = employees_qs.filter(mentors=user)
        
        all_employees = list(employees_qs.values('id', 'name', 'department'))
        total_employees = len(all_employees)
        employee_id_to_name = {e['id']: e['name'] for e in all_employees}
        employee_ids = set(employee_id_to_name.keys())

        # Phase 2: Fetch all attendance records for the target date
        records_qs = AttendanceRecord.objects.filter(date=target_date, employee_id__in=employee_ids)
        all_records = list(records_qs.values('employee_id', 'status', 'type', 'employee__name', 'employee__department'))

        # Phase 3: Categorize in memory (O(N) instead of multiple O(N) queries)
        present_names = []
        leave_names = []
        wfh_names = []
        marked_ids = set()

        surveyors_total = 0
        surveyors_office_names = []
        surveyors_client_names = []
        surveyors_wfh_names = []
        surveyors_leave_names = []
        surveyors_marked_ids = set()

        # Count total surveyors
        for e in all_employees:
            if e['department'] == 'Surveyors':
                surveyors_total += 1

        for r in all_records:
            emp_id = r['employee_id']
            status = r['status']
            att_type = r['type']
            name = r['employee__name']
            dept = r['employee__department']
            
            marked_ids.add(emp_id)
            
            if status in ['present', 'half_day', 'wfh', 'client']:
                present_names.append(name)
            
            if status == 'leave':
                leave_names.append(name)
            
            if status == 'wfh':
                wfh_names.append(name)
                
            # Surveyor specifics
            if dept == 'Surveyors':
                surveyors_marked_ids.add(emp_id)
                if status == 'leave':
                    surveyors_leave_names.append(name)
                elif status == 'client' or (status in ['present', 'half_day'] and att_type == 'client'):
                    surveyors_client_names.append(name)
                elif status == 'wfh' or (status in ['present', 'half_day'] and att_type == 'wfh'):
                    surveyors_wfh_names.append(name)
                elif status in ['present', 'half_day'] and att_type == 'office':
                    surveyors_office_names.append(name)

        # Calculate absent
        absent_names = [e['name'] for e in all_employees if e['id'] not in marked_ids]
        surveyors_absent_names = [e['name'] for e in all_employees if e['department'] == 'Surveyors' and e['id'] not in surveyors_marked_ids]

        surveyors_active = len(surveyors_office_names) + len(surveyors_client_names) + len(surveyors_wfh_names)

        return Response({
            'success': True,
            'date': str(target_date),
            'total_employees': total_employees,
            'present_today': len(present_names),
            'present_names': present_names,
            'absent_today': len(absent_names),
            'absent_names': absent_names,
            'on_leave': len(leave_names),
            'leave_names': leave_names,
            'wfh_today': len(wfh_names),
            'wfh_names': wfh_names,
            'surveyors_total': surveyors_total,
            'surveyors_present': surveyors_active,
            'surveyors_office': len(surveyors_office_names),
            'surveyors_office_names': surveyors_office_names,
            'surveyors_client': len(surveyors_client_names),
            'surveyors_client_names': surveyors_client_names,
            'surveyors_wfh': len(surveyors_wfh_names),
            'surveyors_wfh_names': surveyors_wfh_names,
            'surveyors_leave': len(surveyors_leave_names),
            'surveyors_leave_names': surveyors_leave_names,
            'surveyors_absent': len(surveyors_absent_names),
            'surveyors_absent_names': surveyors_absent_names,
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({
            'success': False,
            'message': f'Failed to fetch admin summary: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def predict_attendance(request):
    """Predict attendance for tomorrow based on historical patterns"""
    try:
        today = date.today()
        tomorrow = today + timedelta(days=1)

        # We look at historical data for the same day of week as tomorrow
        tomorrow_dow = tomorrow.weekday() # 0=Mon, 6=Sun

        # Total active employees
        total_employees = Employee.objects.filter(is_active=True).count()
        if total_employees == 0:
            return Response({'success': True, 'predicted_count': 0, 'confidence': 0, 'trend': 'stable'})

        # Get records for same DOW over last 4 weeks
        history_dates = [tomorrow - timedelta(weeks=i) for i in range(1, 5)]

        counts = []
        for h_date in history_dates:
            present_count = AttendanceRecord.objects.filter(
                date=h_date,
                status__in=['present', 'half_day', 'wfh', 'client']
            ).count()
            if present_count > 0 or AttendanceRecord.objects.filter(date=h_date).exists():
                counts.append(present_count)

        if not counts:
            # Fallback to general daily average if no DOW specific data
            all_recent = AttendanceRecord.objects.filter(
                date__gte=today - timedelta(days=30)
            ).values('date').annotate(count=Count('id', filter=Q(status__in=['present', 'half_day', 'wfh', 'client'])))

            counts = [item['count'] for item in all_recent]

        if not counts:
            return Response({
                'success': True,
                'predicted_count': round(total_employees * 0.8),
                'predicted_percent': 80,
                'confidence': 30,
                'trend': 'stable',
                'message': 'Insufficient data for accurate prediction'
            })

        avg_predicted = sum(counts) / len(counts)
        predicted_percent = (avg_predicted / total_employees) * 100 if total_employees > 0 else 0

        # Calculate Trend: Compare last 7 days vs previous 7 days
        last_7_days = today - timedelta(days=7)
        prev_7_days = today - timedelta(days=14)

        # Formula: Average = Total / Number of working days in a week
        # Over a 7-day period, we assume 5 working days
        current_avg = AttendanceRecord.objects.filter(
            date__gte=last_7_days,
            status__in=['present', 'half_day', 'wfh', 'client']
        ).count() / 5

        previous_avg = AttendanceRecord.objects.filter(
            date__gte=prev_7_days,
            date__lt=last_7_days,
            status__in=['present', 'half_day', 'wfh', 'client']
        ).count() / 5

        if current_avg > previous_avg * 1.05:
            trend = 'up'
        elif current_avg < previous_avg * 0.95:
            trend = 'down'
        else:
            trend = 'stable'

        # Get last 7 days of actual counts for visualization
        recent_history = []
        for i in range(7):
            d = today - timedelta(days=i)
            count = AttendanceRecord.objects.filter(
                date=d,
                status__in=['present', 'half_day', 'wfh', 'client']
            ).count()
            recent_history.append({
                'date': d.strftime('%Y-%m-%d'),
                'day': d.strftime('%a'),
                'count': count
            })
        recent_history.reverse()

        confidence = min(len(counts) * 20 + 20, 95) # Simple confidence score

        return Response({
            'success': True,
            'predicted_count': round(avg_predicted),
            'predicted_percent': round(predicted_percent, 1),
            'confidence': confidence,
            'trend': trend,
            'tomorrow_day': tomorrow.strftime('%A'),
            'recent_history': recent_history,
            'daily_average': round(current_avg, 1)
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def employee_performance_analysis(request, employee_id):
    """Detailed performance and prediction analysis for a single employee"""
    try:
        employee = Employee.objects.get(id=employee_id)
        today = date.today()
        
        # Filtering Logic
        view_type = request.GET.get('view_type', 'period') # period, month, week
        month_param = request.GET.get('month')
        year_param = request.GET.get('year')
        week_param = request.GET.get('week')
        
        is_monthly_view = False
        is_weekly_view = False
        
        if view_type == 'month':
            try:
                view_month = int(month_param) if month_param else today.month
                view_year = int(year_param) if year_param else today.year
                week_idx = request.GET.get('week_idx') # Optional: 1, 2, 3, 4, 5
                
                start_date = date(view_year, view_month, 1)
                if view_month == 12:
                    last_day = (date(view_year + 1, 1, 1) - timedelta(days=1)).day
                else:
                    last_day = (date(view_year, view_month + 1, 1) - timedelta(days=1)).day
                
                end_date = date(view_year, view_month, last_day)

                if week_idx and week_idx != 'all':
                    w = int(week_idx)
                    s_day = (w - 1) * 7 + 1
                    e_day = min(w * 7, last_day)
                    
                    if s_day <= last_day:
                        start_date = date(view_year, view_month, s_day)
                        end_date = date(view_year, view_month, e_day)
                
                is_monthly_view = True
            except (ValueError, TypeError):
                start_date = today - timedelta(days=30)
                end_date = today
        else:
            # Default: period (Last 30 Days)
            start_date = today - timedelta(days=30)
            end_date = today

        # Attendance History for filtered period
        records = AttendanceRecord.objects.filter(
            employee=employee,
            date__range=[start_date, end_date]
        ).order_by('-date')

        history = []
        now_local = timezone.localtime(timezone.now())
        
        for r in records:
            hours = float(r.total_hours)
            
            # If currently checked in but not checked out, calculate hours so far
            if not r.check_out_time and r.check_in_time and r.date == now_local.date():
                try:
                    check_in_t = datetime.strptime(str(r.check_in_time), '%H:%M:%S').time()
                    check_in_dt = timezone.make_aware(datetime.combine(r.date, check_in_t))
                    # Calculate hours since check-in
                    hours = round((now_local - check_in_dt).total_seconds() / 3600, 2)
                    # Cap at a reasonable max (e.g. 14h) to avoid outliers if they forgot to check out yesterday 
                    # (though r.date == now_local.date() handles today)
                    hours = max(0.0, min(hours, 14.0))
                except Exception:
                    pass
            
            history.append({
                'date': r.date.strftime('%Y-%m-%d'),
                'status': r.status,
                'type': r.type,
                'hours': hours
            })

        # 2. Performance Metrics
        num_days = (end_date - start_date).days + 1
        num_weeks = max(num_days / 7.0, 0.1) # Avoid division by zero, min 0.1 weeks
        
        # Calculate working days passed for regularity
        calc_end_date = min(end_date, today)
        if start_date <= calc_end_date:
            passed_days = (calc_end_date - start_date).days + 1
            working_days_passed = sum(1 for d in range(passed_days) if Holiday.is_date_working(start_date + timedelta(days=d)))
        else:
            working_days_passed = 0

        weekday_present_days = records.filter(
            date__week_day__in=[2, 3, 4, 5, 6, 7],
            status__in=['present', 'half_day', 'wfh', 'client']
        ).count()

        
        # Calculate Mon-Fri Avg
        weekday_records = records.filter(
            date__week_day__in=[2, 3, 4, 5, 6, 7] # Mon-Sat (2=Mon... 7=Sat)
        ).aggregate(
            sum_hours=Sum('total_hours')
        )
        # Fixed denominator logic as per user request: (num_weeks * 5)
        total_weekday_hours = float(weekday_records['sum_hours'] or 0)
        weekday_avg = total_weekday_hours / (num_weeks * 6)

        # Calculate Sat-Sun Avg
        weekend_records = records.filter(
            date__week_day__in=[1] # Sun (1) only
        ).aggregate(
            sum_hours=Sum('total_hours')
        )
        # Fixed denominator logic: (num_weeks * 1)
        total_weekend_hours = float(weekend_records['sum_hours'] or 0)
        saturday_avg = total_weekend_hours / (num_weeks * 1)



        summary_stats = records.aggregate(
            total_present=Count('id', filter=Q(status__in=['present', 'half_day', 'wfh', 'client'])),
            sum_hours=Sum('total_hours'),
            wfh_count=Count('id', filter=Q(type='wfh', status__in=['present', 'half_day', 'wfh', 'client'])),
            office_count=Count('id', filter=Q(type='office', status__in=['present', 'half_day', 'wfh', 'client']))
        )

        # Sanitize: cap each present record's hours at 14h before computing any
        # hour-based metric.  This prevents corrupted checkout records (where the
        # fallback logic picked up an unclosed record from a previous day and
        # computed hours spanning multiple calendar days, capped at 99.90h in the
        # DB) from inflating weekly/daily averages into the hundreds.
        MAX_HOURS_PER_DAY = 14.0
        capped_hours_sum = sum(
            min(float(r.total_hours or 0), MAX_HOURS_PER_DAY)
            for r in records
            if r.status in ['present', 'half_day', 'wfh', 'client']
        )
        # Keep the raw DB sum available for other uses (regular/OT split below)
        total_hours_sum = float(summary_stats['sum_hours'] or 0)

        # Weekly average uses the sanitized sum and the actual number of weeks
        # in the filtered period (never hardcoded 4).
        if is_monthly_view:
            weekly_avg_hours = capped_hours_sum / 4.33
        elif is_weekly_view:
            weekly_avg_hours = capped_hours_sum  # Single week: total IS the average
        else:
            weekly_avg_hours = capped_hours_sum / num_weeks

        # Forecast for tomorrow (always uses global patterns)
        tomorrow = date.today() + timedelta(days=1)
        tomorrow_dow = (tomorrow.weekday() + 1) % 7 + 1 
        habit_records = list(AttendanceRecord.objects.filter(
            employee=employee,
            date__week_day=tomorrow_dow
        ).order_by('-date')[:8]) 

        if habit_records:
            present_in_habit = len([r for r in habit_records if r.status in ['present', 'half_day', 'wfh', 'client']])
            prediction_score = (present_in_habit / len(habit_records)) * 100
        else:
            prediction_score = 85.0

        # Create predictive graph data for individual
        predict_days = int(request.GET.get('predict_days', 3))
        history_days = 3
        history_points = []

        graph_dates = []
        for i in range(history_days, 0, -1):
            graph_dates.append(today - timedelta(days=i))
        graph_dates.append(today)

        for d in graph_dates:
            r = AttendanceRecord.objects.filter(employee=employee, date=d).first()
            if r:
                hours = float(r.total_hours or 0)
                if not r.check_out_time and r.check_in_time and r.date == now_local.date():
                    try:
                        check_in_t = datetime.strptime(str(r.check_in_time), '%H:%M:%S').time()
                        check_in_dt = timezone.make_aware(datetime.combine(r.date, check_in_t))
                        hours = round((now_local - check_in_dt).total_seconds() / 3600, 2)
                        hours = max(0.0, min(hours, 14.0))
                    except: pass
                
                if r.status == 'half_day' or getattr(r, 'is_half_day', False):
                    hours = max(hours, 4.5)
            else:
                hours = 0.0

            if d == today:
                day_name = 'Today'
            elif d == today - timedelta(days=1):
                day_name = 'Yesterday'
            else:
                day_name = d.strftime('%A')
            
            history_points.append({
                'date': d.strftime('%Y-%m-%d'),
                'day_name': day_name,
                'hours': hours,
                'is_prediction': False
            })

        # Calculate Prediction Peak scaling from history to avoid "stuck at 8h"
        max_seen = max([p['hours'] for p in history_points] + [8.0])
        limit_hours = min(12.0, max_seen)

        from .intelligence_hub import calculate_multi_day_forecast, IndividualPredictor
        multi_forecast = calculate_multi_day_forecast(predict_days)
        org_forecast_map = {f['date']: f['rate'] for f in multi_forecast}
        individual_engine = IndividualPredictor()

        graph_data = history_points.copy()

        for i in range(1, predict_days + 1):
            target_date = today + timedelta(days=i)
            target_date_str = target_date.strftime('%Y-%m-%d')
            current_org_forecast = org_forecast_map.get(target_date_str, 85.0)
            
            base_prob = individual_engine.predict(employee, current_org_forecast, target_date=target_date)
            
            if target_date.weekday() >= 5:
                pred_hours = 0.0
            else:
                # Real-time logic: Predict hours based on DOW historical pattern + attendance probability
                pred_hours = individual_engine.predict_hours(employee, base_prob, target_date)
            
            graph_data.append({
                'date': target_date_str,
                'day_name': target_date.strftime('%A'),
                'hours': round(pred_hours, 1),
                'is_prediction': True
            })

        # Attendance Habits (Averages for filtered period)
        attendance_with_time = records.filter(check_in_time__isnull=False)
        
        avg_check_in = None
        avg_check_out = None
        
        if attendance_with_time.exists():
            in_seconds = []
            out_seconds = []
            for r in attendance_with_time:
                in_seconds.append(r.check_in_time.hour * 3600 + r.check_in_time.minute * 60 + r.check_in_time.second)
                if r.check_out_time:
                    out_seconds.append(r.check_out_time.hour * 3600 + r.check_out_time.minute * 60 + r.check_out_time.second)
            
            if in_seconds:
                avg_in_sec = sum(in_seconds) / len(in_seconds)
                avg_check_in = f"{int(avg_in_sec // 3600):02d}:{int((avg_in_sec % 3600) // 60):02d}"
            
            if out_seconds:
                avg_out_sec = sum(out_seconds) / len(out_seconds)
                avg_check_out = f"{int(avg_out_sec // 3600):02d}:{int((avg_out_sec % 3600) // 60):02d}"
        
        # Task Management Performance (for filtered period)
        # Improved: Include tasks that were either created in this range, completed in this range, OR are currently active
        tasks_base = Task.objects.filter(assignees=employee).filter(
            Q(created_at__date__range=[start_date, end_date]) |
            Q(completed_at__date__range=[start_date, end_date]) |
            Q(status__in=['todo', 'in_progress'])
        ).distinct()
        completed_tasks = tasks_base.filter(status='completed')
        
        # New Advanced Accuracy Logic
        total_accuracy_points = 0
        tasks_evaluated = 0
        total_span_hours = 0
        spans_counted = 0

        for t in completed_tasks:
            task_score = 0
            
            # 1. Response Speed (Created to Started) - 30% Weight
            if t.started_at:
                response_delta = (t.started_at - t.created_at).total_seconds() / 3600
                if response_delta <= 2: task_score += 30
                elif response_delta <= 6: task_score += 25
                elif response_delta <= 12: task_score += 20
                elif response_delta <= 24: task_score += 15
                else: task_score += 5
            else:
                task_score += 10 # Default minimum

            # 2. Task Span (Started to Completed) - 35% Weight
            if t.started_at and t.completed_at:
                span_delta = (t.completed_at - t.started_at).total_seconds() / 3600
                total_span_hours += span_delta
                spans_counted += 1
                
                if span_delta <= 8: task_score += 35
                elif span_delta <= 24: task_score += 30
                elif span_delta <= 48: task_score += 25
                elif span_delta <= 72: task_score += 15
                else: task_score += 5
            else:
                task_score += 10

            # 3. Deadline Punctuality (Completed to Due Date) - 35% Weight
            if t.due_date and t.completed_at:
                # Treat due_date as end of day
                due_datetime = timezone.make_aware(datetime.combine(t.due_date, time(23, 59, 59)))
                days_diff = (due_datetime - t.completed_at).days
                
                if days_diff >= 2: task_score += 35 # Finished 2+ days early
                elif days_diff >= 1: task_score += 32 # Finished 1 day early
                elif days_diff == 0:
                    if t.completed_at <= due_datetime: task_score += 28 # Finished on due date
                    else: task_score += 15 # Slightly late
                elif days_diff == -1: task_score += 10 # 1 day late
                else: task_score += 0 # 2+ days late
            else:
                task_score += 20 # Neutral score if no due date set

            # Blend with manual Mentor accuracy if it exists (50/50 balance)
            if t.accuracy:
                task_score = (task_score + t.accuracy) / 2

            total_accuracy_points += task_score
            tasks_evaluated += 1

        avg_accuracy = total_accuracy_points / tasks_evaluated if tasks_evaluated > 0 else (70.0 if tasks_base.filter(status='in_progress').exists() else 0.0)
        avg_span_h = total_span_hours / (spans_counted or 1)
        
        total_assigned = tasks_base.count()
        completed_count = completed_tasks.count()
        in_progress_count = tasks_base.filter(status='in_progress').count()

        # Weighted Completion logic: Completed tasks (1.0) + In Progress (0.5)
        weighted_completion = completed_count + (in_progress_count * 0.5)
        completion_rate = (weighted_completion / total_assigned * 100) if total_assigned > 0 else 100.0
        
        work_efficiency = (completion_rate * 0.4) + (avg_accuracy * 0.6)

        task_stats = {
            'total_assigned': total_assigned,
            'todo': tasks_base.filter(status='todo').count(),
            'in_progress': tasks_base.filter(status='in_progress').count(),
            'completed': completed_count,
            'avg_accuracy': round(float(avg_accuracy), 1),
            'work_efficiency': round(float(work_efficiency), 1),
            'avg_span_hours': round(float(avg_span_h), 1)
        }

        # Calculate Regular vs Overtime Hours (Standard 8h)
        total_reg_h = 0
        total_ot_h = 0
        for r in records:
            h = float(r.total_hours or 0)
            reg = min(h, 8.0)
            ot = max(0.0, h - 8.0)
            total_reg_h += reg
            total_ot_h += ot
        
        total_all_h = total_reg_h + total_ot_h
        ot_ratio = round((total_ot_h / total_all_h) * 100, 1) if total_all_h > 0 else 0
        reg_ratio = round((total_reg_h / total_all_h) * 100, 1) if total_all_h > 0 else 0

        profile = employee.profile if hasattr(employee, 'profile') else None

        # Calculate Peak Day (Best Day) based on history (Mon-Fri only)
        dow_counts = {0:0, 1:0, 2:0, 3:0, 4:0, 5:0} # Mon-Sat
        for r in records:
            if r.status in ['present', 'half_day', 'wfh', 'client'] and r.date.weekday() < 6:
                dow_counts[r.date.weekday()] += 1
        
        best_dow = max(dow_counts, key=dow_counts.get) if any(dow_counts.values()) else 0
        day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
        peak_day_individual = day_names[best_dow]

        # Advanced Predictions from Engine
        from .attendance_prediction import AttendancePredictionEngine
        engine = AttendancePredictionEngine(employee.id)
        leave_probs = engine.predict_leaves()
        tomorrow_leave_prob = leave_probs.get(tomorrow.weekday(), 0)
        predicted_hrs = engine.predict_working_hours()

        return Response({
            'success': True,
            'employee_name': employee.name,
            'department': employee.department,
            'email': employee.email,
            'avatar_emoji': profile.avatar_emoji if profile else "👤",
            'history': history,
            'filter': {
                'start_date': str(start_date),
                'end_date': str(end_date),
                'month': start_date.month if is_monthly_view else None,
                'year': start_date.year if is_monthly_view else None,
                'week_idx': request.GET.get('week_idx', 'all'),
                'view_type': view_type
            },
            'metrics': {
                'total_present': summary_stats['total_present'] or 0,
                'avg_hours_present': round(capped_hours_sum / (summary_stats['total_present'] or 1), 1),
                'weekday_avg': round(weekday_avg, 1),
                'saturday_avg': round(saturday_avg, 1),

                'wfh_ratio': round((summary_stats['wfh_count'] / (summary_stats['total_present'] or 1)) * 100, 1) if summary_stats['total_present'] else 0,
                'office_ratio': round((summary_stats['office_count'] / (summary_stats['total_present'] or 1)) * 100, 1) if summary_stats['total_present'] else 0,
                'ot_ratio': ot_ratio,
                'reg_ratio': reg_ratio,
                'total_reg_h': round(total_reg_h, 1),
                'total_ot_h': round(total_ot_h, 1),
                'weekly_avg_hours': round(weekly_avg_hours, 1),
                'avg_check_in': avg_check_in,
                'avg_check_out': avg_check_out,
                'working_days_passed': working_days_passed,
                'weekday_present_days': weekday_present_days
            },
            'tasks': task_stats,
            'prediction': {
                'likelihood': round(prediction_score, 1),
                'tomorrow_day': tomorrow.strftime('%A'),
                'peak_day': peak_day_individual,
                'habit_summary': f"Usually present on {tomorrow.strftime('%A')}s" if prediction_score > 70 else f"Irregular pattern on {tomorrow.strftime('%A')}s",
                'graph_data': graph_data,
                'leave_probability': tomorrow_leave_prob,
                'predicted_daily_hours': predicted_hrs
            }
        })
    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'Employee not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def upcoming_birthdays(request):
    """Get upcoming birthdays for filtered month"""
    try:
        today = date.today()

        try:
            current_month = int(request.GET.get('month', today.month))
            current_year = int(request.GET.get('year', today.year))
        except ValueError:
            current_month = today.month
            current_year = today.year

        # Get employees with birthdays in filtered month
        employees_with_birthdays = EmployeeProfile.objects.filter(
            date_of_birth__month=current_month,
            employee__is_active=True
        ).select_related('employee').order_by('date_of_birth')

        birthdays = []
        for profile in employees_with_birthdays:
            if profile.date_of_birth:
                birth_date = profile.date_of_birth
                # Calculate age based on the viewed year
                age = current_year - birth_date.year
                # If we are viewing a past month in the current year, or future, just straightforward subtraction
                # However, traditionally age is "upcoming age" for that birthday.
                # So if birthday is in that year, the age they turn is year - birth_year.

                # Calculate days until birthday (relative to today, for sorting/urgency)
                # Ensure we construct the date for the viewed year
                try:
                    birthday_on_viewed_year = birth_date.replace(year=current_year)
                except ValueError:
                    # Handle Feb 29 on non-leap years
                    birthday_on_viewed_year = birth_date.replace(year=current_year, day=28)

                days_until = (birthday_on_viewed_year - today).days

                birthdays.append({
                    'id': profile.employee.id,
                    'name': profile.employee.name,
                    'username': profile.employee.username,
                    'department': profile.employee.department,
                    'date_of_birth': str(birth_date),
                    'age': age,
                    'days_until': days_until
                })

        # Sort by day of month
        birthdays.sort(key=lambda x: x['date_of_birth'].split('-')[2])  # Simple sort by day

        return Response({
            'success': True,
            'count': len(birthdays),
            'birthdays': birthdays
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to fetch upcoming birthdays'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def _send_task_notification(user, message, task_id, type="task"):
    """Helper to create a persistent notification in the database and trigger push"""
    from .models import Notification
    try:
        Notification.objects.create(
            user=user,
            type=type,
            message=message,
            link_id=str(task_id)
        )
        
        # Trigger Web Push Notification
        title = "Task Notification"
        if type == "task_comment":
            title = "Comment on Task"
        elif type == "meeting":
            title = "Meeting / MoM"
            
        _trigger_push_notification(user, title, message, f"task_{task_id}")
        
        return True
    except Exception as e:
        print(f"Error creating notification: {e}")
        return False

def _trigger_push_notification(user, title, message, link=None):
    """Internal helper to send browser push notifications via pywebpush"""
    from .models import PushSubscription
    from pywebpush import webpush, WebPushException
    import json
    from django.conf import settings

    subscriptions = PushSubscription.objects.filter(employee=user)
    if not subscriptions.exists():
        return

    vapid_private_key = getattr(settings, 'VAPID_PRIVATE_KEY', None)
    vapid_claims = {"sub": getattr(settings, 'VAPID_CLAIMS_SUB', 'mailto:admin@example.com')}

    # Fix VAPID private key if it has literal \n or extra quotes (common in .env)
    if isinstance(vapid_private_key, str):
        if "\\n" in vapid_private_key:
            vapid_private_key = vapid_private_key.replace("\\n", "\n")
        vapid_private_key = vapid_private_key.strip('"\'')

    if not vapid_private_key:
        print("[Push] VAPID private key is missing — skipping.")
        return

    data = json.dumps({
        "title": title,
        "body": message,
        "url": link  # Changed from "data": {"link": link} to match sw.js expectation
    })

    success_count = 0
    for sub in subscriptions:
        try:
            webpush(
                subscription_info={
                    "endpoint": sub.endpoint,
                    "keys": {
                        "p256dh": sub.p256dh,
                        "auth": sub.auth
                    }
                },
                data=data,
                vapid_private_key=vapid_private_key,
                vapid_claims=vapid_claims
            )
            success_count += 1
        except WebPushException as ex:
            print(f"[Push] WebPushException for {user.username}: {ex}")
            if ex.response and ex.response.status_code == 410:
                sub.delete()
        except Exception as e:
            print(f"[Push] Unexpected error for {user.username}: {e}")
    
    if success_count > 0:
        print(f"[Push] Successfully sent {success_count} notifications to {user.username}")


@api_view(['GET'])
@require_gated_token_api
def get_notifications(request):
    """Get notifications for the current user"""
    user_id = request.GET.get('user_id')
    if not user_id:
        return Response({'success': False, 'message': 'User ID required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        user = Employee.objects.get(id=user_id)
    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'User not found'}, status=status.HTTP_404_NOT_FOUND)

    notifications = []

    # 0. Received Birthday Wishes
    received_wishes = BirthdayWish.objects.filter(
        receiver_id=user_id,
        is_read=False
    ).select_related('sender').order_by('-created_at')

    for wish in received_wishes:
        notifications.append({
            'type': 'wish',
            'icon': '🎈',
            'message': f"{wish.sender.name}: {wish.message}",
            'time': wish.created_at.strftime('%I:%M %p'),
            'id': f'wish_{wish.id}'
        })

    # Note: Synthetic birthday and task notifications are removed here 
    # as they kept reappearing after "Mark all as read". 
    # Real notifications are handled by the Notification model.

    # 3. Pending requests (for admins/mentors)
    if user.role == 'admin':
        pending_requests_count = EmployeeRequest.objects.filter(
            status='pending'
        ).count()

        if pending_requests_count > 0:
            notifications.append({
                'type': 'request',
                'icon': '📋',
                'message': f'{pending_requests_count} pending approval(s)',
                'time': 'Now',
                'id': 'pending_requests'
            })
    
    # 3.1 Task Requests for Mentors
    is_mentor = user.role == 'mentor' or user.subordinates.exists()
    if is_mentor:
        task_requests = EmployeeRequest.objects.filter(
            request_type='task_request',
            status='pending',
            employee__mentors=user
        ).select_related('employee')
        
        for treq in task_requests:
            notifications.append({
                'type': 'task_request',
                'icon': '⚠️',
                'message': f'{treq.employee.name} is requesting a task',
                'time': treq.created_at.strftime('%I:%M %p') if treq.created_at else 'Now',
                'id': f'task_req_{treq.id}',
                'employee_id': treq.employee.id,
                'employee_name': treq.employee.name
            })
            
        # Idle Subordinates (No tasks) - Summary Notification
        idle_count = 0
        try:
            subordinates = user.subordinates.all()
            for sub in subordinates:
                # Check if this subordinate has any tasks in progress or todo
                has_active = Task.objects.filter(
                    assignees=sub, 
                    status__in=['todo', 'in_progress']
                ).exists()
                if not has_active:
                    idle_count += 1
        except Exception:
            idle_count = 0
        
    # Note: Synthetic/Calculated notifications are removed here to ensure "Mark all as read"
    # functions correctly. Users only see alerts that can be dismissed (tracked in DB).

    # 4. Persistence Notifications from DB
    from .models import Notification
    db_notifs = Notification.objects.filter(user=user, is_read=False).order_by('-created_at')[:15]
    for dn in db_notifs:
        # Avoid duplicating recent comments if they are already added above
        if dn.type == 'task_comment' and any(n.get('id') == f'comment_{dn.link_id}' for n in notifications):
            continue
            
        icon = '🔔'
        if dn.type == 'task_comment': icon = '💬'
        elif dn.type == 'task': icon = '📝'
        elif dn.type == 'meeting': icon = '🤝'
        elif dn.type == 'request': icon = '📋'
        
        notifications.append({
            'id': f'dn_{dn.id}',
            'type': dn.type,
            'icon': icon,
            'message': dn.message,
            'time': dn.created_at.strftime('%I:%M %p'),
            'task_id': dn.link_id
        })

    return Response({
        'success': True,
        'notifications': notifications,
        'unread_count': len(notifications)
    })

@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def request_new_task(request):
    """Employee requests a new task from their mentor"""
    user_id = request.data.get('user_id')
    if not user_id:
        return Response({'success': False, 'message': 'User ID required'}, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        user = Employee.objects.get(id=user_id)
        
        # Check if already has a pending task request
        existing = EmployeeRequest.objects.filter(
            employee=user,
            request_type='task_request',
            status='pending'
        ).exists()
        
        if existing:
            return Response({'success': False, 'message': 'You already have a pending task request sent to your mentor.'})
            
        req = EmployeeRequest.objects.create(
            employee=user,
            request_type='task_request',
            start_date=timezone.now().date(),
            end_date=timezone.now().date(),
            reason='New task request from Task Manager V2',
            status='pending'
        )

        # Notify mentors or admins
        mentors = user.mentors.all()
        if not mentors.exists():
            # Fallback to admins
            mentors = Employee.objects.filter(role='admin', is_active=True)
        
        notif_msg = f"{user.name} has requested a new task assignment."
        for mentor in mentors:
            _send_task_notification(
                user=mentor,
                message=notif_msg,
                task_id=f"req_{req.id}",
                type="request"
            )
        
        return Response({'success': True, 'message': 'Task request sent successfully to your mentor/admin!'})
        
    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def mark_notifications_read(request):
    """Mark all notifications or a specific one as read"""
    user_id = request.data.get('user_id')
    notification_id = request.data.get('notification_id') # Optional: if we want to mark specific

    if not user_id:
        return Response({'success': False, 'message': 'User ID required'}, status=status.HTTP_400_BAD_REQUEST)

    # Handle BirthdayWishes persistence
    wishes = BirthdayWish.objects.filter(receiver_id=user_id, is_read=False)
    if notification_id and str(notification_id).startswith('wish_'):
        wish_id = str(notification_id).replace('wish_', '')
        wishes = wishes.filter(id=wish_id)
    wishes.update(is_read=True)

    # Handle persistent Notification model
    from .models import Notification
    db_notifs = Notification.objects.filter(user_id=user_id, is_read=False)
    if notification_id and str(notification_id).startswith('dn_'):
        notif_id = str(notification_id).replace('dn_', '')
        db_notifs = db_notifs.filter(id=notif_id)
    db_notifs.update(is_read=True)

    return Response({'success': True, 'message': 'Notifications marked as read'})


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def send_birthday_wish(request):
    """Send a birthday wish to an employee"""
    sender_id = request.data.get('sender_id')
    receiver_id = request.data.get('receiver_id')
    message = request.data.get('message', 'Wishing you a very Happy Birthday! 🎂')

    if not all([sender_id, receiver_id]):
        return Response({'success': False, 'message': 'Sender and Receiver IDs required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        sender = Employee.objects.get(id=sender_id)
        receiver = Employee.objects.get(id=receiver_id)

        # Prevent duplicate wishes for same day
        today_start = timezone.now().replace(hour=0, minute=0, second=0, microsecond=0)
        existing_wish = BirthdayWish.objects.filter(
            sender=sender,
            receiver=receiver,
            created_at__gte=today_start
        ).exists()

        if existing_wish:
            return Response({'success': False, 'message': 'You have already sent a wish today!'})

        wish = BirthdayWish.objects.create(
            sender=sender,
            receiver=receiver,
            message=message
        )
        return Response({'success': True, 'message': 'Birthday wish sent successfully!'})

    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'User not found'}, status=status.HTTP_404_NOT_FOUND)



@api_view(['GET'])
@require_gated_token_api
def pending_requests(request):
    """Get pending or history (approved/rejected) WFH and leave requests"""
    user_id = request.GET.get('user_id')
    user = Employee.objects.filter(id=user_id).first() if user_id else None
    # Include de-facto Mentors (any user who has subordinates)
    is_mentor = user and (user.role == 'mentor' or (user.role != 'admin' and user.subordinates.exists()))

    try:
        status_param = request.GET.get('status', 'pending')
        
        # Buffer logic for WFH requests
        today = timezone.now().date()
        buffer_date = today - timedelta(days=2)
        
        if status_param == 'history':
            # Get approved and rejected requests, plus expired WFH requests
            requests_obj = EmployeeRequest.objects.filter(
                Q(status__in=['approved', 'rejected']) |
                Q(request_type='wfh', end_date__lt=buffer_date)
            ).select_related('employee').order_by('-start_date')
            
            # If we are specifically in history, we should exclude Approved WFH that are still within buffer
            # (since they are shown in Active)
            requests_obj = requests_obj.exclude(
                status='approved', request_type='wfh', end_date__gte=buffer_date
            )
        else:
            # Get pending requests, plus recently approved WFH requests
            requests_obj = EmployeeRequest.objects.filter(
                Q(status='pending') |
                Q(status='approved', request_type='wfh', end_date__gte=buffer_date)
            ).select_related('employee').order_by('start_date')
            
            # Exclude WFH requests that have expired (even if pending)
            requests_obj = requests_obj.exclude(
                request_type='wfh', end_date__lt=buffer_date
            )

        if is_mentor:
            requests_obj = requests_obj.filter(employee__mentors=user)

        requests_data = []
        from .models import Task
        for req in requests_obj:
            task_info = None
            if req.request_type == 'wfh':
                today_tasks = Task.objects.filter(assignees=req.employee, due_date=req.start_date)
                total = today_tasks.count()
                completed = today_tasks.filter(status='completed').count()
                task_info = {
                    'total': total,
                    'completed': completed,
                    'percent': round((completed / total * 100), 1) if total > 0 else 0,
                    'summary': f"{completed}/{total} tasks completed" if total > 0 else "No tasks due"
                }

            requests_data.append({
                'id': req.id,
                'employee_id': req.employee.id,
                'employee_name': req.employee.name,
                'username': req.employee.username,
                'type': req.request_type,
                'date': str(req.start_date), # Frontend uses this key currently
                'start_date': str(req.start_date),
                'end_date': str(req.end_date),
                'reason': req.reason,
                'status': req.status,
                'task_info': task_info,
                'reviewed_by_name': req.reviewed_by.name if req.reviewed_by else None,
                'is_mentor': req.reviewed_by in req.employee.mentors.all() if req.reviewed_by else False,
                'created_at': req.created_at.isoformat()
            })

        return Response({
            'success': True,
            'count': len(requests_data),
            'requests': requests_data
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to fetch pending requests'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



@api_view(['GET'])
@require_gated_token_api
def my_requests(request):
    """Get request history for an employee"""
    try:
        employee_id = request.GET.get('employee_id')
        user = get_current_user(request, requested_id=employee_id)
        if not user:
            return Response({'success': False, 'message': 'Unauthorized'}, status=403)

        # Get all requests for employee
        requests_obj = EmployeeRequest.objects.filter(
            employee=user
        ).order_by('-created_at')

        requests_data = []
        for req in requests_obj:
            requests_data.append({
                'id': req.id,
                'type': req.request_type,
                'start_date': str(req.start_date),
                'end_date': str(req.end_date),
                'reason': req.reason,
                'status': req.status,
                'admin_response': req.admin_response,
                'reviewed_by_name': req.reviewed_by.name if req.reviewed_by else None,
                'is_mentor': req.reviewed_by in req.employee.mentors.all() if req.reviewed_by else False,
                'created_at': req.created_at.isoformat()
            })

        return Response({
            'success': True,
            'count': len(requests_data),
            'requests': requests_data
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to fetch request history'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def active_tasks(request):
    """Get count of in-progress tasks"""
    try:
        # Use request.user which is already validated by @require_gated_token_api
        caller = getattr(request, 'user', None)
        
        # In development, request.user might be AnonymousUser if token is missing
        if not caller or not caller.is_authenticated or not hasattr(caller, 'role'):
            return Response({'success': False, 'message': 'Unauthorized or Session Expired'}, status=403)
            
        requested_id = request.GET.get('employee_id')
        
        # Determine which user's tasks we are looking at
        if caller.role.lower() == 'admin':
            # Admin can see any user's count, or all tasks if requested_id is missing or points to them
            if requested_id and str(requested_id) != str(caller.id):
                user = Employee.objects.filter(id=requested_id).first() or caller
            else:
                user = caller
        else:
            # Non-admins can ONLY see their own count
            user = caller

        query = Task.objects.filter(status='in_progress')
        
        # Safety check for role
        role = (user.role or '').lower()
        
        # Treat de-facto Mentors (employees with subordinates) like official Mentors
        is_emp_mentor = role == 'mentor' or (role != 'admin' and user.subordinates.exists())
        
        if role == 'admin':
            # Admin sees all in-progress tasks
            pass
        elif is_emp_mentor:
            query = query.filter(Q(assignees__mentors=user) | Q(mentor=user) | Q(created_by=user) | Q(assignees=user)).distinct()
        else:
            query = query.filter(Q(assignees=user) | Q(mentor=user)).distinct()

        active_count = query.count()

        return Response({
            'success': True,
            'count': active_count
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to fetch active tasks count'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def _get_admin_task_mentor_data():
    """Helper: Get all tasks for Admin Task Mentor"""
    tasks = Task.objects.select_related('created_by', 'mentor').prefetch_related('assignees').order_by('-created_at')
    return _serialize_tasks(tasks)

def _get_employee_my_tasks_data(employee):
    """Helper: Get assigned tasks + overseen tasks for Employee My Tasks"""
    tasks = Task.objects.filter(
        Q(assignees=employee) | Q(mentor=employee) | Q(overseers=employee)
    ).distinct().select_related('created_by', 'mentor').prefetch_related('assignees').order_by('-created_at')
    return _serialize_tasks(tasks)

def _get_mentor_employees_tasks_data(mentor):
    """Helper: Get tasks for employees reporting to this mentor + tasks explicitly managed by them"""
    # Exclude tasks where the mentor themselves is an assignee to keep Team Tasks focused on management
    query = (Q(assignees__mentors=mentor) | Q(mentor=mentor) | Q(overseers=mentor))
    tasks = Task.objects.filter(query).exclude(assignees=mentor).distinct().select_related('created_by', 'mentor').prefetch_related('assignees').order_by('-created_at')
    return _serialize_tasks(tasks)

def _serialize_tasks(tasks):
    """Helper: Serialize task list with comments"""
    data = []
    for task in tasks:
        # Get comments for each task
        comments = []
        for comment in task.comments.all().select_related('author'):
            comments.append({
                'id': comment.id,
                'author_name': comment.author.name,
                'content': comment.content,
                'created_at': comment.created_at.isoformat()
            })

        # Get all assignees
        assignees_info = []
        for assignee in task.assignees.all():
            assignees_info.append({
                'id': assignee.id,
                'name': assignee.name
            })

        # Get steps
        steps = []
        for step in task.steps.all():
            steps.append({
                'id': step.id,
                'text': step.text,
                'is_completed': step.is_completed
            })

        # Get history
        history_log = []
        for h in task.history.all().select_related('changed_by'):
            history_log.append({
                'id': h.id,
                'field': h.field_changed,
                'old': h.old_value,
                'new': h.new_value,
                'by': h.changed_by.name if h.changed_by else 'System',
                'at': h.changed_at.isoformat()
            })

        # Get attachments
        attachments = []
        for a in task.attachments.all():
            attachments.append({
                'id': a.id,
                'name': a.file.name.split('/')[-1],
                'url': a.file.url
            })

        data.append({
            'id': task.id,
            'title': task.title,
            'description': task.description,
            'status': task.status,
            'priority': task.priority,
            'assignees': assignees_info,
            'overseers': [{'id': o.id, 'name': o.name} for o in task.overseers.all()],
            'Mentor_id': task.mentor.id if task.mentor else None,
            'Mentor_name': task.mentor.name if task.mentor else None,
            'mentor_id': task.mentor.id if task.mentor else None,
            'mentor_name': task.mentor.name if task.mentor else None,
            'created_by': task.created_by.id,
            'created_by_name': task.created_by.name,
            'start_date': str(task.start_date) if task.start_date else None,
            'due_date': str(task.due_date) if task.due_date else None,
            'started_at': task.started_at.isoformat() if task.started_at else None,
            'completed_at': task.completed_at.isoformat() if task.completed_at else None,
            'created_at': task.created_at.isoformat(),
            'updated_at': task.updated_at.isoformat(),
            'comments': comments,
            'steps': steps,
            'history': history_log,
            'attachments': attachments
        })
    return data

def _create_task_admin(data, creator, files=None):
    """Helper: Admin creates a task"""
    required_fields = ['title'] # assignees handled below
    for field in required_fields:
        if not data.get(field):
            raise ValueError(f'{field} is required')

    import json
    assigned_input = data.get('assignees') or data.get('assigned_to')
    if isinstance(assigned_input, str):
        try:
            assigned_input = json.loads(assigned_input)
        except json.JSONDecodeError:
            assigned_input = [assigned_input]
    assigned_ids = assigned_input if isinstance(assigned_input, list) else [assigned_input] if assigned_input else []
    
    mentor_id = data.get('mentor_id')
    overseer_ids = data.get('overseer_ids') or []
    if isinstance(overseer_ids, str):
        try:
            overseer_ids = json.loads(overseer_ids)
        except json.JSONDecodeError:
            overseer_ids = [overseer_ids]
            
    if not overseer_ids and mentor_id and mentor_id != 'none':
        overseer_ids = [mentor_id]

    Mentor_employee = None
    if mentor_id and mentor_id != 'none':
        try:
            Mentor_employee = Employee.objects.get(id=mentor_id)
        except:
            pass
    elif overseer_ids:
        try:
            Mentor_employee = Employee.objects.get(id=overseer_ids[0])
        except:
            pass

    start_date = data.get('start_date')
    if not start_date:
        start_date = None

    due_date = data.get('due_date')
    if not due_date:
        due_date = None

    task = Task.objects.create(
        title=data['title'],
        description=data.get('description', ''),
        status=data.get('status', 'todo'),
        priority=data.get('priority', 'medium'),
        mentor=Mentor_employee,
        created_by=creator,
        start_date=start_date,
        due_date=due_date
    )
    
    task.assignees.set(Employee.objects.filter(id__in=assigned_ids))
    if overseer_ids:
        task.overseers.set(Employee.objects.filter(id__in=overseer_ids))

    EmployeeRequest.objects.filter(
        employee_id__in=assigned_ids,
        request_type='task_request',
        status='pending'
    ).update(status='approved', admin_response=f'Task "{task.title}" assigned.')

    if files:
        attachments = files.getlist('attachments')
        for f in attachments:
            TaskAttachment.objects.create(task=task, file=f)

    # Notification Logic (Consolidated)
    task_title = str(task.title or "Untitled Task")
    is_mom = any(kw in task_title.upper() for kw in ["MOM", "MEETING"]) or task_title.startswith("MoM Tasks")

    assignee_objs = Employee.objects.filter(id__in=assigned_ids)
    for assignee in assignee_objs:
        msg = f"New Minutes/Task: {task_title}" if is_mom else f"New task assigned: {task_title}"
        notif_type = "meeting" if is_mom else "task"
        try:
            Notification.objects.create(
                user_id=assignee.id,
                type=notif_type,
                message=msg,
                link_id=str(task.id)
            )
            # Sync Global Utility (Defined in this file)
            _trigger_push_notification(assignee, "Meeting MoM" if is_mom else "Task Assignment", msg, f"task_{task.id}")
        except Exception as e:
            print(f"Failed notif for {assignee.id}: {e}")

    return task

@api_view(['GET', 'POST'])
@require_gated_token_api
@parser_classes([JSONParser, MultiPartParser, FormParser])
def tasks_api(request):
    """Get all tasks or create a new task (Separated Admin/Employee Logic)"""
    if request.method == 'GET':
        try:
            employee_id = request.GET.get('employee_id')

            if not employee_id:
                # Security default
                return Response({'success': True, 'tasks': []})

            try:
                emp = Employee.objects.get(id=employee_id)

                if emp.role == 'admin':
                    # ADMIN PATH - Now respects scope for consistency with Mentor view
                    scope = request.GET.get('scope')
                    if scope == 'my':
                        # Tasks assigned TO the admin
                        tasks = Task.objects.filter(assignees=emp).distinct().select_related('created_by', 'mentor').prefetch_related('assignees').order_by('-created_at')
                        tasks_data = _serialize_tasks(tasks)
                    else:
                        # Full view for admin (Team/All)
                        tasks_data = _get_admin_task_mentor_data()
                elif emp.role == 'mentor':
                    # Mentor PATH - Separated based on scope
                    scope = request.GET.get('scope')
                    if scope == 'my':
                        # Strictly tasks assigned TO the mentor
                        tasks = Task.objects.filter(assignees=emp).distinct().select_related('created_by', 'mentor').prefetch_related('assignees').order_by('-created_at')
                        tasks_data = _serialize_tasks(tasks)
                    elif scope == 'team':
                        # Subordinates' tasks (Excludes mentor's personal tasks)
                        tasks_data = _get_mentor_employees_tasks_data(emp)
                    else:
                        # Legacy merged view (if no scope provided)
                        own_tasks = _get_employee_my_tasks_data(emp)
                        subordinate_tasks = _get_mentor_employees_tasks_data(emp)
                        tasks_data = own_tasks + [t for t in subordinate_tasks if t['id'] not in [ot['id'] for ot in own_tasks]]
                else:
                    # EMPLOYEE PATH
                    scope = request.GET.get('scope')
                    if scope == 'my':
                        # Strictly tasks assigned TO the employee
                        tasks = Task.objects.filter(assignees=emp).distinct().select_related('created_by', 'mentor').prefetch_related('assignees').order_by('-created_at')
                        tasks_data = _serialize_tasks(tasks)
                    else:
                        tasks_data = _get_employee_my_tasks_data(emp)

                return Response({
                    'success': True,
                    'tasks': tasks_data
                })

            except Employee.DoesNotExist:
                return Response({'success': True, 'tasks': []})

        except Exception as e:
            return Response({
                'success': False,
                'message': f'Failed to fetch tasks: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    elif request.method == 'POST':
        try:
            data = request.data
            creator_id = data.get('created_by')

            # Identify creator
            if creator_id:
                creator = Employee.objects.get(id=creator_id)
            else:
                creator = Employee.objects.filter(role='admin').first()
                if not creator:
                    return Response({'success': False, 'message': 'No creator found'}, status=status.HTTP_400_BAD_REQUEST)

            # Dispatch creation logic
            if creator.role == 'admin':
                task = _create_task_admin(data, creator, request.FILES)
            else:
                # Re-use admin logic for now as employee creation wasn't strictly defined different yet, 
                # but valid separation point.
                task = _create_task_admin(data, creator, request.FILES) 

            return Response({
                'success': True,
                'message': 'Task created successfully',
                'task_id': task.id
            })

        except ValueError as e:
            return Response({'success': False, 'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        except Employee.DoesNotExist:
            return Response({'success': False, 'message': 'Assigned employee or Mentor not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            # More helpful error for debugging
            return Response({
                'success': False,
                'message': f'Failed to create task: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# Alias: URL conf references 'views.create_task' for the explicit create endpoint
create_task = tasks_api

def _update_task_admin(task, data, user=None):
    """Helper: Admin/Overseer/Reporting Mentor updates task details"""
    user_role = str(user.role).lower() if user else 'none'
    is_admin = user_role == 'admin'
    is_overseer = task.mentor and user and task.mentor.id == user.id
    
    # Mentor check: user manages at least one of the assignees
    is_reporting_mentor = False
    if user:
        is_reporting_mentor = task.assignees.filter(mentors=user).exists()

    if task.status == 'completed' and not (is_admin or is_overseer or is_reporting_mentor):
        # We allow Admins and the Overseer to bypass this for correction/reopening
        raise ValueError(f"Cannot modify a completed task.")

    if 'status' in data:
        new_status = data['status']
        if new_status == 'in_progress' and not task.started_at:
            task.started_at = timezone.now()
        elif new_status == 'completed' and not task.completed_at:
            task.completed_at = timezone.now()
        task.status = new_status
    if 'priority' in data:
        old_priority = str(task.priority).lower() if task.priority else 'medium'
        new_priority = str(data['priority']).lower()
        if old_priority != new_priority:
            from .models import TaskHistory
            
            last_24h = timezone.now() - timedelta(hours=24)
            recent_change = TaskHistory.objects.filter(
                task=task, field_changed='priority', changed_at__gte=last_24h
            ).order_by('-changed_at').first()
            
            if recent_change:
                recent_change.new_value = new_priority
                recent_change.changed_by = user
                recent_change.changed_at = timezone.now()
                recent_change.save()
            else:
                TaskHistory.objects.create(
                    task=task,
                    field_changed='priority',
                    old_value=old_priority,
                    new_value=new_priority,
                    changed_by=user
                )
        task.priority = new_priority
    if 'title' in data:
        old_title = str(task.title) if task.title else ''
        new_title = str(data['title'])
        if old_title != new_title:
            from .models import TaskHistory
            TaskHistory.objects.create(
                task=task,
                field_changed='title',
                old_value=old_title,
                new_value=new_title,
                changed_by=user
            )
        task.title = new_title
    if 'description' in data:
        old_desc = str(task.description) if task.description else ''
        new_desc = str(data['description'])
        if old_desc != new_desc:
            from .models import TaskHistory
            TaskHistory.objects.create(
                task=task,
                field_changed='description',
                old_value=old_desc,
                new_value=new_desc,
                changed_by=user
            )
        task.description = new_desc
    if 'start_date' in data:
        old_start = str(task.start_date) if task.start_date else 'None'
        new_start = str(data['start_date'])
        if old_start != new_start:
            from .models import TaskHistory
            TaskHistory.objects.create(
                task=task,
                field_changed='start_date',
                old_value=old_start,
                new_value=new_start,
                changed_by=user
            )
        task.start_date = data['start_date']
    if 'due_date' in data:
        old_due = str(task.due_date) if task.due_date else 'None'
        new_due = str(data['due_date'])
        if old_due != new_due:
            from .models import TaskHistory
            TaskHistory.objects.create(
                task=task,
                field_changed='due_date',
                old_value=old_due,
                new_value=new_due,
                changed_by=user
            )
        task.due_date = data['due_date']
        
    if 'steps' in data:
        from .models import TaskStep
        # Expecting data['steps'] to be a list of {id?: int, text: string, is_completed: bool}
        # Simplified: overwrite OR update. Let's do update logic.
        incoming_steps = data['steps']
        current_step_ids = []
        for s_data in incoming_steps:
            if s_data.get('id'):
                step = TaskStep.objects.get(id=s_data['id'], task=task)
                step.text = s_data['text']
                step.is_completed = s_data['is_completed']
                step.save()
                current_step_ids.append(step.id)
            else:
                new_step = TaskStep.objects.create(
                    task=task,
                    text=s_data['text'],
                    is_completed=s_data.get('is_completed', False)
                )
                current_step_ids.append(new_step.id)
        # Remove steps not in incoming data? For now keep them simple unless requested.
        
    import json
    if 'assignees' in data or 'assigned_to' in data:
        assigned_input = data.get('assignees') or data.get('assigned_to')
        if isinstance(assigned_input, str):
            try:
                assigned_input = json.loads(assigned_input)
            except json.JSONDecodeError:
                assigned_input = [assigned_input]
        assigned_ids = assigned_input if isinstance(assigned_input, list) else [assigned_input] if assigned_input else []
        task.assignees.set(Employee.objects.filter(id__in=assigned_ids))

    if 'mentor_id' in data or 'overseer_ids' in data:
        overseer_ids = data.get('overseer_ids')
        if isinstance(overseer_ids, str):
            try:
                overseer_ids = json.loads(overseer_ids)
            except json.JSONDecodeError:
                overseer_ids = [overseer_ids]
        mentor_id = data.get('mentor_id')
        
        if overseer_ids is not None:
            # Multi-overseer update
            overseer_qs = Employee.objects.filter(id__in=overseer_ids)
            task.overseers.set(overseer_qs)
            # Sync backward compatibility field 'Mentor'
            if overseer_qs.exists():
                task.mentor = overseer_qs.first()
            else:
                task.mentor = None
        elif mentor_id:
            if mentor_id == 'none':
                task.mentor = None
                task.overseers.clear()
            else:
                try:
                    Mentor_emp = Employee.objects.get(id=mentor_id)
                    task.mentor = Mentor_emp
                    task.overseers.set([Mentor_emp])
                except:
                    pass
    task.save()
    return True

def _update_task_employee(task, data, user=None):
    """Helper: Employee updates task (limited access - mostly status)"""
    user_role = str(user.role).lower() if user else 'none'
    # Employee typically only updates status or adds comments (comments not implemented yet)
    if task.status == 'completed' and user_role != 'admin':
        # STRICTLY BLOCK for generic updates
        # Exception: If user is trying to reopen? "it can't be changed" implies NO.
        # Exception: If user is trying to reopen? "it can't be changed" implies NO.
        # return False - REMOVED to allow raising exception
        raise ValueError(f"Cannot modify a completed task (ReqID: {user.id if user else '?'})")

    if 'title' in data:
        old_title = str(task.title) if task.title else ''
        new_title = str(data['title'])
        if old_title != new_title:
            from .models import TaskHistory
            TaskHistory.objects.create(
                task=task,
                field_changed='title',
                old_value=old_title,
                new_value=new_title,
                changed_by=user
            )
        task.title = new_title

    if 'description' in data:
        old_desc = str(task.description) if task.description else ''
        new_desc = str(data['description'])
        if old_desc != new_desc:
            from .models import TaskHistory
            TaskHistory.objects.create(
                task=task,
                field_changed='description',
                old_value=old_desc,
                new_value=new_desc,
                changed_by=user
            )
        task.description = new_desc

    if 'due_date' in data:
        old_due = str(task.due_date) if task.due_date else 'None'
        new_due = str(data['due_date'])
        if old_due != new_due:
            from .models import TaskHistory
            TaskHistory.objects.create(
                task=task,
                field_changed='due_date',
                old_value=old_due,
                new_value=new_due,
                changed_by=user
            )
        task.due_date = data['due_date']
        
    if 'priority' in data:
        old_priority = str(task.priority).lower() if task.priority else 'medium'
        new_priority = str(data['priority']).lower()
        if old_priority != new_priority:
            from .models import TaskHistory
            
            last_24h = timezone.now() - timedelta(hours=24)
            recent_change = TaskHistory.objects.filter(
                task=task, field_changed='priority', changed_at__gte=last_24h
            ).order_by('-changed_at').first()
            
            if recent_change:
                recent_change.new_value = new_priority
                recent_change.changed_by = user
                recent_change.changed_at = timezone.now()
                recent_change.save()
            else:
                TaskHistory.objects.create(
                    task=task,
                    field_changed='priority',
                    old_value=old_priority,
                    new_value=new_priority,
                    changed_by=user
                )
        task.priority = new_priority

    if 'status' in data:
        new_status = data['status']
        if new_status == 'in_progress' and not task.started_at:
            task.started_at = timezone.now()
        elif new_status == 'completed' and not task.completed_at:
            task.completed_at = timezone.now()
        task.status = new_status

    if 'steps' in data:
        from .models import TaskStep
        incoming_steps = data['steps']
        for s_data in incoming_steps:
            if s_data.get('id'):
                try:
                    step = TaskStep.objects.get(id=s_data['id'], task=task)
                    step.text = s_data['text']
                    step.is_completed = s_data['is_completed']
                    step.save()
                except TaskStep.DoesNotExist:
                    pass
            else:
                TaskStep.objects.create(
                    task=task,
                    text=s_data['text'],
                    is_completed=s_data.get('is_completed', False)
                )

    task.save()
    return True

@api_view(['GET', 'POST'])
@require_gated_token_api
@parser_classes([JSONParser, MultiPartParser, FormParser])
def meetings_api(request):
    """List or create meetings"""
    if request.method == 'GET':
        try:
            employee_id = request.GET.get('employee_id')
            if not employee_id:
                return Response({'success': False, 'message': 'Employee ID required'}, status=status.HTTP_400_BAD_REQUEST)
            
            from .models import Meeting
            # Get meetings where user is participant OR creator
            meetings = Meeting.objects.filter(
                Q(participants__id=employee_id) | Q(created_by_id=employee_id)
            ).distinct().prefetch_related('participants', 'created_by').order_by('-date', '-start_time')[:20]
            
            data = []
            for m in meetings:
                import re as _re, json as _json
                raw_desc = m.description or ''
                # Extract embedded steps_json
                steps_data = []
                steps_match = _re.search(r'__steps_json__(.*?)__end_steps__', raw_desc, _re.DOTALL)
                if steps_match:
                    try:
                        steps_data = _json.loads(steps_match.group(1))
                    except Exception:
                        steps_data = []
                # Clean display description
                display_desc = _re.sub(r'\n\n__steps_json__.*?__end_steps__', '', raw_desc, flags=_re.DOTALL).strip()

                data.append({
                    'id': m.id,
                    'title': m.title,
                    'description': display_desc,
                    'steps': steps_data,
                    'date': m.date.strftime('%Y-%m-%d'), 
                    'display_date': m.date.strftime('%d-%m-%Y'),
                    'start_time': m.start_time.strftime('%H:%M') if m.start_time else '', 
                    'display_time': m.start_time.strftime('%I:%M %p') if m.start_time else '',
                    'created_by_name': m.created_by.name,
                    'created_by_id': m.created_by.id,
                    'participants': [{'id': p.id, 'name': p.name} for p in m.participants.all()]
                })
            
            return Response({'success': True, 'meetings': data})
        except Exception as e:
            return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    elif request.method == 'POST':
        try:
            from .models import Meeting, Employee
            data = request.data
            creator_id = data.get('created_by')
            if not creator_id:
                return Response({'success': False, 'message': 'Creator ID required'}, status=status.HTTP_400_BAD_REQUEST)
            
            creator = Employee.objects.get(id=creator_id)

            # Store steps_json in description for later editing
            description = data.get('description', '')
            steps_json = data.get('steps_json')
            if steps_json:
                description = f"{description}\n\n__steps_json__{steps_json}__end_steps__"

            meeting = Meeting.objects.create(
                title=data.get('title'),
                description=description,
                date=data.get('date'),
                start_time=data.get('start_time') if data.get('start_time') else None,
                created_by=creator
            )
            
            participant_ids = data.get('participants', [])
            if participant_ids:
                meeting.participants.set(Employee.objects.filter(id__in=participant_ids))
            
            return Response({'success': True, 'meeting_id': meeting.id})
        except Exception as e:
            return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET', 'PATCH', 'DELETE'])
@require_gated_token_api
@parser_classes([JSONParser])
def meeting_detail_api(request, meeting_id):
    """Update or delete a meeting"""
    try:
        from .models import Meeting, Employee
        meeting = Meeting.objects.get(id=meeting_id)
        
        if request.method == 'DELETE':
            meeting.delete()
            return Response({'success': True, 'message': 'Meeting deleted'})
            
        elif request.method == 'PATCH':
            data = request.data
            if 'title' in data: meeting.title = data['title']
            if 'date' in data: meeting.date = data['date']
            if 'start_time' in data: 
                meeting.start_time = data['start_time'] if data['start_time'] else None
            
            # Update description, embedding steps_json if provided
            new_desc = data.get('description', meeting.description or '')
            steps_json = data.get('steps_json')
            if steps_json:
                import re
                new_desc = re.sub(r'\n\n__steps_json__.*?__end_steps__', '', new_desc or '', flags=re.DOTALL)
                new_desc = f"{new_desc}\n\n__steps_json__{steps_json}__end_steps__"
            meeting.description = new_desc

            if 'participants' in data:
                meeting.participants.set(Employee.objects.filter(id__in=data['participants']))
            
            meeting.save()
            return Response({'success': True, 'message': 'Meeting updated'})
            
    except Meeting.DoesNotExist:
        return Response({'success': False, 'message': 'Meeting not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET', 'POST', 'PATCH', 'DELETE'])
@require_gated_token_api
@parser_classes([JSONParser, MultiPartParser, FormParser])
def task_detail_api(request, task_id):
    """Update, delete or fetch a task (Separated Admin/Employee Logic)"""
    try:
        task = Task.objects.prefetch_related('assignees').select_related('mentor').get(id=task_id)
    except Task.DoesNotExist:
        return Response({
            'success': False,
            'message': 'Task not found'
        }, status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        # Serialize task using existing helper
        task_data = _serialize_tasks([task])[0]
        return Response({
            'success': True,
            'task': task_data
        })

    data = request.data
    requesting_user_id = data.get('user_id') # Must be passed from frontend

    if not requesting_user_id:
        return Response({'success': False, 'message': 'User verification required'}, status=status.HTTP_403_FORBIDDEN)

    try:
        requesting_user = Employee.objects.get(id=requesting_user_id)

        # Check permissions and dispatch
        if request.method in ['POST', 'PATCH']:
            # Check for DELETE method simulation
            if data.get('_method') == 'DELETE':
                is_task_mentor = task.mentor and task.mentor.id == requesting_user.id
                if requesting_user.role != 'admin' and not is_task_mentor:
                    return Response({'success': False, 'message': 'Unauthorized to delete this task'}, status=status.HTTP_403_FORBIDDEN)

                task.delete()
                return Response({'success': True, 'message': 'Task deleted'})

            # Update Logic
            role = str(requesting_user.role).lower()
            is_assignee = task.assignees.filter(id=requesting_user.id).exists()
            is_mentor_of_assignee = task.assignees.filter(mentors=requesting_user).exists()

            if role == 'admin':
                _update_task_admin(task, data, requesting_user)
                return Response({'success': True, 'message': 'Task updated (Admin)'})

            elif task.mentor and task.mentor.id == requesting_user.id:
                # Task Overseer can also perform full updates
                _update_task_admin(task, data, requesting_user)
                return Response({'success': True, 'message': 'Task updated (Overseer)'})

            elif is_mentor_of_assignee:
                # Assignee's Reporting Mentor can also perform full updates
                _update_task_admin(task, data, requesting_user)
                return Response({'success': True, 'message': 'Task updated (Mentor)'})

            elif is_assignee:
                _update_task_employee(task, data, requesting_user)
                return Response({'success': True, 'message': 'Task updated (Employee)'})

            else:
                return Response({'success': False, 'message': 'Unauthorized'}, status=status.HTTP_403_FORBIDDEN)

    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'User not found'}, status=status.HTTP_403_FORBIDDEN)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def bulk_update_tasks(request):
    """Update multiple tasks at once (primarily for priority ranking or step toggling)"""
    data = request.data
    updates = data.get('updates', [])
    user_id = data.get('user_id')
    task_ids = data.get('task_ids', [])

    if not user_id:
        return Response({'success': False, 'message': 'User ID required'}, status=status.HTTP_403_FORBIDDEN)

    try:
        user = Employee.objects.get(id=user_id)
        
        # If it's a dictionary, it's a special V2 single-task update (steps, etc.)
        if isinstance(updates, dict):
            if not task_ids:
                return Response({'success': False, 'message': 'task_ids required for this update type'}, status=status.HTTP_400_BAD_REQUEST)
            
            tasks = Task.objects.filter(id__in=task_ids)
            if not tasks.exists():
                return Response({'success': False, 'message': 'Tasks not found'}, status=status.HTTP_404_NOT_FOUND)

            for task in tasks:
                try:
                    # Permission check: Admin, Mentor, or Assignee
                    is_assignee = task.assignees.filter(id=user.id).exists()
                    is_admin_mentor = str(user.role).lower() in ['admin', 'mentor'] or user.subordinates.exists()
                    
                    if not (is_assignee or is_admin_mentor):
                        continue  # Skip tasks user has no permission for

                    # Handle Step Toggles
                    if 'steps_toggle' in updates:
                        for step_data in updates['steps_toggle']:
                            step = TaskStep.objects.filter(id=step_data['id'], task=task).first()
                            if step:
                                old_val = "Completed" if step.is_completed else "Pending"
                                new_val = "Completed" if step_data['is_completed'] else "Pending"
                                if step.is_completed != step_data['is_completed']:
                                    step.is_completed = step_data['is_completed']
                                    step.save()
                                    TaskHistory.objects.create(
                                        task=task,
                                        field_changed=f"Step: {step.text[:30]}",
                                        old_value=old_val,
                                        new_value=new_val,
                                        changed_by=user
                                    )

                    # Handle Add Step
                    if 'add_step' in updates:
                        new_step = TaskStep.objects.create(task=task, text=updates['add_step'])
                        TaskHistory.objects.create(
                            task=task,
                            field_changed="Added Step",
                            old_value="",
                            new_value=new_step.text,
                            changed_by=user
                        )

                    # Handle Status Update
                    if 'status' in updates:
                        new_status = updates['status']
                        old_status = task.status
                        if old_status != new_status:
                            task.status = new_status
                            # Set timestamps on transition
                            if new_status == 'in_progress' and not task.started_at:
                                task.started_at = timezone.now()
                            elif new_status == 'completed' and not task.completed_at:
                                task.completed_at = timezone.now()
                            task.save()
                            TaskHistory.objects.create(
                                task=task,
                                field_changed="status",
                                old_value=old_status,
                                new_value=new_status,
                                changed_by=user
                            )

                    # Auto Status Transitions from step changes (only if status not manually set)
                    elif 'steps_toggle' in updates or 'add_step' in updates:
                        all_steps = list(task.steps.all())
                        if all_steps:
                            total_steps = len(all_steps)
                            completed_steps = sum(1 for s in all_steps if s.is_completed)
                            old_status = task.status
                            if completed_steps == total_steps:
                                task.status = 'completed'
                                if not task.completed_at:
                                    task.completed_at = timezone.now()
                            elif completed_steps > 0:
                                task.status = 'in_progress'
                                if not task.started_at:
                                    task.started_at = timezone.now()
                            if task.status != old_status:
                                task.save()
                                TaskHistory.objects.create(
                                    task=task,
                                    field_changed="status",
                                    old_value=old_status,
                                    new_value=task.status,
                                    changed_by=user
                                )

                except Exception as task_err:
                    return Response({'success': False, 'message': f'Error updating task {task.id}: {str(task_err)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


            return Response({'success': True, 'message': 'Task(s) updated successfully'})

        # Legacy list-based updates (Priority ranking)
        # Only allow Admin/Mentor for this
        if user.role != 'admin' and user.role != 'Mentor':
            return Response({'success': False, 'message': 'Admin or Mentor access required for priority updates'}, status=status.HTTP_403_FORBIDDEN)

        for item in updates:
            task_id = item.get('id')
            priority = item.get('priority')
            if task_id and priority:
                task = Task.objects.filter(id=task_id).first()
                if task and task.priority != priority:
                    old_p = task.priority
                    task.priority = priority
                    task.save()
                    TaskHistory.objects.create(
                        task=task,
                        field_changed="priority",
                        old_value=old_p,
                        new_value=priority,
                        changed_by=user
                    )
        
        return Response({'success': True, 'message': f'Updated {len(updates)} tasks'})
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def task_comment_api(request):
    """Add a comment to a task"""
    data = request.data
    task_id = data.get('task_id')
    author_id = data.get('user_id') or data.get('author_id')
    content = data.get('content')

    if not all([task_id, author_id, content]):
        return Response({
            'success': False,
            'message': 'task_id, author_id, and content are required'
        }, status=status.HTTP_400_BAD_REQUEST)

    try:
        task = Task.objects.prefetch_related('assignees').select_related('mentor').get(id=task_id)
        author = Employee.objects.get(id=author_id)

        # Updated permission checks for hybrid assignment
        is_assignee = task.assignees.filter(id=author.id).exists()
        is_mentor_of_assignee = task.assignees.filter(mentors=author).exists()

        can_comment = False
        role = str(author.role).lower()
        if role == 'admin':
            can_comment = True
        elif task.mentor and task.mentor.id == author.id:
            can_comment = True
        elif is_assignee:
            can_comment = True
        elif is_mentor_of_assignee:
            can_comment = True

        if not can_comment:
            return Response({
                'success': False,
                'message': 'You do not have permission to comment on this task'
            }, status=status.HTTP_403_FORBIDDEN)

        comment = TaskComment.objects.create(
            task=task,
            author=author,
            content=content
        )

        # Trigger notifications for mentor and assignees
        try:
            # Notify assignees (if comment is not by them)
            for assignee in task.assignees.all():
                if assignee.id != author.id:
                    _send_task_notification(assignee, f"New comment by {author.name} on: {task.title}", task.id, "task_comment")
            
            # Notify overseers
            for overseer in task.overseers.all():
                if overseer.id != author.id:
                    _send_task_notification(overseer, f"New comment on task: {task.title}", task.id, "task_comment")
            
            if task.mentor and task.mentor.id != author.id and not task.overseers.filter(id=task.mentor.id).exists():
                _send_task_notification(task.mentor, f"New comment on task: {task.title}", task.id, "task_comment")

        except Exception as e:
            print(f"Notification error: {e}")

        return Response({
            'success': True,
            'message': 'Comment added successfully',
            'comment': {
                'id': comment.id,
                'author_name': author.name,
                'content': comment.content,
                'created_at': comment.created_at.isoformat()
            }
        })

    except Task.DoesNotExist:
        return Response({'success': False, 'message': 'Task not found'}, status=status.HTTP_404_NOT_FOUND)
    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'Author not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
@require_gated_token_api
def task_attach_api(request):
    """Attach files to an existing task"""
    task_id = request.data.get('task_id')
    if not task_id:
        return Response({'success': False, 'message': 'task_id required'}, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        task = Task.objects.get(id=task_id)
        files = request.FILES.getlist('attachments')
        for f in files:
            TaskAttachment.objects.create(task=task, file=f)
        return Response({'success': True, 'message': f'Attached {len(files)} files'})
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)





@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def wfh_request_reject(request):
    """Reject WFH request"""
    data = request.data
    request_id = data.get('request_id')
    reason = data.get('reason', '')

    if not request_id:
        return Response({
            'success': False,
            'message': 'Request ID is required'
        }, status=status.HTTP_400_BAD_REQUEST)

    try:
        from .models import EmployeeRequest
        wfh_request = EmployeeRequest.objects.get(id=request_id)
        wfh_request.status = 'rejected'
        wfh_request.admin_response = reason
        wfh_request.reviewed_at = timezone.now()
        # Set reviewed_by to admin user
        admin_user = Employee.objects.filter(role='admin').first()
        if admin_user:
            wfh_request.reviewed_by = admin_user
        wfh_request.save()

        # Notify Employee of rejection
        reviewer_name = wfh_request.reviewed_by.name if wfh_request.reviewed_by else "Admin"
        notif_msg = f"Your WFH request for {wfh_request.start_date} has been rejected by {reviewer_name}."
        if reason:
            notif_msg += f" Note: {reason}"
        _send_task_notification(wfh_request.employee, notif_msg, wfh_request.id, type="request")

        return Response({
            'success': True,
            'message': 'WFH request rejected'
        })
    except EmployeeRequest.DoesNotExist:
        return Response({
            'success': False,
            'message': 'WFH request not found'
        }, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to reject request'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
@require_gated_token_api
def employees_simple_list(request):
    """Get simple list of employees for dropdowns"""
    try:
        employees_qs = Employee.objects.filter(is_active=True).order_by('name')
        employees_data = []
        for emp in employees_qs:
            employees_data.append({
                'id': emp.id,
                'name': emp.name,
                'role': emp.role,
                'mentor_ids': [m.id for m in emp.mentors.all()]
            })
            
        return Response({
            'success': True,
            'employees': employees_data
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': 'Failed to fetch employees'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def wfh_request_approve(request):
    """Approve or reject a Request (WFH or Leave)"""
    try:
        data = request.data
        request_id = data.get('request_id')
        status_val = data.get('status', 'approved')
        admin_response = data.get('admin_response', '')
        reviewer_id = data.get('reviewed_by') 

        try:
            request_obj = EmployeeRequest.objects.get(id=request_id)
        except EmployeeRequest.DoesNotExist:
            return Response({'success': False, 'message': 'Request not found'}, status=status.HTTP_404_NOT_FOUND)

        request_obj.status = status_val
        request_obj.admin_response = admin_response
        request_obj.reviewed_at = timezone.now()

        if reviewer_id:
            try:
                request_obj.reviewed_by = Employee.objects.get(id=reviewer_id)
            except:
                pass

        if not request_obj.reviewed_by:
            admin_user = Employee.objects.filter(role='admin').first()
            if admin_user:
                request_obj.reviewed_by = admin_user

        request_obj.save()

        # Notify Employee of the decision
        reviewer_name = request_obj.reviewed_by.name if request_obj.reviewed_by else "Admin"
        notif_msg = f"Your {request_obj.request_type.upper()} request for {request_obj.start_date} has been {status_val} by {reviewer_name}."
        if admin_response:
            notif_msg += f" Note: {admin_response}"
        _send_task_notification(request_obj.employee, notif_msg, request_obj.id, type="request")

        # If approved, handle based on request type
        if status_val == 'approved':
            req_type = request_obj.request_type
            
            # 1. Determine attendance record settings
            if req_type == 'wfh':
                attendance_status = 'wfh'
                attendance_type = 'wfh'
                is_half = False
            elif req_type == 'full_day':
                attendance_status = 'leave'
                attendance_type = 'office'
                is_half = False
            elif req_type == 'half_day':
                attendance_status = 'half_day'
                attendance_type = 'office'
                is_half = True
            elif req_type == 'optional_holiday':
                # 'holiday' is not in STATUS_CHOICES, use 'leave' with notes
                attendance_status = 'leave'
                attendance_type = 'office'
                is_half = False
            else:
                attendance_status = 'leave'
                attendance_type = 'office'
                is_half = False

            # 2. Update attendance records for the date range
            if req_type != 'unblock_attendance' and req_type != 'task_request':
                from datetime import timedelta
                current_date = request_obj.start_date
                while current_date <= request_obj.end_date:
                    AttendanceRecord.objects.update_or_create(
                        employee=request_obj.employee,
                        date=current_date,
                        defaults={
                            'type': attendance_type,
                            'status': attendance_status,
                            'is_half_day': is_half,
                            'notes': f'Approved request ({req_type})',
                        }
                    )
                    current_date += timedelta(days=1)
            
            # 3. Update Leave Balance for CL if applicable
            if req_type in ['full_day', 'half_day']:
                try:
                    profile = request_obj.employee.profile
                    increment = 1.0 if req_type == 'full_day' else 0.5
                    num_days = (request_obj.end_date - request_obj.start_date).days + 1
                    profile.taken_cl += (increment * num_days)
                    profile.save()
                except Exception as e:
                    print(f"Error updating leave balance: {e}")

        elif status_val == 'rejected':
            # Revert attendance records for the range if they were previously auto-created/updated
            req_type = request_obj.request_type
            if req_type != 'unblock_attendance' and req_type != 'task_request':
                from datetime import timedelta
                current_date = request_obj.start_date
                while current_date <= request_obj.end_date:
                    # Only revert if it matches the current request's expected approved status
                    # to avoid accidentally reverting a 'present' record if someone worked anyway
                    AttendanceRecord.objects.filter(
                        employee=request_obj.employee,
                        date=current_date,
                        status__in=['wfh', 'leave', 'half_day']
                    ).update(status='absent', notes=f'Request rejected ({req_type})')
                    current_date += timedelta(days=1)

        return Response({
            'success': True,
            'message': f'Request {status_val}'
        })
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return Response({
            'success': False,
            'message': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def unblock_attendance(request):
    """Create a request to unblock attendance after 3 consecutive missed checkouts"""
    try:
        data = request.data
        employee_id = data.get('employee_id')
        
        if not employee_id:
            return Response({'success': False, 'message': 'Missing employee ID'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            employee = Employee.objects.get(id=employee_id)
        except Employee.DoesNotExist:
            return Response({'success': False, 'message': 'Employee not found'}, status=status.HTTP_404_NOT_FOUND)

        # Check if an active pending/approved request already exists
        now_local = timezone.localtime(timezone.now()).date()
        recent_req = EmployeeRequest.objects.filter(
            employee=employee,
            request_type='unblock_attendance'
        ).order_by('-created_at').first()

        if recent_req and recent_req.status == 'pending':
            return Response({'success': False, 'message': 'An unblock request is already pending.'}, status=status.HTTP_400_BAD_REQUEST)
        if recent_req and recent_req.status == 'approved' and recent_req.created_at.date() == now_local:
            return Response({'success': False, 'message': 'Attendance is already unblocked for today.'}, status=status.HTTP_400_BAD_REQUEST)

        # Create request using today's date for start and end date as placeholders
        req = EmployeeRequest.objects.create(
            employee=employee,
            request_type='unblock_attendance',
            start_date=now_local,
            end_date=now_local,
            reason='Automated request: 3 consecutive missed check-outs.',
            status='pending'
        )

        # Notify admins
        admins = Employee.objects.filter(role='admin')
        for admin in admins:
            _send_task_notification(admin, f"Unblock request for {employee.name} (Missed Check-outs)", req.id, type="request")

        return Response({'success': True, 'message': 'Unblock request submitted to Admin successfully.'})
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def leave_request(request):
    """Create a new leave request (Full Day or Half Day)"""
    try:
        data = request.data
        employee_id = data.get('employee_id')
        date_str = data.get('date')
        dates_list = data.get('dates', []) # New:支持列表
        r_type = data.get('type') # 'full_day', 'half_day', 'wfh'
        reason = data.get('reason')
        period = data.get('period') # 'first_half', 'second_half'

        if not employee_id or not r_type:
            return Response({'success': False, 'message': 'Missing fields'}, status=status.HTTP_400_BAD_REQUEST)

        # Build list of dates to process
        target_dates = []
        if dates_list:
            target_dates = dates_list
        elif date_str:
            target_dates = [date_str]
        else:
            return Response({'success': False, 'message': 'Date(s) required'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            employee = Employee.objects.get(id=employee_id)
        except Employee.DoesNotExist:
            return Response({'success': False, 'message': 'Employee not found'}, status=status.HTTP_404_NOT_FOUND)

        created_count = 0
        skipped_count = 0

        for d_str in target_dates:
            try:
                req_date = datetime.strptime(d_str, '%Y-%m-%d').date()
                
                # Check existing
                existing = EmployeeRequest.objects.filter(employee=employee, start_date=req_date).first()
                if existing:
                    skipped_count += 1
                    continue

                req = EmployeeRequest.objects.create(
                    employee=employee,
                    request_type=r_type,
                    start_date=req_date,
                    end_date=req_date,
                    reason=reason,
                    status='pending',
                    half_day_period=period if r_type == 'half_day' else None
                )
                
                # Notify mentors
                for mentor in employee.mentors.all():
                    _send_task_notification(mentor, f"{employee.name} requested {r_type.replace('_', ' ')} for {req_date}", req.id, type="request")
                created_count += 1
            except Exception:
                skipped_count += 1

        if created_count == 0 and skipped_count > 0:
            return Response({'success': False, 'message': 'Requests already exist for selected date(s)'}, status=status.HTTP_400_BAD_REQUEST)
        
        return Response({
            'success': True, 
            'message': f'Submitted {created_count} request(s). {skipped_count} skipped.',
            'created_count': created_count,
            'skipped_count': skipped_count
        })
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def leave_request_approve(request):
    """Approve or reject a leave request (admin/mentor only)"""
    # Verify caller is admin or mentor
    caller = get_current_user(request)
    if not caller or caller.role not in ('admin', 'mentor'):
        # Also allow de-facto mentors who have subordinates
        if not caller or not caller.subordinates.exists():
            return Response({'success': False, 'message': 'Admin or Mentor access required'}, status=403)

    try:
        data = request.data
        request_id = data.get('request_id')
        status_val = data.get('status', 'approved') # approved or rejected
        admin_response = data.get('admin_response', '')
        reviewer_id = data.get('reviewed_by')

        try:
            req = EmployeeRequest.objects.get(id=request_id)
        except EmployeeRequest.DoesNotExist:
            return Response({'success': False, 'message': 'Request not found'}, status=status.HTTP_404_NOT_FOUND)

        req.status = status_val
        req.admin_response = admin_response
        req.reviewed_at = timezone.now()
        
        if reviewer_id:
            try:
                req.reviewed_by = Employee.objects.get(id=reviewer_id)
            except:
                pass

        if not req.reviewed_by:
            admin_user = Employee.objects.filter(role='admin').first()
            if admin_user:
                req.reviewed_by = admin_user

        req.save()

        # Notify Employee of the decision
        reviewer_name = req.reviewed_by.name if req.reviewed_by else "Admin"
        notif_msg = f"Your {req.request_type.upper()} request for {req.start_date} has been {status_val} by {reviewer_name}."
        if admin_response:
            notif_msg += f" Note: {admin_response}"
        _send_task_notification(req.employee, notif_msg, req.id, type="request")

        # If approved, create or update AttendanceRecord to reflect in calendar
        if status_val == 'approved':
            req_type = req.request_type
            
            # Unblock requests do not create attendance records, they just lift the check-in block
            if req_type != 'unblock_attendance' and req_type != 'task_request':
                if req_type == 'wfh':
                    attendance_status = 'wfh'
                    attendance_type = 'wfh'
                elif req_type == 'full_day':
                    attendance_status = 'leave'
                    attendance_type = 'office'
                elif req_type == 'half_day':
                    attendance_status = 'half_day'
                    attendance_type = 'office'
                elif req_type == 'optional_holiday':
                    attendance_status = 'leave' # 'holiday' is not in STATUS_CHOICES
                    attendance_type = 'office'
                else:
                    attendance_status = 'leave'
                    attendance_type = 'office'

                # Create or update attendance record for each day in the request date range
                from datetime import timedelta
                current_date = req.start_date
                while current_date <= req.end_date:
                    AttendanceRecord.objects.update_or_create(
                        employee=req.employee,
                        date=current_date,
                        defaults={
                            'type': attendance_type,
                            'status': attendance_status,
                            'is_half_day': (req_type == 'half_day'),
                            'notes': f'Approved {req_type} request',
                        }
                    )
                    current_date += timedelta(days=1)

                # Update Leave Balance for CL
                if req_type in ['full_day', 'half_day']:
                    try:
                        profile = req.employee.profile
                        increment = 1.0 if req_type == 'full_day' else 0.5
                        num_days = (req.end_date - req.start_date).days + 1
                        profile.taken_cl += (increment * num_days)
                        profile.save()
                    except Exception as e:
                        print(f"Error updating leave balance: {e}")
        
        elif status_val == 'rejected':
            # Revert attendance records if previously approved
            req_type = req.request_type
            if req_type != 'unblock_attendance' and req_type != 'task_request':
                from datetime import timedelta
                current_date = req.start_date
                while current_date <= req.end_date:
                    AttendanceRecord.objects.filter(
                        employee=req.employee,
                        date=current_date,
                        status__in=['wfh', 'leave', 'half_day']
                    ).update(status='absent', notes=f'Request rejected ({req_type})')
                    current_date += timedelta(days=1)

        return Response({'success': True, 'message': f'Request {status_val}'})
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



@api_view(['GET'])
@require_gated_token_api
def attendance_predictions(request):
    """Get AI-powered attendance predictions for all employees (Admin only)"""
    try:
        # Check if user is admin
        employee_id = request.GET.get('employee_id')
        if not employee_id:
            return Response({
                'success': False,
                'message': 'Employee ID is required'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            employee = Employee.objects.get(id=employee_id)
            if employee.role != 'admin':
                return Response({
                    'success': False,
                    'message': 'Unauthorized. Admin access required.'
                }, status=status.HTTP_403_FORBIDDEN)
        except Employee.DoesNotExist:
            return Response({
                'success': False,
                'message': 'Employee not found'
            }, status=status.HTTP_404_NOT_FOUND)
        
        # Import prediction engine
        from .attendance_prediction import get_all_employees_predictions
        
        # Get predictions for all employees
        predictions = get_all_employees_predictions()
        
        return Response({
            'success': True,
            'count': len(predictions),
            'predictions': predictions
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({
            'success': False,
            'message': f'Failed to generate predictions: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# ========== Company Predictive Report API ==========

@api_view(['GET'])
@require_gated_token_api
def company_predictive_report(request):
    """
    Generate a company-wide predictive attendance report.
    Query params:
        start_date  (YYYY-MM-DD)  — default 30 days ago
        end_date    (YYYY-MM-DD)  — default today
    Returns summary + per-employee predictive metrics.
    """
    try:
        from datetime import date, timedelta, datetime as dt
        from .models import Employee, AttendanceRecord

        today = timezone.localtime(timezone.now()).date()

        start_str = request.GET.get('start_date')
        end_str   = request.GET.get('end_date')

        try:
            start_date = dt.strptime(start_str, '%Y-%m-%d').date() if start_str else today - timedelta(days=30)
            end_date   = dt.strptime(end_str,   '%Y-%m-%d').date() if end_str   else today
        except ValueError:
            return Response({'success': False, 'message': 'Invalid date format. Use YYYY-MM-DD.'}, status=400)

        # Enforce max 30-day window
        if (end_date - start_date).days > 30:
            return Response({'success': False, 'message': 'Date range cannot exceed 30 days.'}, status=400)
        if end_date > today:
            end_date = today

        total_days   = (end_date - start_date).days + 1
        working_days = sum(1 for i in range(total_days)
                           if Holiday.is_date_working(start_date + timedelta(days=i)))

        employees = Employee.objects.filter(is_active=True).exclude(role='admin').order_by('department', 'name')

        # Aggregate all records in range in one query
        records_qs = AttendanceRecord.objects.filter(
            date__range=[start_date, end_date],
            employee__is_active=True
        ).select_related('employee').values(
            'employee_id', 'employee__name', 'employee__department',
            'date', 'status', 'type', 'check_in_time', 'check_out_time', 'total_hours'
        )

        # Group by employee
        from collections import defaultdict
        emp_records = defaultdict(list)
        for r in records_qs:
            emp_records[r['employee_id']].append(r)

        # Company totals
        company_present = 0
        company_wfh     = 0
        company_absent  = 0
        company_leave   = 0
        company_half    = 0

        per_employee = []

        for emp in employees:
            recs = emp_records.get(emp.id, [])
            present  = sum(1 for r in recs if r['status'] == 'present')
            wfh      = sum(1 for r in recs if r['status'] == 'wfh')
            half_day = sum(1 for r in recs if r['status'] == 'half_day')
            leave    = sum(1 for r in recs if r['status'] == 'leave')
            absent   = sum(1 for r in recs if r['status'] == 'absent')
            attended = present + wfh + half_day

            att_rate = round((attended / working_days) * 100) if working_days > 0 else 0

            # Check-in times for avg
            check_in_times = [r['check_in_time'] for r in recs if r['check_in_time']]
            avg_checkin = None
            if check_in_times:
                total_mins = sum(t.hour * 60 + t.minute for t in check_in_times)
                avg_m = total_mins // len(check_in_times)
                avg_checkin = f"{avg_m // 60:02d}:{avg_m % 60:02d}"

            # Total hours
            total_hours = sum(float(r['total_hours'] or 0) for r in recs)
            avg_hours   = round(total_hours / attended, 1) if attended > 0 else 0

            # Simple trend: compare first half vs second half of period
            mid = start_date + timedelta(days=total_days // 2)
            first_half  = sum(1 for r in recs if r['date'] < mid and r['status'] in ('present','wfh','half_day'))
            second_half = sum(1 for r in recs if r['date'] >= mid and r['status'] in ('present','wfh','half_day'))
            trend = 'up' if second_half > first_half else ('down' if second_half < first_half else 'stable')

            # Predictive likelihood (simple heuristic based on recent rate)
            likelihood = min(int(att_rate * 1.05), 100)

            # Daily attendance rate time-series (for mini chart)
            day_series = []
            for i in range(total_days):
                d = start_date + timedelta(days=i)
                if d.weekday() >= 6:  # skip Sunday
                    continue
                day_rec = next((r for r in recs if r['date'] == d), None)
                status  = day_rec['status'] if day_rec else 'absent'
                day_series.append({
                    'date':    str(d),
                    'day':     d.strftime('%a'),
                    'status':  status,
                    'present': 1 if status in ('present', 'wfh', 'half_day') else 0
                })

            company_present += present
            company_wfh     += wfh
            company_absent  += absent
            company_leave   += leave
            company_half    += half_day

            # Advanced Predictions from Engine
            from .attendance_prediction import AttendancePredictionEngine
            engine = AttendancePredictionEngine(emp.id)
            leave_probs = engine.predict_leaves()
            # Average leave probability for this employee
            avg_leave_prob = round(sum(leave_probs.values()) / 7, 2)
            predicted_hrs_val = engine.predict_working_hours()

            per_employee.append({
                'id':          emp.id,
                'name':        emp.name,
                'department':  emp.department,
                'role':        emp.role,
                'att_rate':    att_rate,
                'attended':    attended,
                'present':     present,
                'wfh':         wfh,
                'half_day':    half_day,
                'leave':       leave,
                'absent':      absent,
                'avg_checkin': avg_checkin,
                'avg_hours':   avg_hours,
                'total_hours': round(total_hours, 1),
                'trend':       trend,
                'likelihood':  likelihood,
                'day_series':  day_series,
                'leave_probability': avg_leave_prob,
                'predicted_hours': predicted_hrs_val
            })

        total_emp  = len(per_employee)
        avg_att    = round(sum(e['att_rate'] for e in per_employee) / total_emp) if total_emp else 0

        return Response({
            'success': True,
            'period': {
                'start_date':    str(start_date),
                'end_date':      str(end_date),
                'total_days':    total_days,
                'working_days':  working_days,
            },
            'company_summary': {
                'total_employees': total_emp,
                'avg_attendance':  avg_att,
                'total_present':   company_present,
                'total_wfh':       company_wfh,
                'total_absent':    company_absent,
                'total_leave':     company_leave,
                'total_half_day':  company_half,
            },
            'employees': per_employee,
        })
    except Exception as e:
        import traceback; traceback.print_exc()
        return Response({'success': False, 'message': str(e)}, status=500)


# ========== Intelligence Hub API Endpoints ==========

@api_view(['GET'])
@require_gated_token_api
def intelligence_hub_forecast(request):
    """Get current attendance forecast with confidence and trend"""
    try:
        from .intelligence_hub import calculate_forecast, get_current_day_name, load_model_state, SLMInsightGenerator
        
        forecast, confidence, trend = calculate_forecast()
        day_name = get_current_day_name()
        model_state = load_model_state()
        
        # Note: Model training is now handled by a system-level Cron job at 6:30 PM daily.
        # Check /scripts/train_model.sh for details.
        
        employee_id = request.GET.get('employee_id')
        if employee_id:
            try:
                from .models import AttendanceRecord
                from datetime import date, timedelta
                
                today = date.today()
                start_date = today - timedelta(days=30)
                
                records = AttendanceRecord.objects.filter(
                    employee_id=employee_id,
                    date__range=[start_date, today]
                )
                
                passed_days = (today - start_date).days + 1
                # Optimization: Fetch all holidays in range once to avoid 30 queries in a loop
                holiday_map = {h.date: h for h in Holiday.objects.filter(date__range=[start_date, today])}
                working_days_passed = 0
                for d in range(passed_days):
                    curr_date = start_date + timedelta(days=d)
                    h = holiday_map.get(curr_date)
                    is_working = False
                    if h:
                        if h.is_working_day or h.is_optional:
                            is_working = True
                    else:
                        is_working = curr_date.weekday() < 6
                    if is_working:
                        working_days_passed += 1
                
                weekday_present_days = records.filter(
                    date__week_day__in=[2, 3, 4, 5, 6, 7],
                    status__in=['present', 'half_day', 'wfh', 'client']
                ).count()
                
                if working_days_passed > 0:
                    forecast = round((weekday_present_days / working_days_passed) * 100)
                else:
                    forecast = 0
            except Exception as e:
                import logging
                logging.getLogger('attendance').error(f"Error calculating personal forecast: {e}")
        
        return Response({
            'success': True,
            'forecast': {
                'percentage': forecast,
                'confidence': confidence,
                'trend': trend,
                'day_name': day_name,
                'subtitle': f"{day_name}'s Forecast",
                'model_state': model_state,
                'ai_insight': SLMInsightGenerator.generate_insight({
                    'forecast': forecast,
                    'confidence': confidence,
                    'trend': trend,
                    'attendance_streak': model_state.get('attendance_streak', 0) if model_state else 0
                })
            }
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({
            'success': False,
            'message': f'Failed to calculate forecast: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def intelligence_hub_trends(request):
    """Get 30-day trend data with comprehensive company overview"""
    try:
        from .intelligence_hub import get_company_overview
        
        days = int(request.GET.get('days', 30))
        predict_days = int(request.GET.get('predict_days', 3))
        overview_data = get_company_overview(days, predict_days)
        
        return Response({
            'success': True,
            **overview_data  # Unpacks summary, departments, employees, trends
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({
            'success': False,
            'message': f'Failed to get trend data: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def intelligence_hub_search(request):
    """Search personnel with attendance predictions"""
    try:
        from .intelligence_hub import search_personnel
        
        data = request.data
        query = data.get('query')
        department = data.get('department')
        min_attendance = data.get('min_attendance')
        max_attendance = data.get('max_attendance')
        mentor_id = data.get('mentor_id')
        
        results = search_personnel(query, department, min_attendance, max_attendance, mentor_id)
        
        return Response({
            'success': True,
            'count': len(results),
            'results': results
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({
            'success': False,
            'message': f'Failed to search personnel: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def employee_hr_report(request, employee_id):
    """
    Generate a comprehensive HR attendance report for a single employee
    over a given date range (max 31 days). Returns all the metrics
    needed for a professional PDF report.
    """
    try:
        employee = Employee.objects.get(id=employee_id)
        today = date.today()

        # Parse date range — max 31 days, default last 30 days
        start_str = request.GET.get('start_date')
        end_str   = request.GET.get('end_date')

        try:
            start_date = datetime.strptime(start_str, '%Y-%m-%d').date() if start_str else today - timedelta(days=30)
            end_date   = datetime.strptime(end_str,   '%Y-%m-%d').date() if end_str   else today
        except ValueError:
            start_date = today - timedelta(days=30)
            end_date   = today

        # Cap the range to 31 days
        if (end_date - start_date).days > 30:
            start_date = end_date - timedelta(days=30)

        # Clamp end_date to today
        if end_date > today:
            end_date = today

        # ── Attendance Records ──────────────────────────────────────────────
        records = AttendanceRecord.objects.filter(
            employee=employee,
            date__range=[start_date, end_date]
        ).order_by('date')

        now_local = timezone.localtime(timezone.now())

        # Count breakdown and build daily log
        daily_log   = []
        total_hours = 0.0
        status_counts = {
            'present': 0,
            'absent':  0,
            'leave':   0,
            'wfh':     0,
            'half_day':0,
        }

        check_in_seconds_list  = []
        check_out_seconds_list = []
        late_days = 0
        PUNCTUAL_THRESHOLD_H  = 10   # After 10:00 AM = late
        PUNCTUAL_THRESHOLD_M  = 0

        for r in records:
            hours = float(r.total_hours or 0)
            if not r.check_out_time and r.check_in_time and r.date == now_local.date():
                try:
                    ci = datetime.strptime(str(r.check_in_time), '%H:%M:%S').time()
                    ci_dt = timezone.make_aware(datetime.combine(r.date, ci))
                    hours = round(min(max(0.0, (now_local - ci_dt).total_seconds() / 3600), 14.0), 2)
                except Exception:
                    pass

            total_hours += hours

            # Status bucket
            if r.type == 'wfh' and r.status in ['present', 'half_day', 'wfh', 'client']:
                status_counts['wfh'] += 1
            elif r.status == 'present':
                status_counts['present'] += 1
            elif r.status == 'absent':
                status_counts['absent'] += 1
            elif r.status == 'leave':
                status_counts['leave'] += 1
            elif r.status == 'half_day':
                status_counts['half_day'] += 1
            else:
                status_counts['absent'] += 1

            # Punctuality
            if r.check_in_time:
                sec = r.check_in_time.hour * 3600 + r.check_in_time.minute * 60 + r.check_in_time.second
                check_in_seconds_list.append(sec)
                threshold_sec = PUNCTUAL_THRESHOLD_H * 3600 + PUNCTUAL_THRESHOLD_M * 60
                if sec > threshold_sec:
                    late_days += 1

            if r.check_out_time:
                sec = r.check_out_time.hour * 3600 + r.check_out_time.minute * 60 + r.check_out_time.second
                check_out_seconds_list.append(sec)

            daily_log.append({
                'date':        r.date.strftime('%d %b %Y'),
                'day':         r.date.strftime('%A'),
                'status':      r.status,
                'type':        r.type,
                'check_in':    str(r.check_in_time)[:5] if r.check_in_time else '—',
                'check_out':   str(r.check_out_time)[:5] if r.check_out_time else '—',
                'hours':       round(hours, 2),
            })

        # ── Summary Metrics ─────────────────────────────────────────────────
        total_days  = (end_date - start_date).days + 1
        working_days = sum(
            1 for i in range(total_days)
            if Holiday.is_date_working(start_date + timedelta(days=i))
        )

        raw_attended = status_counts['present'] + status_counts['wfh'] + status_counts['half_day']
        attended_days = min(working_days, raw_attended)
        attendance_rate = min(100.0, round((raw_attended / working_days * 100), 1)) if working_days else 0

        avg_check_in  = None
        avg_check_out = None
        if check_in_seconds_list:
            s = sum(check_in_seconds_list) / len(check_in_seconds_list)
            avg_check_in = f"{int(s//3600):02d}:{int((s%3600)//60):02d}"
        if check_out_seconds_list:
            s = sum(check_out_seconds_list) / len(check_out_seconds_list)
            avg_check_out = f"{int(s//3600):02d}:{int((s%3600)//60):02d}"

        punctual_days = attended_days - late_days
        punctuality_rate = round((punctual_days / attended_days * 100), 1) if attended_days else 0

        avg_hours_per_day = round(total_hours / attended_days, 2) if attended_days else 0

        # ── Task Performance (optional, best-effort) ────────────────────────
        try:
            tasks_qs = Task.objects.filter(assignees=employee).filter(
                Q(created_at__date__range=[start_date, end_date]) |
                Q(completed_at__date__range=[start_date, end_date]) |
                Q(status__in=['todo', 'in_progress'])
            ).distinct()
            total_tasks     = tasks_qs.count()
            completed_tasks = tasks_qs.filter(status='completed').count()
            task_completion_rate = round((completed_tasks / total_tasks * 100), 1) if total_tasks else 0
        except Exception:
            total_tasks = completed_tasks = task_completion_rate = 0

        # ── Profile ─────────────────────────────────────────────────────────
        profile = getattr(employee, 'profile', None)
        designation = getattr(profile, 'designation', '') or ''

        return Response({
            'success': True,
            'report': {
                'employee': {
                    'id':          employee.id,
                    'name':        employee.name,
                    'username':    employee.username,
                    'email':       employee.email,
                    'department':  employee.department,
                    'designation': designation,
                    'avatar_emoji': profile.avatar_emoji if profile else '👤',
                },
                'period': {
                    'start_date':  str(start_date),
                    'end_date':    str(end_date),
                    'total_days':  total_days,
                    'working_days': working_days,
                },
                'summary': {
                    'present':          status_counts['present'],
                    'wfh':              status_counts['wfh'],
                    'half_day':         status_counts['half_day'],
                    'leave':            status_counts['leave'],
                    'absent':           status_counts['absent'],
                    'attended_days':    attended_days,
                    'attendance_rate':  attendance_rate,
                    'total_hours':      round(total_hours, 1),
                    'avg_hours_per_day': avg_hours_per_day,
                    'avg_check_in':     avg_check_in or '—',
                    'avg_check_out':    avg_check_out or '—',
                    'late_days':        late_days,
                    'punctual_days':    punctual_days,
                    'punctuality_rate': punctuality_rate,
                    'task_total':       total_tasks,
                    'task_completed':   completed_tasks,
                    'task_completion_rate': task_completion_rate,
                },
                'daily_log': daily_log,
            }
        })

    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'Employee not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@require_gated_token_api
def intelligence_hub_train(request):
    """Trigger training of the forecast model using all historical data"""
    try:
        from .intelligence_hub import train_forecast_model
        
        user_id = request.data.get('user_id')
        user = Employee.objects.filter(id=user_id).first()
        
        result = train_forecast_model()
        
        if result['success']:
            # Create a localized log entry
            summary = result['summary']
            TrainingLog.objects.create(
                trained_by=user,
                data_points=summary.get('data_points', 0),
                average_rate=summary.get('average_rate', 0.0),
                stability_factor=summary.get('stability_factor', 0.0),
                logs=result.get('logs', []),
                summary=summary
            )
            
            return Response({
                'success': True,
                'message': 'Model trained successfully',
                'summary': summary,
                'logs': result.get('logs', [])
            })
        else:
            return Response({
                'success': False,
                'message': result['message'],
                'logs': result.get('logs', [])
            }, status=status.HTTP_400_BAD_REQUEST)
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({
            'success': False,
            'message': f'Training failed: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@require_gated_token_api
def intelligence_hub_training_history(request):
    """Fetch recent model training history"""
    try:
        logs = TrainingLog.objects.all().select_related('trained_by')[:10]
        data = [{
            'id': log.id,
            'timestamp': log.timestamp.strftime('%Y-%m-%d %H:%M:%S'),
            'trained_by_name': log.trained_by.name if log.trained_by else 'System',
            'data_points': log.data_points,
            'average_rate': log.average_rate,
            'stability_factor': log.stability_factor,
            'summary': log.summary
        } for log in logs]
        
        return Response({
            'success': True,
            'history': data
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': f'Failed to fetch history: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def clear_training_history(request):
    """Clear all model training history"""
    try:
        user_id = request.data.get('user_id')
        if not user_id:
            return Response({'success': False, 'message': 'User ID required'}, status=status.HTTP_400_BAD_REQUEST)
        
        user = Employee.objects.filter(id=user_id).first()
        if not user or user.role != 'admin':
            return Response({'success': False, 'message': 'Unauthorized to clear history'}, status=status.HTTP_403_FORBIDDEN)
        
        # Delete all training logs
        TrainingLog.objects.all().delete()
        
        return Response({
            'success': True,
            'message': 'Training history cleared successfully'
        })
    except Exception as e:
        return Response({
            'success': False,
            'message': f'Failed to clear history: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



@api_view(['GET', 'POST', 'DELETE'])
@require_gated_token_api
@parser_classes([JSONParser])
def temporary_tags_api(request):
    """API for managing temporary tags"""
    print(f"DEBUG: temporary_tags_api method={request.method}")
    if request.method == 'GET':
        employee_id = request.query_params.get('employee_id')
        tags = TemporaryTag.objects.all().select_related('employee')
        if employee_id:
            tags = tags.filter(employee_id=employee_id)
        
        tags_data = [{
            'id': tag.id,
            'employee_id': tag.employee.id,
            'employee_username': tag.employee.username,
            'employee_name': tag.employee.name,
            'department': tag.department,
            'role': tag.role,
            'start_date': str(tag.start_date),
            'end_date': str(tag.end_date),
            'created_at': tag.created_at.isoformat(),
        } for tag in tags.order_by('-created_at')]
        
        return Response({'success': True, 'tags': tags_data})

    elif request.method == 'POST':
        data = request.data
        print(f"DEBUG: temporary_tags_api POST data={data}")
        try:
            employee_id = data.get('employee_id')
            department = data.get('department')
            role = data.get('role')
            start_date = data.get('start_date')
            end_date = data.get('end_date')
            
            print(f"DEBUG: Creating tag for employee_id={employee_id}, dept={department}, role={role}, range={start_date} to {end_date}")
            
            employee = Employee.objects.get(id=employee_id)
            tag = TemporaryTag.objects.create(
                employee=employee,
                department=data.get('department'),
                role=data.get('role'),
                start_date=data.get('start_date'),
                end_date=data.get('end_date')
            )
            return Response({
                'success': True,
                'message': 'Temporary tag created successfully',
                'tag_id': tag.id
            })
        except Employee.DoesNotExist:
            return Response({'success': False, 'message': 'Employee not found'}, status=404)
        except Exception as e:
            return Response({'success': False, 'message': str(e)}, status=400)

    elif request.method == 'DELETE':
        tag_id = request.query_params.get('id') or request.data.get('id')
        try:
            tag = TemporaryTag.objects.get(id=tag_id)
            tag.delete()
            return Response({'success': True, 'message': 'Temporary tag deleted successfully'})
        except TemporaryTag.DoesNotExist:
            return Response({'success': False, 'message': 'Tag not found'}, status=404)
        except Exception as e:
            return Response({'success': False, 'message': str(e)}, status=400)


@api_view(['POST'])
@parser_classes([JSONParser])
def verify_token(request):
    """Verify attendance token from portal and return user data if valid"""
    token = request.data.get('token')
    if not token:
        return Response({'success': False, 'message': 'Token is required'}, status=status.HTTP_400_BAD_REQUEST)

    # 1. Try Gated Access verification (itsdangerous) - Priority for auto-login
    from .security import validate_gated_token
    success, result = validate_gated_token(token)
    if success:
        user_id = result.get('user_id')
        username = result.get('username')
        employee = None
        
        if username:
            employee = Employee.objects.filter(username=username, is_active=True).first()
        if not employee and user_id:
            employee = Employee.objects.filter(id=user_id, is_active=True).first()

        if employee:
            profile = EmployeeProfile.objects.filter(employee=employee).first()
            assignment = employee.get_current_assignment()
            user_data = {
                'id': employee.id,
                'username': employee.username,
                'name': employee.name,
                'email': employee.email,
                'phone': employee.phone,
                'department': assignment['department'],
                'primary_office': employee.primary_office,
                'role': assignment['role'],
                'is_temporary': assignment['is_temporary'],
                'has_subordinates': employee.subordinates.exists(),
                'gender': profile.gender if profile else None,
                'date_of_birth': str(profile.date_of_birth) if profile and profile.date_of_birth else None,
                'avatar_emoji': profile.avatar_emoji if profile else "👤",
                'avatar_url': profile.avatar_url if profile else None,
                'theme_settings': profile.theme_settings if profile else {},
                'mentors': [{'id': m.id, 'name': m.name} for m in employee.mentors.all()],
                'total_cl': profile.total_cl if profile else 12,
                'taken_cl': profile.taken_cl if profile else 0,
            }
            return Response({
                'success': True, 
                'message': 'Token verified (Gated)',
                'user': user_data
            })
        else:
            return Response({'success': False, 'message': 'User associated with token not found'}, status=404)

    # 2. Try HMAC verification (The legacy way)
    secret = getattr(settings, "ATTENDANCE_SECRET_KEY", "hanuai-attendance-secret-shared-key").encode()
    message = timezone.localtime(timezone.now()).strftime("%Y-%m-%d").encode()
    expected_hmac = hmac.new(secret, message, hashlib.sha256).hexdigest()
    
    if hmac.compare_digest(token, expected_hmac):
        return Response({'success': True, 'message': 'Token verified (HMAC)'})
    
    return Response({'success': False, 'message': f'Invalid token: {result}'}, status=status.HTTP_401_UNAUTHORIZED)


def error_400_view(request, exception=None):
    """Custom 400 Bad Request handler"""
    return render(request, '400.html', status=400)


def error_403_view(request, exception=None, message="None Provided"):
    """Custom 403 Forbidden handler with diagnostic info"""
    return render(request, '403.html', {'error_message': message}, status=403)


def error_404_view(request, exception=None):
    """Custom 404 Not Found handler"""
    return render(request, '404.html', status=404)


def error_500_view(request):
    """Custom 500 Server Error handler"""
    return render(request, '500.html', status=500)


@require_valid_token
def spa_view(request):
    """Protected view to serve the SPA index.html."""
    host = request.get_host()
    # Check if we are running in development (localhost/127.0.0.1)
    is_development = '127.0.0.1' in host or 'localhost' in host
    
    # Check if user is attached by the decorator
    is_authenticated = hasattr(request, 'user') and isinstance(request.user, Employee)
    
    context = {
        'maps_api_key': settings.MAPS_API_KEY,
        'gated_token': request.GET.get('token'),
        'is_development': is_development,
        'is_authenticated': is_authenticated
    }
    return render(request, 'index.html', context)



@require_valid_token
def gated_dashboard(request):
    """Entry point for gated access with token from portal."""
    token_str = request.GET.get('token')
    from .security import validate_gated_token
    success, data = validate_gated_token(token_str)
    
    host = request.get_host()
    is_development = '127.0.0.1' in host or 'localhost' in host
    
    # User should already be attached by @require_valid_token
    is_authenticated = hasattr(request, 'user') and isinstance(request.user, Employee)

    context = {
        'maps_api_key': settings.MAPS_API_KEY,
        'gated_token': token_str,
        'is_development': is_development,
        'is_authenticated': is_authenticated
    }
    
    if success:
        context['gated_user_id'] = data.get('user_id')
        context['is_gated'] = True
        
    return render(request, 'index.html', context)


@api_view(['GET'])
@require_gated_token_api
def employee_list_summary(request):
    """
    Returns a simplified list of active employees including:
    name, department, phone, and role.
    """
    try:
        employees = Employee.objects.filter(is_active=True).values(
            'name', 'department', 'phone', 'role'
        ).order_by('name')
        return Response({
            'success': True,
            'employees': list(employees)
        }, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({
            'success': False,
            'message': f'Failed to fetch employee list: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
@require_gated_token_api
def avatar_assets_list(request):
    """Returns available assets grouped by category"""
    category = request.GET.get('category')
    assets = AvatarAsset.objects.filter(is_active=True)
    if category:
        assets = assets.filter(category=category)
    
    serializer = AvatarAssetSerializer(assets, many=True)
    return Response({
        'success': True,
        'assets': serializer.data
    })

@api_view(['GET', 'POST', 'PUT'])
@require_gated_token_api
def user_memoji_api(request, user_id=None):
    """Retrieve or update a user's memoji configuration"""
    if not user_id:
        user_id = request.data.get('user_id') or request.GET.get('user_id')
    
    if not user_id:
        return Response({'success': False, 'message': 'User ID required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        employee = Employee.objects.get(id=user_id)
    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'Employee not found'}, status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        memoji = Memoji.objects.filter(employee=employee).first()
        if not memoji:
            return Response({'success': True, 'memoji': None})
        serializer = MemojiSerializer(memoji)
        return Response({'success': True, 'memoji': serializer.data})

    elif request.method in ['POST', 'PUT']:
        avatar_config = request.data.get('avatar_config')
        avatar_url = request.data.get('avatar_url')
        
        if not avatar_config and not avatar_url:
            return Response({'success': False, 'message': 'Avatar configuration or URL required'}, status=status.HTTP_400_BAD_REQUEST)
        
        defaults = {}
        if avatar_config:
            defaults['avatar_config'] = avatar_config
        if avatar_url:
            defaults['avatar_url'] = avatar_url
            
        memoji, created = Memoji.objects.update_or_create(
            employee=employee,
            defaults=defaults
        )
        
        # Sync with Employee profile avatar fields for global rendering
        profile = getattr(employee, 'profile', None)
        if profile:
            if avatar_url:
                profile.avatar_url = avatar_url
            if avatar_config:
                import json
                profile.theme_settings['avatar_config'] = avatar_config
            profile.save()
        
        serializer = MemojiSerializer(memoji)
        return Response({
            'success': True, 
            'message': 'Avatar saved successfully',
            'memoji': serializer.data
        })


@api_view(['POST'])
@require_gated_token_api
@csrf_exempt
@parser_classes([MultiPartParser, FormParser])
def upload_avatar(request):
    """Upload a custom photo for user avatar"""
    employee_id = request.POST.get('employee_id')
    if not employee_id:
        return Response({'success': False, 'message': 'Employee ID is required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        employee = Employee.objects.get(id=employee_id)
        profile, _ = EmployeeProfile.objects.get_or_create(employee=employee)
    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'Employee not found'}, status=status.HTTP_404_NOT_FOUND)

    if 'photo' not in request.FILES:
        return Response({'success': False, 'message': 'No photo file provided'}, status=status.HTTP_400_BAD_REQUEST)

    photo = request.FILES['photo']
    
    # Validation
    if photo.size > 2 * 1024 * 1024: # 2MB limit
        return Response({'success': False, 'message': 'Photo size exceeds 2MB limit'}, status=status.HTTP_400_BAD_REQUEST)
    
    if not photo.content_type.startswith('image/'):
        return Response({'success': False, 'message': 'Invalid file type. Please upload an image.'}, status=status.HTTP_400_BAD_REQUEST)

    # Save file
    upload_dir = os.path.join(settings.MEDIA_ROOT, 'avatars')
    os.makedirs(upload_dir, exist_ok=True)
    
    ext = os.path.splitext(photo.name)[1].lower()
    filename = f"avatar_{employee_id}_{uuid.uuid4().hex[:8]}{ext}"
    file_path = os.path.join(upload_dir, filename)
    
    with open(file_path, 'wb') as f:
        for chunk in photo.chunks():
            f.write(chunk)
            
    # Update profile
    previous_photo = profile.avatar_url
    profile.avatar_url = f'avatars/{filename}'
    profile.avatar_emoji = "👤" # Reset emoji if used
    profile.theme_settings.pop('avatar_config', None) # Clear 3D config
    profile.save()
    
    # Delete previous custom avatar if exists and is not default
    if previous_photo and os.path.exists(os.path.join(settings.MEDIA_ROOT, previous_photo)):
        try:
            os.remove(os.path.join(settings.MEDIA_ROOT, previous_photo))
        except:
            pass

    return Response({
        'success': True,
        'message': 'Avatar photo uploaded successfully',
        'avatar_url': profile.avatar_url
    })


@api_view(['GET'])
@require_gated_token_api
@parser_classes([JSONParser])
def mentor_status(request):
    """Get today's attendance status for all mentors of the current employee"""
    employee_id = request.query_params.get('employee_id')
    
    if not employee_id:
        return Response({'success': False, 'message': 'Employee ID is required'}, status=status.HTTP_400_BAD_REQUEST)
        
    try:
        employee = Employee.objects.get(id=employee_id)
        mentors = employee.mentors.all()
        
        mentor_data = []
        today = timezone.localtime(timezone.now()).date()
        
        for mentor in mentors:
            status = 'Absent'
            record = AttendanceRecord.objects.filter(employee=mentor, date=today).first()
            if record:
                status = record.status
                
            try:
                profile = EmployeeProfile.objects.get(employee=mentor)
                avatar = profile.avatar_url or profile.avatar_emoji or '👤'
                bg = profile.theme_settings.get('avatarTextBg', '#2563eb') if profile.theme_settings else '#2563eb'
            except EmployeeProfile.DoesNotExist:
                avatar = '👤'
                bg = '#2563eb'
                
            mentor_data.append({
                'id': mentor.id,
                'name': mentor.name,
                'role': mentor.role,
                'status': status,
                'avatar': avatar,
                'bg': bg
            })
            
        return Response({
            'success': True,
            'mentors': mentor_data
        })
        
    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'Employee not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# ========== Web Push Notification API ==========

@api_view(['GET'])
@require_gated_token_api
def get_vapid_public_key(request):
    """Return the VAPID public key so the browser can subscribe to push notifications."""
    from django.conf import settings
    key = getattr(settings, 'VAPID_PUBLIC_KEY', '')
    if not key:
        return Response({'success': False, 'message': 'VAPID not configured'}, status=status.HTTP_503_SERVICE_UNAVAILABLE)
    return Response({'success': True, 'public_key': key})


@api_view(['POST', 'DELETE'])
@require_gated_token_api
@parser_classes([JSONParser])
def save_push_subscription(request):
    """
    POST  → Save (or update) a browser push subscription for an employee.
    DELETE → Remove a subscription (for unsubscribe / logout).
    Body: { employee_id, endpoint, p256dh, auth }
    """
    from attendance.models import PushSubscription

    data = request.data
    employee_id = data.get('employee_id')
    endpoint = data.get('endpoint')

    if not employee_id or not endpoint:
        return Response({'success': False, 'message': 'employee_id and endpoint are required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        employee = Employee.objects.get(id=employee_id)
    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'Employee not found'}, status=status.HTTP_404_NOT_FOUND)

    if request.method == 'DELETE':
        PushSubscription.objects.filter(employee=employee, endpoint=endpoint).delete()
        return Response({'success': True, 'message': 'Push subscription removed'})

    # POST — save or update
    p256dh = data.get('p256dh', '')
    auth   = data.get('auth', '')
    ua     = request.META.get('HTTP_USER_AGENT', '')[:255]

    PushSubscription.objects.update_or_create(
        endpoint=endpoint,
        defaults={
            'employee': employee,
            'p256dh': p256dh,
            'auth': auth,
            'user_agent': ua,
        }
    )
    return Response({'success': True, 'message': 'Push subscription saved'})


@csrf_exempt
def service_worker_view(request):
    """
    Serve the Service Worker script from the root with the required 
    Service-Worker-Allowed header to allow it to control the entire site scope.
    """
    try:
        #sw_path = os.path.join(settings.BASE_DIR, 'static', 'sw.js')
        # Optimized: use STATICFILES_DIRS if available
        if hasattr(settings, 'STATICFILES_DIRS') and settings.STATICFILES_DIRS:
            sw_path = os.path.join(settings.STATICFILES_DIRS[0], 'sw.js')
        else:
            sw_path = os.path.join(settings.BASE_DIR, 'static', 'sw.js')
            
        with open(sw_path, 'rb') as f:
            content = f.read()
        
        response = HttpResponse(content, content_type='application/javascript')
        response['Service-Worker-Allowed'] = '/'
        return response
    except Exception as e:
        return HttpResponse(f"Service Worker not found: {str(e)}", status=404)


# ─────────────────────────────────────────────────────────────────────────────
#  HOLIDAY MANAGEMENT VIEWS
# ─────────────────────────────────────────────────────────────────────────────

from .models import Holiday, HolidayUpload, UserHoliday


def _parse_holiday_pdf(file_bytes):
    """Parse a PDF file and extract holiday rows using pdfplumber."""
    try:
        import pdfplumber
        import io
        import re
        rows = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            # Try to find year in text
            full_text = ""
            for page in pdf.pages:
                full_text += (page.extract_text() or "") + "\n"
            year_match = re.search(r'\b(20\d{2})\b', full_text)
            doc_year = int(year_match.group(1)) if year_match else None

            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    if not table: continue
                    header = [str(c).lower().strip() if c else '' for c in table[0]]
                    for row in table[1:]:
                        if row and any(cell for cell in row):
                            rows.append({'raw': row, 'header': header, 'doc_year': doc_year})
        return rows
    except Exception as e:
        raise ValueError(f"PDF parsing failed: {e}")


def _parse_holiday_docx(file_bytes):
    """Parse a DOCX file and extract holiday rows with correct paragraph context."""
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        import io, re
        
        doc = Document(io.BytesIO(file_bytes))
        rows = []
        
        # Capture year from full text
        full_text = "\n".join([p.text for p in doc.paragraphs])
        year_match = re.search(r'\b(20\d{2})\b', full_text)
        doc_year = int(year_match.group(1)) if year_match else None

        # Helper to iterate elements in order
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P

        def iter_block_items(parent):
            # docx.Document is a factory; the actual class is docx.document.Document
            if hasattr(parent, 'element') and hasattr(parent.element, 'body'):
                parent_elm = parent.element.body
            else:
                parent_elm = parent._element
            
            for child in parent_elm.iterchildren():
                tag = child.tag.lower()
                if tag.endswith('}p'):
                    yield Paragraph(child, parent)
                elif tag.endswith('}tbl'):
                    yield Table(child, parent)

        current_context = ""
        for item in iter_block_items(doc):
            if isinstance(item, Paragraph):
                txt = item.text.strip()
                if txt and len(txt) < 200:
                    current_context = txt
            elif isinstance(item, Table):
                if not item.rows: continue
                is_optional_table = "optional" in current_context.lower()
                header = [cell.text.lower().strip() for cell in item.rows[0].cells]
                for row in item.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        rows.append({
                            'raw': cells, 
                            'header': header, 
                            'doc_year': doc_year,
                            'is_optional_context': is_optional_table,
                            'context_text': current_context
                        })

        if not rows:
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    rows.append({'raw': t.split(), 'header': None, 'doc_year': doc_year})
        return rows
    except Exception as e:
        raise ValueError(f"DOCX parsing failed: {e}")


def _parse_holiday_excel(file_bytes, file_ext):
    """Parse an Excel/CSV file using pandas."""
    try:
        import pandas as pd
        import io
        import re
        if file_ext in ['.xlsx', '.xls']:
            df = pd.read_excel(io.BytesIO(file_bytes), dtype=str)
        else:
            df = pd.read_csv(io.BytesIO(file_bytes), dtype=str)
        df.columns = [str(c).lower().strip() for c in df.columns]
        
        all_text = " ".join(df.astype(str).values.flatten())
        year_match = re.search(r'\b(20\d{2})\b', all_text)
        doc_year = int(year_match.group(1)) if year_match else None
        
        rows = []
        for _, row in df.iterrows():
            rows.append({'raw': list(row.values), 'header': list(df.columns), 'doc_year': doc_year})
        return rows
    except Exception as e:
        raise ValueError(f"Excel/CSV parsing failed: {e}")


def _parse_holiday_txt(file_bytes):
    """Parse a plain-text file into rows."""
    try:
        text = file_bytes.decode('utf-8', errors='replace')
        rows = []
        for line in text.split('\n'):
            line = line.strip()
            if line and not line.startswith('#'):
                parts = [p.strip() for p in line.replace('\t', ',').split(',')]
                rows.append({'raw': parts, 'header': None})
        return rows
    except Exception as e:
        raise ValueError(f"TXT parsing failed: {e}")


def _normalize_rows(raw_rows):
    """Normalise parsed rows into Holiday objects."""
    import re
    from datetime import datetime as dt
    from django.utils import timezone

    current_year = timezone.now().year

    def clean_date_str(s):
        if not s: return ""
        # Remove ordinal suffixes (1st -> 1, 22nd -> 22)
        s = re.sub(r'(\d+)(st|nd|rd|th)', r'\1', str(s), flags=re.IGNORECASE)
        # Remove special characters like dots or commas
        s = re.sub(r'[,.\(\)]', ' ', s)
        return s.strip()

    def smart_parse_date(s, year=None, month_hint=None):
        s = clean_date_str(s)
        if not s or s.lower() in ('nan', 'none', ''):
            return None
        
        # Heuristic: If it's just a bare number (1-31), it's likely a day within a month
        is_bare_number = re.fullmatch(r'\d{1,2}', s)
        if is_bare_number:
            if not month_hint: return None
            s = f"{s} {month_hint}"

        has_year = re.search(r'\d{4}', s)
        has_digits = re.search(r'\d+', s)
        has_month_name = any(m in s.lower() for m in ['jan','feb','mar','apr','may','jun','jul','aug','sep','oct','nov','dec'])
        
        # Prefer strings that actually look like dates (have digits and month or are complex)
        if not has_digits: return None

        formats = [
            '%d %B %Y', '%d %b %Y', '%B %d %Y', '%b %d %Y',
            '%d-%m-%Y', '%d/%m/%Y', '%Y-%m-%d',
            '%d %m %Y', '%Y %m %d',
            '%d %B', '%B %d', '%b %d', '%d %b'
        ]
        
        target_year = year or current_year
        valid_date = None

        for fmt in formats:
            try:
                if '%' not in fmt or 'Y' not in fmt:
                    d = dt.strptime(f"{s} {target_year}", f"{fmt} %Y").date()
                else:
                    d = dt.strptime(s, fmt).date()
                valid_date = d
                break
            except ValueError:
                continue

        if not valid_date:
            try:
                from dateutil import parser as dutil
                default_dt = dt(target_year, 1, 1)
                res = dutil.parse(s, default=default_dt, dayfirst=True).date()
                if not has_digits: return None
                valid_date = res
            except Exception:
                pass
        
        if valid_date and year and valid_date.year != year:
            valid_date = valid_date.replace(year=year)
            
        return valid_date

    holidays = []
    seen_keys = set()
    MONTH_NAMES = ['january','february','march','april','may','june','july','august','september','october','november','december']

    for item in raw_rows:
        raw_data = item.get('raw', [])
        header = item.get('header') or []
        doc_year = item.get('doc_year') or current_year
        # Handle different item keys for is_optional
        is_optional_context = item.get('is_optional', item.get('is_optional_context', False))

        if not raw_data: continue
        
        # Decide if raw_data is a single row or a list of rows
        if isinstance(raw_data, list) and len(raw_data) > 0 and not isinstance(raw_data[0], (list, tuple)):
            rows_to_process = [raw_data]
        elif isinstance(raw_data, list):
            rows_to_process = raw_data
        else:
            continue

        header_text = ' '.join(str(h).lower() for h in header)
        context_is_optional = is_optional_context or any(w in header_text for w in ['optional', 'restricted', 'rh'])

        for raw in rows_to_process:
            if not isinstance(raw, (list, tuple)): continue
            row_text = ' '.join(str(c).lower() for c in raw if c)
            if not row_text: continue

            month_hint = None
            for m in MONTH_NAMES:
                if m in row_text or m[:3] in row_text:
                    month_hint = m
                    break
            
            name = None
            date_val = None
            is_optional = context_is_optional or any(w in row_text for w in ['optional', 'restricted', 'rh'])
            is_working = any(w in row_text for w in ['working day', 'non-holiday', 'work day'])
            
            # Step 1: Specific Column Mapping
            col_map = {str(h).lower().strip(): i for i, h in enumerate(header)}
            def get_col(*keys):
                for k in keys:
                    for hk, idx in col_map.items():
                        if k in hk and idx < len(raw):
                            return str(raw[idx]).strip()
                return None

            h_name = get_col('holiday', 'festival', 'occasion', 'name')
            h_date = get_col('date', 'on_date')
            h_month = get_col('month')
            
            if h_date:
                date_val = smart_parse_date(h_date, year=doc_year, month_hint=(h_month or month_hint))
            if h_name:
                name = h_name

            # Step 2: Advanced Row Scanning
            potential_dates = []
            potential_names = []

            for i, cell in enumerate(raw):
                cs = str(cell).strip()
                if not cs or cs.lower() in ('nan', ''): continue
                
                # Rule out simple serial numbers (column 0, numeric, small)
                if i == 0 and re.fullmatch(r'\d{1,2}', cs):
                    continue
                
                # Try as date
                d = smart_parse_date(cs, year=doc_year, month_hint=month_hint)
                if d:
                    priority = 0
                    if any(m in cs.lower() for m in ['jan','feb','mar','apr','may','jun','jul','aug','sep','oct','nov','dec']):
                        priority += 2
                    if not re.fullmatch(r'\d{1,2}', cs):
                        priority += 1
                    potential_dates.append((d, priority))
                
                # Try as name
                cs_low = cs.lower()
                is_month = cs_low in MONTH_NAMES or any(m[:3] == cs_low for m in MONTH_NAMES if len(cs_low)==3)
                if not is_month and not any(wd in cs_low for wd in ['monday','tuesday','wednesday','thursday','friday','saturday','sunday']):
                    if len(cs) > 2 and not re.search(r'\d', cs):
                        potential_names.append(cs)
                    elif len(cs) > 3 and not re.match(r'^\d+$', cs) and not smart_parse_date(cs, year=doc_year):
                        potential_names.append(cs)

            if not date_val and potential_dates:
                potential_dates.sort(key=lambda x: x[1], reverse=True)
                date_val = potential_dates[0][0]
            
            if not name and potential_names:
                name_candidates = [n for n in potential_names if n.lower() not in MONTH_NAMES]
                if name_candidates:
                    name = max(name_candidates, key=len)

            if not name or not date_val: continue

            # Cleanup
            name = re.sub(r'^\d+[\.\s]*-?\s*', '', name).strip()
            if any(x in name.lower() for x in ['holiday', 'festival', 'sr no']): continue
            if name.lower() in MONTH_NAMES: continue

            import calendar as cal_mod
            day_str = cal_mod.day_name[date_val.weekday()]

            key = (name.lower(), str(date_val))
            if key in seen_keys: continue
            seen_keys.add(key)

            holidays.append({
                'name': name[:200],
                'date': str(date_val),
                'day': day_str,
                'is_optional': is_optional,
                'is_working_day': is_working,
                'year': date_val.year,
            })

    return sorted(holidays, key=lambda h: (h['date'], h['name']))



@api_view(['POST'])
@require_gated_token_api
@parser_classes([MultiPartParser, FormParser])
def holiday_upload_parse(request):
    """
    Admin uploads a holiday document.
    Parses it and returns a preview — no DB save yet.
    """
    user_id = request.data.get('user_id')
    employee = Employee.objects.filter(id=user_id, role='admin').first()
    if not employee:
        return Response({'success': False, 'message': 'Admin access required.'}, status=403)

    uploaded_file = request.FILES.get('file')
    if not uploaded_file:
        return Response({'success': False, 'message': 'No file provided.'}, status=400)

    ext = os.path.splitext(uploaded_file.name)[1].lower()
    allowed = ['.pdf', '.docx', '.xlsx', '.xls', '.csv', '.txt']
    if ext not in allowed:
        return Response({
            'success': False,
            'message': f'Unsupported format. Allowed: {", ".join(allowed)}'
        }, status=400)

    try:
        if not uploaded_file:
            return Response({'success': False, 'message': 'No file provided.'}, status=400)

        file_bytes = uploaded_file.read()
        raw_rows = []

        if ext == '.pdf':
            raw_rows = _parse_holiday_pdf(file_bytes)
        elif ext == '.docx':
            raw_rows = _parse_holiday_docx(file_bytes)
        elif ext in ['.xlsx', '.xls', '.csv']:
            raw_rows = _parse_holiday_excel(file_bytes, ext)
        else:  # .txt
            raw_rows = _parse_holiday_txt(file_bytes)

        holidays = _normalize_rows(raw_rows)

        # Create HolidayUpload audit record
        upload_obj = HolidayUpload.objects.create(
            file_name=uploaded_file.name[:255],
            uploaded_by=employee,
            status='parsed',
            parsed_count=len(holidays),
        )
        
        # Save file to media storage
        try:
            uploaded_file.seek(0)
            upload_obj.file.save(uploaded_file.name, uploaded_file, save=True)
        except Exception as file_err:
            print(f"Warning: Could not save holiday file to storage: {file_err}")

        return Response({
            'success': True,
            'upload_id': upload_obj.id,
            'file_name': uploaded_file.name,
            'parsed_count': len(holidays),
            'holidays': holidays,
        })

    except Exception as e:
        import traceback
        error_msg = f"Holiday Parse Error: {str(e)}"
        print(error_msg)
        print(traceback.format_exc())
        return Response({
            'success': False, 
            'message': error_msg,
            'debug_trace': traceback.format_exc() if settings.DEBUG else None
        }, status=500)


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def holiday_save(request):
    """
    Admin approves parsed holidays and saves them to DB.
    Accepts { user_id, upload_id, holidays: [...] }
    Handles duplicate-date conflicts gracefully.
    """
    user_id = request.data.get('user_id')
    employee = Employee.objects.filter(id=user_id, role='admin').first()
    if not employee:
        return Response({'success': False, 'message': 'Admin access required.'}, status=403)

    upload_id = request.data.get('upload_id')
    holidays_data = request.data.get('holidays', [])

    created = 0
    updated = 0
    skipped = 0
    errors = []

    from datetime import datetime as dt
    for h in holidays_data:
        try:
            date_val = dt.strptime(h['date'], '%Y-%m-%d').date()
            obj, created_flag = Holiday.objects.update_or_create(
                date=date_val,
                defaults={
                    'name': h.get('name', 'Holiday')[:200],
                    'day': h.get('day', ''),
                    'is_optional': bool(h.get('is_optional', False)),
                    'is_working_day': bool(h.get('is_working_day', False)),
                    'year': date_val.year,
                    'description': h.get('description', ''),
                }
            )
            if created_flag:
                created += 1
            else:
                updated += 1
        except Exception as e:
            skipped += 1
            errors.append(str(e))

    # mark upload as approved
    if upload_id:
        HolidayUpload.objects.filter(id=upload_id).update(
            status='approved',
            saved_count=created + updated
        )

    # Audit log notification to admin
    Notification.objects.create(
        user=employee,
        type='holiday_upload',
        message=f'Holiday list saved: {created} added, {updated} updated, {skipped} skipped.',
    )

    return Response({
        'success': True,
        'created': created,
        'updated': updated,
        'skipped': skipped,
        'errors': errors[:5],
    })


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def update_holiday(request):
    """
    Update an existing holiday. Admin only.
    """
    user_id = request.data.get('user_id')
    employee = Employee.objects.filter(id=user_id, role='admin').first()
    if not employee:
        return Response({'success': False, 'message': 'Admin access required.'}, status=403)

    holiday_id = request.data.get('holiday_id')
    name = request.data.get('name')
    date_str = request.data.get('date')
    day_str = request.data.get('day')
    is_optional = request.data.get('is_optional')
    description = request.data.get('description', '')

    if not holiday_id:
        return Response({'success': False, 'message': 'holiday_id is required.'}, status=400)

    try:
        holiday = Holiday.objects.get(id=holiday_id)
        if name: holiday.name = name
        if date_str:
            new_date = datetime.strptime(date_str, '%Y-%m-%d').date()
            if Holiday.objects.filter(date=new_date).exclude(id=holiday_id).exists():
                return Response({'success': False, 'message': f'A holiday already exists on {date_str}'}, status=400)
            holiday.date = new_date
            holiday.year = new_date.year
            
            if day_str:
                holiday.day = day_str
            else:
                # derive weekday if not provided
                import calendar as cal_mod
                holiday.day = cal_mod.day_name[new_date.weekday()]
        elif day_str:
            holiday.day = day_str
        
        if is_optional is not None:
            holiday.is_optional = bool(is_optional)
        
        is_working_day = request.data.get('is_working_day')
        if is_working_day is not None:
            holiday.is_working_day = bool(is_working_day)
        
        holiday.description = description
        holiday.save()

        return Response({'success': True, 'message': 'Holiday updated successfully.'})
    except Holiday.DoesNotExist:
        return Response({'success': False, 'message': 'Holiday not found.'}, status=404)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=500)


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def delete_holiday(request):
    """
    Delete a holiday. Admin only.
    """
    user_id = request.data.get('user_id')
    employee = Employee.objects.filter(id=user_id, role='admin').first()
    if not employee:
        return Response({'success': False, 'message': 'Admin access required.'}, status=403)

    holiday_id = request.data.get('holiday_id')
    if not holiday_id:
        return Response({'success': False, 'message': 'holiday_id is required.'}, status=400)

    try:
        holiday = Holiday.objects.get(id=holiday_id)
        holiday.delete()
        return Response({'success': True, 'message': 'Holiday deleted successfully.'})
    except Holiday.DoesNotExist:
        return Response({'success': False, 'message': 'Holiday not found.'}, status=404)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=500)



@api_view(['GET'])
@require_gated_token_api
def get_holidays(request):
    """
    Returns holidays for a given year (or current year).
    Optionally includes user's optional holiday selections.
    Query params: year, user_id
    """
    year = request.GET.get('year', timezone.now().year)
    user_id = request.GET.get('user_id')

    try:
        year = int(year)
    except (ValueError, TypeError):
        year = timezone.now().year

    holidays = Holiday.objects.filter(year=year)

    # get this user's selected optional holidays
    selected_ids = set()
    if user_id:
        selected_ids = set(
            UserHoliday.objects.filter(
                user_id=user_id, holiday__year=year
            ).values_list('holiday_id', flat=True)
        )

    data = []
    for h in holidays:
        data.append({
            'id': h.id,
            'name': h.name,
            'date': str(h.date),
            'day': h.day,
            'is_optional': h.is_optional,
            'is_working_day': h.is_working_day,
            'year': h.year,
            'description': h.description or '',
            'user_selected': h.id in selected_ids,
        })

    return Response({'success': True, 'holidays': data, 'year': year})


@api_view(['POST'])
@require_gated_token_api
def manage_date(request):
    """
    Directly manage a date's holiday/working status.
    Admin only.
    """
    user_id = request.data.get('user_id')
    employee = Employee.objects.filter(id=user_id, role='admin').first()
    if not employee:
        return Response({'success': False, 'message': 'Admin access required.'}, status=403)

    date_str = request.data.get('date')
    date_type = request.data.get('type') # 'working', 'holiday', 'optional'
    reason = request.data.get('reason', '')

    if not date_str:
        return Response({'success': False, 'message': 'Date is required.'}, status=400)

    try:
        from datetime import datetime
        target_date = datetime.strptime(date_str, '%Y-%m-%d').date()
        
        # Determine status flags
        is_optional = (date_type == 'optional')
        is_working = (date_type == 'working')
        
        existing_holiday = Holiday.objects.filter(date=target_date).first()
        
        msg = ""
        import calendar as cal_mod
        day_name = cal_mod.day_name[target_date.weekday()]

        if date_type == 'working':
            Holiday.objects.update_or_create(
                date=target_date,
                defaults={
                    'name': reason or "Working Day",
                    'is_optional': False,
                    'is_working_day': True,
                    'day': day_name,
                    'year': target_date.year
                }
            )
            msg = f"marked as Working Day: {reason or 'Regular'}"
        else:
            # Set as Holiday or Optional
            Holiday.objects.update_or_create(
                date=target_date,
                defaults={
                    'name': reason or ('Optional Holiday' if is_optional else 'Public Holiday'),
                    'is_optional': is_optional,
                    'is_working_day': False,
                    'day': day_name,
                    'year': target_date.year
                }
            )
            type_label = "Optional Holiday" if is_optional else "Mandatory Holiday"
            msg = f"set as {type_label}: {reason or 'Scheduled'}"

        # Notify EVERYONE
        all_employees = Employee.objects.filter(is_active=True)
        notifications = []
        for emp in all_employees:
            notifications.append(Notification(
                user=emp,
                type='holiday_update',
                message=f"Admin updated calendar for {date_str}: Now {msg}."
            ))
        Notification.objects.bulk_create(notifications)

        return Response({'success': True, 'message': f"Date {date_str} {msg}."})
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=500)


@api_view(['POST'])
@require_gated_token_api
@parser_classes([JSONParser])
def select_optional_holiday(request):
    """
    Employee selects / deselects an optional holiday.
    Enforces max 2 optional selections per year.
    Body: { user_id, holiday_id, action: 'select'|'deselect' }
    """
    MAX_OPTIONAL = 2

    user_id = request.data.get('user_id')
    holiday_id = request.data.get('holiday_id')
    action = request.data.get('action', 'select')

    if not user_id or not holiday_id:
        return Response({'success': False, 'message': 'user_id and holiday_id are required.'}, status=400)

    try:
        employee = Employee.objects.get(id=user_id)
        holiday = Holiday.objects.get(id=holiday_id, is_optional=True)
    except Employee.DoesNotExist:
        return Response({'success': False, 'message': 'Employee not found.'}, status=404)
    except Holiday.DoesNotExist:
        return Response({'success': False, 'message': 'Optional holiday not found.'}, status=404)

    if action == 'deselect':
        UserHoliday.objects.filter(user=employee, holiday=holiday).delete()
        return Response({'success': True, 'message': 'Holiday deselected.'})

    # Check limit
    existing_count = UserHoliday.objects.filter(
        user=employee, holiday__year=holiday.year
    ).count()
    if existing_count >= MAX_OPTIONAL:
        return Response({
            'success': False,
            'message': f'You can select at most {MAX_OPTIONAL} optional holidays per year.'
        }, status=400)

    uh, created = UserHoliday.objects.get_or_create(user=employee, holiday=holiday)
    if not created:
        return Response({'success': False, 'message': 'Holiday already selected.'}, status=400)

    return Response({'success': True, 'message': 'Optional holiday selected successfully.'})


@api_view(['GET'])
@require_gated_token_api
def export_holidays_ics(request):
    """
    Export holidays as an ICS calendar file.
    Query params: year, user_id (if provided, includes user's optional selections)
    """
    year = int(request.GET.get('year', timezone.now().year))
    user_id = request.GET.get('user_id')

    try:
        from icalendar import Calendar, Event as ICSEvent
        from datetime import datetime as dt
        import pytz

        cal = Calendar()
        cal.add('prodid', '-//HanuAI Holiday Calendar//EN')
        cal.add('version', '2.0')
        cal.add('calscale', 'GREGORIAN')
        cal.add('x-wr-calname', f'Holidays {year}')

        holidays = Holiday.objects.filter(year=year)
        selected_ids = set()
        if user_id:
            selected_ids = set(
                UserHoliday.objects.filter(
                    user_id=user_id, holiday__year=year
                ).values_list('holiday_id', flat=True)
            )

        for h in holidays:
            # Skip optional holidays unless user selected them (or no user_id – export all)
            if h.is_optional and user_id and h.id not in selected_ids:
                continue

            event = ICSEvent()
            event.add('summary', h.name)
            event.add('dtstart', h.date)
            event.add('dtend', h.date)
            event.add('description', f"{'Optional' if h.is_optional else 'Holiday'}")
            cal.add_component(event)

        ics_bytes = cal.to_ical()
        response = HttpResponse(ics_bytes, content_type='text/calendar; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="holidays_{year}.ics"'
        return response

    except ImportError:
        return Response({'success': False, 'message': 'icalendar library not installed.'}, status=500)
    except Exception as e:
        return Response({'success': False, 'message': str(e)}, status=500)


@api_view(['GET'])
@require_gated_token_api
def holiday_upload_history(request):
    """
    Returns the last 20 holiday upload audit records (admin only).
    """
    user_id = request.GET.get('user_id')
    employee = Employee.objects.filter(id=user_id, role='admin').first()
    if not employee:
        return Response({'success': False, 'message': 'Admin access required.'}, status=403)

    uploads = HolidayUpload.objects.all()[:20]
    data = []
    for u in uploads:
        data.append({
            'id': u.id,
            'file_name': u.file_name,
            'uploaded_by': u.uploaded_by.name if u.uploaded_by else 'Unknown',
            'uploaded_at': u.uploaded_at.strftime('%Y-%m-%d %H:%M'),
            'status': u.status,
            'parsed_count': u.parsed_count,
            'saved_count': u.saved_count,
        })
    return Response({'success': True, 'uploads': data})
